#!/usr/bin/env python3
"""General non-modality CLI helpers and command handlers."""

from __future__ import annotations

import glob
import importlib.metadata
import importlib.util
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

from cli_anything.onlyoffice.core.command_registry import (
    CATEGORY_COUNTS,
    CLI_SCHEMA_VERSION,
    TOTAL_COMMANDS,
    VERSION,
    build_capability_metadata,
    build_help_payload,
    command_usage,
)
from cli_anything.onlyoffice.core.parse_utils import parse_float, parse_int


OPEN_COMPATIBILITY_NAMESPACES = {
    "document",
    "doc",
    "spreadsheet",
    "sheet",
    "workbook",
    "presentation",
    "slides",
    "slide",
    "pdf",
    "text",
    "file",
    "office",
}

ONLYOFFICE_DOCKER_CONTAINER = "onlyoffice-documentserver"
ONLYOFFICE_X2T_PATH = (
    "/var/www/onlyoffice/documentserver/server/FileConverter/bin/x2t"
)

PYTHON_REQUIREMENTS = [
    {
        "key": "python_docx",
        "package": "python-docx",
        "import_name": "docx",
        "minimum_version": "1.1.0",
        "purpose": "DOCX create/read/edit/format/preflight support",
    },
    {
        "key": "openpyxl",
        "package": "openpyxl",
        "import_name": "openpyxl",
        "minimum_version": "3.1.2",
        "purpose": "XLSX workbook, chart, CSV, and validation support",
    },
    {
        "key": "python_pptx",
        "package": "python-pptx",
        "import_name": "pptx",
        "minimum_version": "0.6.23",
        "purpose": "PPTX creation, editing, shape, notes, and media support",
    },
    {
        "key": "requests",
        "package": "requests",
        "import_name": "requests",
        "minimum_version": "2.31.0",
        "purpose": "DocumentServer HTTP client support",
    },
    {
        "key": "scipy",
        "package": "scipy",
        "import_name": "scipy",
        "minimum_version": "1.11.0",
        "purpose": "Spreadsheet statistical tests",
    },
    {
        "key": "rdflib",
        "package": "rdflib",
        "import_name": "rdflib",
        "minimum_version": "7.0.0",
        "purpose": "RDF graph parsing, querying, and serialization",
    },
    {
        "key": "lxml",
        "package": "lxml",
        "import_name": "lxml",
        "minimum_version": "4.9.0",
        "purpose": "OOXML/XML support",
    },
    {
        "key": "pymupdf",
        "package": "PyMuPDF",
        "import_name": "fitz",
        "minimum_version": "1.24.0",
        "purpose": "PDF native block reading, image extraction, and rendering",
    },
    {
        "key": "pillow",
        "package": "Pillow",
        "import_name": "PIL",
        "minimum_version": "10.0.0",
        "purpose": "Image decoding and format conversion",
    },
    {
        "key": "pyshacl",
        "package": "pyshacl",
        "import_name": "pyshacl",
        "minimum_version": "0.25.0",
        "purpose": "SHACL validation for rdf-validate",
    },
]


def _version_tuple(value: object) -> Tuple[int, ...]:
    parts = re.findall(r"\d+", str(value or ""))
    if not parts:
        return ()
    return tuple(int(part) for part in parts[:4])


def _version_ok(installed: Optional[str], minimum: Optional[str]) -> bool:
    if not installed or not minimum:
        return bool(installed)
    installed_tuple = _version_tuple(installed)
    minimum_tuple = _version_tuple(minimum)
    if not installed_tuple or not minimum_tuple:
        return True
    width = max(len(installed_tuple), len(minimum_tuple))
    return installed_tuple + (0,) * (width - len(installed_tuple)) >= minimum_tuple + (
        0,
    ) * (width - len(minimum_tuple))


def detect_python_requirements(
    *,
    requirements: Optional[List[Dict[str, str]]] = None,
    find_spec: Callable[[str], object] = importlib.util.find_spec,
    version_lookup: Callable[[str], str] = importlib.metadata.version,
) -> List[Dict[str, object]]:
    """Return install-time status for required Python distributions."""
    results = []
    for req in requirements or PYTHON_REQUIREMENTS:
        import_name = req["import_name"]
        package = req["package"]
        available = find_spec(import_name) is not None
        installed_version = None
        version_error = None
        try:
            installed_version = version_lookup(package)
        except importlib.metadata.PackageNotFoundError:
            if available:
                installed_version = None
        except Exception as exc:
            version_error = str(exc)
        minimum = req.get("minimum_version")
        version_satisfied = bool(available and _version_ok(installed_version, minimum))
        status = "pass" if available and version_satisfied else "fail"
        entry = {
            "key": req["key"],
            "package": package,
            "import_name": import_name,
            "required": True,
            "available": available,
            "installed_version": installed_version,
            "minimum_version": minimum,
            "version_satisfied": version_satisfied,
            "status": status,
            "purpose": req.get("purpose", ""),
            "install_requirement": f"{package}>={minimum}" if minimum else package,
        }
        if version_error:
            entry["version_error"] = version_error
        results.append(entry)
    return results


def build_installation_check(
    *,
    doc_server=None,
    docx_available: bool,
    openpyxl_available: bool,
    pptx_available: bool,
    live_smoke: bool = False,
    conversion_detector: Optional[Callable[[], Dict[str, object]]] = None,
    python_detector: Callable[[], List[Dict[str, object]]] = detect_python_requirements,
) -> Dict[str, object]:
    """Build a strict post-clone/post-pull install readiness report."""
    if conversion_detector is None:
        conversion_detector = detect_conversion_capability
    python_dependencies = python_detector()
    conversion = conversion_detector()
    docker_info = conversion.get("docker", {})
    x2t_info = conversion.get("x2t", {})
    external_dependencies = [
        {
            "key": "docker",
            "label": "Docker CLI",
            "required": True,
            "available": bool(docker_info.get("available")),
            "path": docker_info.get("path"),
            "status": "pass" if docker_info.get("available") else "fail",
            "purpose": "Runs the OnlyOffice DocumentServer conversion probe.",
            "install_hint": "Install Docker and ensure the current user can run docker commands.",
        },
        {
            "key": "onlyoffice_x2t",
            "label": "OnlyOffice DocumentServer x2t",
            "required": True,
            "available": x2t_info.get("available") is True,
            "container": x2t_info.get("container"),
            "path": x2t_info.get("path"),
            "checked": x2t_info.get("checked"),
            "status": "pass" if x2t_info.get("available") is True else "fail",
            "purpose": "Converts Office documents to PDF/images for render-aware audits.",
            "install_hint": "Start the onlyoffice-documentserver container with x2t available.",
        },
    ]
    missing_python = [
        item["install_requirement"]
        for item in python_dependencies
        if not item.get("available")
    ]
    outdated_python = [
        item["install_requirement"]
        for item in python_dependencies
        if item.get("available") and not item.get("version_satisfied")
    ]
    missing_external = [
        item["key"] for item in external_dependencies if not item.get("available")
    ]
    python_ok = not missing_python and not outdated_python
    external_ok = not missing_external
    install_ready = bool(python_ok and external_ok and doc_server is not None)
    live_smoke_result = None
    if live_smoke:
        live_smoke_result = run_live_docx_pdf_smoke(doc_server) if install_ready else {
            "success": False,
            "skipped": True,
            "reason": "Base dependency checks did not pass.",
        }
        install_ready = install_ready and bool(live_smoke_result.get("success"))
    install_hints = []
    if missing_python or outdated_python:
        install_hints.append(
            "Run the project venv interpreter and reinstall after pulling: "
            "python -m pip install -e ."
        )
    if "docker" in missing_external:
        install_hints.append("Install Docker and confirm `docker ps` works for this user.")
    if "onlyoffice_x2t" in missing_external:
        install_hints.append(
            "Start or reinstall the onlyoffice-documentserver container so x2t exists."
        )
    if doc_server is None:
        install_hints.append("OnlyOffice client failed to initialize; inspect earlier import errors.")

    return {
        "success": install_ready,
        "schema_version": CLI_SCHEMA_VERSION,
        "version": VERSION,
        "python": sys.executable,
        "mode": "setup-check",
        "install_ready": install_ready,
        "live_smoke_requested": live_smoke,
        "live_smoke": live_smoke_result,
        "python_dependencies_ok": python_ok,
        "external_dependencies_ok": external_ok,
        "client_available": doc_server is not None,
        "runtime_import_flags": {
            "python_docx": docx_available,
            "openpyxl": openpyxl_available,
            "python_pptx": pptx_available,
        },
        "python_dependencies": python_dependencies,
        "external_dependencies": external_dependencies,
        "conversion": conversion,
        "missing_python": missing_python,
        "outdated_python": outdated_python,
        "missing_external": missing_external,
        "install_hints": install_hints,
    }


def run_live_docx_pdf_smoke(doc_server) -> Dict[str, object]:
    """Run an optional real DOCX->PDF smoke through the installed converter."""
    if doc_server is None:
        return {"success": False, "error": "OnlyOffice client is unavailable."}
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        return {"success": False, "error": f"python-docx unavailable for smoke test: {exc}"}

    sentinel = "ONLYOFFICE_LIVE_SMOKE_SENTINEL"
    try:
        with tempfile.TemporaryDirectory(prefix="onlyoffice-live-smoke-") as tmpdir:
            docx_path = os.path.join(tmpdir, "smoke.docx")
            pdf_path = os.path.join(tmpdir, "smoke.pdf")
            doc = Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(sentinel)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            doc.save(docx_path)

            conversion = doc_server.doc_to_pdf(docx_path, output_path=pdf_path)
            pdf_exists = os.path.exists(pdf_path)
            pdf_header_ok = False
            if pdf_exists:
                with open(pdf_path, "rb") as handle:
                    pdf_header_ok = handle.read(5) == b"%PDF-"
            blocks = (
                doc_server.pdf_read_blocks(
                    pdf_path,
                    include_spans=True,
                    include_images=False,
                    include_empty=False,
                )
                if conversion.get("success") and pdf_header_ok
                else {"success": False, "error": "PDF conversion/header check failed."}
            )
            text = ""
            span_has_font = False
            span_has_size = False
            if blocks.get("success"):
                for page in blocks.get("pages", []):
                    for block in page.get("blocks", []):
                        text += "\n" + str(block.get("text") or "")
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                if span.get("font"):
                                    span_has_font = True
                                if span.get("size") is not None:
                                    span_has_size = True
            checks = {
                "conversion_success": bool(conversion.get("success")),
                "pdf_exists": pdf_exists,
                "pdf_header_ok": pdf_header_ok,
                "pdf_blocks_read": bool(blocks.get("success")),
                "sentinel_text_present": sentinel in text,
                "span_font_metadata_present": span_has_font,
                "span_size_metadata_present": span_has_size,
            }
            return {
                "success": all(checks.values()),
                "checks": checks,
                "conversion": conversion,
                "blocks_summary": {
                    "success": blocks.get("success"),
                    "total_pages": blocks.get("total_pages"),
                    "pages_scanned": blocks.get("pages_scanned"),
                    "span_count": blocks.get("span_count"),
                    "error": blocks.get("error"),
                },
            }
    except Exception as exc:
        return {"success": False, "error": str(exc)}


def detect_conversion_capability(
    *,
    which: Callable[[str], Optional[str]] = shutil.which,
    run: Callable[..., subprocess.CompletedProcess] = subprocess.run,
    timeout: float = 2.0,
) -> Dict[str, object]:
    """Return structured, failure-tolerant Docker/x2t conversion status."""
    docker_path = which("docker")
    status: Dict[str, object] = {
        "available": False,
        "docker": {
            "available": bool(docker_path),
            "path": docker_path,
        },
        "x2t": {
            "available": None,
            "container": ONLYOFFICE_DOCKER_CONTAINER,
            "path": ONLYOFFICE_X2T_PATH,
            "checked": False,
        },
    }
    if not docker_path:
        status["x2t"]["error"] = "docker executable not found"
        return status

    try:
        result = run(
            [
                docker_path,
                "exec",
                ONLYOFFICE_DOCKER_CONTAINER,
                "test",
                "-x",
                ONLYOFFICE_X2T_PATH,
            ],
            capture_output=True,
            timeout=timeout,
        )
    except subprocess.TimeoutExpired:
        status["x2t"].update({"available": False, "checked": True, "error": "probe timed out"})
        return status
    except Exception as exc:
        status["x2t"].update({"available": False, "checked": True, "error": str(exc)})
        return status

    x2t_available = result.returncode == 0
    status["x2t"].update(
        {
            "available": x2t_available,
            "checked": True,
            "exit_code": result.returncode,
        }
    )
    if not x2t_available:
        stderr_raw = result.stderr or b""
        stderr = (
            stderr_raw.decode("utf-8", errors="replace")
            if isinstance(stderr_raw, bytes)
            else str(stderr_raw)
        ).strip()
        if stderr:
            status["x2t"]["error"] = stderr
    status["available"] = x2t_available
    return status


def detect_onlyoffice_file_type(file_path: str) -> str:
    """Best-effort type detection for files OnlyOffice commonly opens."""
    ext = Path(file_path).suffix.lower()
    if ext in {".xlsx", ".xls", ".xlsm", ".xltx", ".xltm", ".ods", ".csv", ".tsv"}:
        return "spreadsheet"
    if ext in {".pptx", ".ppt", ".pptm", ".ppsx", ".odp"}:
        return "presentation"
    if ext == ".pdf":
        return "pdf"
    if ext in {
        ".docx",
        ".doc",
        ".docm",
        ".dotx",
        ".odt",
        ".rtf",
        ".txt",
        ".md",
        ".html",
        ".htm",
        ".epub",
    }:
        return "document"
    return "file"


def normalize_command_alias(command: str) -> Tuple[str, Optional[Dict[str, str]]]:
    """Map agent-style dotted compatibility commands onto the canonical CLI surface."""
    if "." not in command:
        return command, None
    namespace, action = command.rsplit(".", 1)
    if action not in {"open", "watch", "info"}:
        return command, None
    if namespace not in OPEN_COMPATIBILITY_NAMESPACES:
        return command, None
    return action, {
        "requested_command": command,
        "resolved_command": action,
        "command_namespace": namespace,
    }


def apply_alias_meta(
    result: Dict[str, object], alias_meta: Optional[Dict[str, str]]
) -> Dict[str, object]:
    """Attach compatibility alias metadata to a result payload."""
    if alias_meta:
        merged = dict(result)
        merged.update(alias_meta)
        return merged
    return result


def cmd_list(
    *,
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    docx_available: bool,
    openpyxl_available: bool,
) -> None:
    """List recent documents, spreadsheets, and presentations."""
    files = []
    patterns = [
        "~/Documents/*.docx",
        "~/Documents/*.xlsx",
        "~/Documents/*.pptx",
        "~/Documents/*.txt",
        "~/Downloads/*.docx",
        "~/Downloads/*.xlsx",
        "~/Downloads/*.pptx",
        "~/Downloads/*.csv",
    ]

    for pattern in patterns:
        for filepath in glob.glob(os.path.expanduser(pattern)):
            try:
                stat = os.stat(filepath)
                ext = Path(filepath).suffix.lower()
                if ext in [".xlsx", ".csv"]:
                    ftype = "spreadsheet"
                elif ext == ".pptx":
                    ftype = "presentation"
                else:
                    ftype = "document"
                files.append(
                    {
                        "path": filepath,
                        "name": os.path.basename(filepath),
                        "type": ftype,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "size": stat.st_size,
                    }
                )
            except OSError:
                continue

    files.sort(key=lambda item: item["modified"], reverse=True)
    print_result(
        {
            "success": True,
            "count": len(files),
            "files": files[:20],
            "python_docx": docx_available,
            "openpyxl": openpyxl_available,
        },
        json_output,
    )


def cmd_open(
    file_path: str,
    mode: str = "gui",
    *,
    json_output: bool,
    alias_meta: Optional[Dict[str, str]],
    print_result: Callable[[dict, bool], None],
) -> None:
    """Open a document in OnlyOffice GUI or web viewer."""
    try:
        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            print_result(
                apply_alias_meta(
                    {"success": False, "error": f"File not found: {file_path}"},
                    alias_meta,
                ),
                json_output,
            )
            return

        detected_type = detect_onlyoffice_file_type(abs_path)
        if mode == "gui":
            subprocess.Popen(
                ["onlyoffice-desktopeditors", abs_path], start_new_session=True
            )
            result = {
                "success": True,
                "file": abs_path,
                "mode": "gui",
                "detected_type": detected_type,
                "message": "Opened in OnlyOffice Desktop Editors",
            }
        elif mode == "web":
            result = {
                "success": True,
                "file": abs_path,
                "mode": "web",
                "detected_type": detected_type,
                "message": "Document Server URL: http://localhost:8080 (configure shared folder for web access)",
                "note": "For web viewing, copy file to Document Server documents folder",
            }
        else:
            result = {
                "success": False,
                "error": f"Unknown mode: {mode}. Use 'gui' or 'web'",
            }

        print_result(apply_alias_meta(result, alias_meta), json_output)
    except Exception as exc:
        print_result(
            apply_alias_meta({"success": False, "error": str(exc)}, alias_meta),
            json_output,
        )


def cmd_watch(
    file_path: str,
    mode: str = "gui",
    *,
    json_output: bool,
    alias_meta: Optional[Dict[str, str]],
    print_result: Callable[[dict, bool], None],
) -> None:
    """Watch a file and auto-open it in GUI for real-time viewing."""
    try:
        import time

        abs_path = os.path.abspath(file_path)
        if not os.path.exists(abs_path):
            print_result(
                apply_alias_meta(
                    {"success": False, "error": f"File not found: {file_path}"},
                    alias_meta,
                ),
                json_output,
            )
            return

        start_payload = apply_alias_meta(
            {
                "success": True,
                "watching": abs_path,
                "mode": mode,
                "detected_type": detect_onlyoffice_file_type(abs_path),
                "message": f"Started watching {abs_path}. Press Ctrl+C to stop.",
            },
            alias_meta,
        )

        if json_output:
            print(json.dumps(start_payload, indent=2))
        else:
            print(f"Watching: {abs_path}")
            print(f"Mode: {mode}")
            print("Press Ctrl+C to stop")

        if mode == "gui":
            subprocess.Popen(
                ["onlyoffice-desktopeditors", abs_path], start_new_session=True
            )

        last_mtime = os.stat(abs_path).st_mtime
        try:
            while True:
                time.sleep(1)
                if os.path.exists(abs_path):
                    current_mtime = os.stat(abs_path).st_mtime
                    if current_mtime != last_mtime:
                        last_mtime = current_mtime
                        if not json_output:
                            print(
                                f"  [Updated] File modified at {datetime.now().strftime('%H:%M:%S')}"
                            )
        except KeyboardInterrupt:
            if not json_output:
                print("\nWatch stopped.")
            print_result(
                apply_alias_meta(
                    {
                        "success": True,
                        "watching": abs_path,
                        "stopped": True,
                        "detected_type": detect_onlyoffice_file_type(abs_path),
                    },
                    alias_meta,
                ),
                json_output,
            )
    except Exception as exc:
        print_result(
            apply_alias_meta({"success": False, "error": str(exc)}, alias_meta),
            json_output,
        )


def cmd_info(
    file_path: str,
    *,
    json_output: bool,
    alias_meta: Optional[Dict[str, str]],
    print_result: Callable[[dict, bool], None],
    doc_server,
) -> None:
    """Get file information."""
    result = (
        doc_server.get_document_info(file_path)
        if doc_server
        else {"success": False, "error": "Client not available"}
    )
    print_result(apply_alias_meta(result, alias_meta), json_output)


def cmd_editor_session(
    file_path: str,
    *,
    open_if_needed: bool,
    wait_seconds: float,
    activate: bool,
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    doc_server,
) -> None:
    """Locate or open a native OnlyOffice desktop editor session."""
    result = (
        doc_server.editor_session(
            file_path,
            open_if_needed=open_if_needed,
            wait_seconds=wait_seconds,
            activate=activate,
        )
        if doc_server
        else {"success": False, "error": "Client not available"}
    )
    print_result(result, json_output)


def cmd_editor_capture(
    file_path: str,
    output_path: str,
    *,
    backend: str,
    open_if_needed: bool,
    page: Optional[int],
    cell_range: Optional[str],
    slide: Optional[int],
    zoom_reset: bool,
    zoom_in_steps: int,
    zoom_out_steps: int,
    crop: Optional[str],
    settle_ms: int,
    wait_seconds: float,
    dpi: int,
    fmt: Optional[str],
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    doc_server,
) -> None:
    """Capture a live editor viewport or rendered fallback image."""
    result = (
        doc_server.capture_editor_view(
            file_path,
            output_path,
            backend=backend,
            open_if_needed=open_if_needed,
            page=page,
            cell_range=cell_range,
            slide=slide,
            zoom_reset=zoom_reset,
            zoom_in_steps=zoom_in_steps,
            zoom_out_steps=zoom_out_steps,
            crop=crop,
            settle_ms=settle_ms,
            wait_seconds=wait_seconds,
            dpi=dpi,
            fmt=fmt,
        )
        if doc_server
        else {"success": False, "error": "Client not available"}
    )
    print_result(result, json_output)


def cmd_backup_list(
    file_path: str,
    *,
    limit: int,
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    doc_server,
) -> None:
    """List backups for a target file."""
    if not doc_server:
        print_result({"success": False, "error": "Client not available"}, json_output)
        return
    result = doc_server.list_backups(file_path, limit=parse_int(limit, "--limit"))
    print_result(result, json_output)


def cmd_backup_prune(
    *,
    file_path: Optional[str],
    keep: int,
    days: Optional[int],
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    doc_server,
) -> None:
    """Prune backups by retention rules."""
    if not doc_server:
        print_result({"success": False, "error": "Client not available"}, json_output)
        return
    result = doc_server.prune_backups(
        file_path=file_path,
        keep=parse_int(keep, "--keep"),
        older_than_days=parse_int(days, "--days") if days is not None else None,
    )
    print_result(result, json_output)


def cmd_backup_restore(
    file_path: str,
    *,
    backup: Optional[str],
    latest: bool,
    dry_run: bool,
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    doc_server,
) -> None:
    """Restore a file from backup."""
    if not doc_server:
        print_result({"success": False, "error": "Client not available"}, json_output)
        return
    result = doc_server.restore_backup(
        file_path=file_path,
        backup=backup,
        latest=latest,
        dry_run=dry_run,
    )
    print_result(result, json_output)


def cmd_status(
    *,
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    doc_server,
    docx_available: bool,
    openpyxl_available: bool,
    pptx_available: bool,
) -> None:
    """Check installation status."""
    try:
        import rdflib

        rdflib_available = True
        rdflib_version = rdflib.__version__
    except ImportError:
        rdflib_available = False
        rdflib_version = None

    try:
        import pyshacl

        shacl_available = True
    except ImportError:
        shacl_available = False

    conversion = detect_conversion_capability()
    conversion_available = bool(conversion.get("available"))
    docker_info = conversion.get("docker", {})
    x2t_info = conversion.get("x2t", {})
    dependencies = {
        "python_docx": docx_available,
        "openpyxl": openpyxl_available,
        "python_pptx": pptx_available,
        "rdflib": rdflib_available,
        "pyshacl": shacl_available,
        "docker": bool(docker_info.get("available")),
        "onlyoffice_x2t": x2t_info.get("available") is True,
    }
    install_check = build_installation_check(
        doc_server=doc_server,
        docx_available=docx_available,
        openpyxl_available=openpyxl_available,
        pptx_available=pptx_available,
        conversion_detector=lambda: conversion,
    )
    capabilities = {
        "docx_create": docx_available,
        "docx_read": docx_available,
        "docx_edit": docx_available,
        "docx_tables": docx_available,
        "docx_formatting": docx_available,
        "docx_submission": docx_available,
        "docx_references": docx_available,
        "xlsx_create": openpyxl_available,
        "xlsx_read": openpyxl_available,
        "xlsx_edit": openpyxl_available,
        "xlsx_formulas": openpyxl_available,
        "xlsx_charts": openpyxl_available,
        "xlsx_stats": openpyxl_available,
        "xlsx_csv": openpyxl_available,
        "pptx_create": pptx_available,
        "pptx_read": pptx_available,
        "pptx_edit": pptx_available,
        "pptx_notes": pptx_available,
        "rdf_create": rdflib_available,
        "rdf_query": rdflib_available,
        "rdf_validate": shacl_available,
        "office_to_pdf": conversion_available,
        "docx_to_pdf": docx_available and conversion_available,
        "xlsx_to_pdf": openpyxl_available and conversion_available,
        "pptx_to_pdf": pptx_available and conversion_available,
    }
    print_result(
        {
            "success": True,
            "schema_version": CLI_SCHEMA_VERSION,
            "version": VERSION,
            "python": sys.executable,
            "document_server": {
                "healthy": doc_server.check_health() if doc_server else False
            },
            "registry": {
                "version": VERSION,
                "schema_version": CLI_SCHEMA_VERSION,
                "total_commands": TOTAL_COMMANDS,
                "category_counts": dict(CATEGORY_COUNTS),
            },
            "python_docx": docx_available,
            "openpyxl": openpyxl_available,
            "python_pptx": pptx_available,
            "rdflib": rdflib_available,
            "rdflib_version": rdflib_version,
            "pyshacl": shacl_available,
            "conversion": conversion,
            "dependencies": dependencies,
            "install_check": {
                "install_ready": install_check["install_ready"],
                "python_dependencies_ok": install_check["python_dependencies_ok"],
                "external_dependencies_ok": install_check["external_dependencies_ok"],
                "client_available": install_check["client_available"],
                "missing_python": install_check["missing_python"],
                "outdated_python": install_check["outdated_python"],
                "missing_external": install_check["missing_external"],
                "install_hints": install_check["install_hints"],
            },
            "capabilities": capabilities,
            "capability_metadata": build_capability_metadata(
                {**dependencies, **capabilities}
            ),
            "total_commands": TOTAL_COMMANDS,
            "command_count": TOTAL_COMMANDS,
            "category_counts": dict(CATEGORY_COUNTS),
        },
        json_output,
    )


def cmd_setup_check(
    *,
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    doc_server,
    docx_available: bool,
    openpyxl_available: bool,
    pptx_available: bool,
    live_smoke: bool = False,
) -> None:
    """Strict install/update readiness check for freshly cloned or pulled checkouts."""
    print_result(
        build_installation_check(
            doc_server=doc_server,
            docx_available=docx_available,
            openpyxl_available=openpyxl_available,
            pptx_available=pptx_available,
            live_smoke=live_smoke,
        ),
        json_output,
    )


def cmd_help(
    *,
    json_output: bool,
    print_result: Callable[[dict, bool], None],
    docx_available: bool,
    openpyxl_available: bool,
    pptx_available: bool,
) -> None:
    """Show help."""
    try:
        import rdflib

        rdflib_available = True
    except ImportError:
        rdflib_available = False

    result = build_help_payload(
        {
            "python_docx": docx_available,
            "openpyxl": openpyxl_available,
            "python_pptx": pptx_available,
            "rdflib": rdflib_available,
        }
    )
    print_result(result, json_output)


def handle_general_command(
    command: str,
    raw_args: List[str],
    *,
    json_output: bool,
    alias_meta: Optional[Dict[str, str]],
    print_result: Callable[[dict, bool], None],
    doc_server,
    docx_available: bool,
    openpyxl_available: bool,
    pptx_available: bool,
) -> bool:
    """Handle general non-prefixed CLI commands and return True when recognised."""
    if command == "list":
        cmd_list(
            json_output=json_output,
            print_result=print_result,
            docx_available=docx_available,
            openpyxl_available=openpyxl_available,
        )
        return True

    if command == "open":
        if not raw_args:
            print_result({"success": False, "error": command_usage("open")}, json_output)
        else:
            mode = raw_args[1] if len(raw_args) > 1 else "gui"
            cmd_open(
                raw_args[0],
                mode,
                json_output=json_output,
                alias_meta=alias_meta,
                print_result=print_result,
            )
        return True

    if command == "watch":
        if not raw_args:
            print_result({"success": False, "error": command_usage("watch")}, json_output)
        else:
            mode = raw_args[1] if len(raw_args) > 1 else "gui"
            cmd_watch(
                raw_args[0],
                mode,
                json_output=json_output,
                alias_meta=alias_meta,
                print_result=print_result,
            )
        return True

    if command == "info":
        if not raw_args:
            print_result({"success": False, "error": command_usage("info")}, json_output)
        else:
            cmd_info(
                raw_args[0],
                json_output=json_output,
                alias_meta=alias_meta,
                print_result=print_result,
                doc_server=doc_server,
            )
        return True

    if command == "editor-session":
        if not raw_args:
            print_result(
                {"success": False, "error": command_usage("editor-session")},
                json_output,
            )
        else:
            open_if_needed = False
            wait_seconds = 10.0
            activate = False
            i = 1
            while i < len(raw_args):
                if raw_args[i] == "--open":
                    open_if_needed = True
                    i += 1
                elif raw_args[i] == "--wait" and i + 1 < len(raw_args):
                    wait_seconds = parse_float(raw_args[i + 1], "--wait")
                    i += 2
                elif raw_args[i] == "--activate":
                    activate = True
                    i += 1
                else:
                    i += 1
            cmd_editor_session(
                raw_args[0],
                open_if_needed=open_if_needed,
                wait_seconds=wait_seconds,
                activate=activate,
                json_output=json_output,
                print_result=print_result,
                doc_server=doc_server,
            )
        return True

    if command == "editor-capture":
        if len(raw_args) < 2:
            print_result(
                {"success": False, "error": command_usage("editor-capture")},
                json_output,
            )
        else:
            backend = "auto"
            open_if_needed = False
            page = None
            cell_range = None
            slide = None
            zoom_reset = False
            zoom_in_steps = 0
            zoom_out_steps = 0
            crop = None
            settle_ms = 800
            wait_seconds = 10.0
            dpi = 150
            fmt = None
            i = 2
            while i < len(raw_args):
                if raw_args[i] == "--backend" and i + 1 < len(raw_args):
                    backend = raw_args[i + 1]
                    i += 2
                elif raw_args[i] == "--open":
                    open_if_needed = True
                    i += 1
                elif raw_args[i] == "--page" and i + 1 < len(raw_args):
                    page = parse_int(raw_args[i + 1], "--page")
                    i += 2
                elif raw_args[i] == "--range" and i + 1 < len(raw_args):
                    cell_range = raw_args[i + 1]
                    i += 2
                elif raw_args[i] == "--slide" and i + 1 < len(raw_args):
                    slide = parse_int(raw_args[i + 1], "--slide")
                    i += 2
                elif raw_args[i] == "--zoom-reset":
                    zoom_reset = True
                    i += 1
                elif raw_args[i] == "--zoom-in" and i + 1 < len(raw_args):
                    zoom_in_steps = parse_int(raw_args[i + 1], "--zoom-in")
                    i += 2
                elif raw_args[i] == "--zoom-out" and i + 1 < len(raw_args):
                    zoom_out_steps = parse_int(raw_args[i + 1], "--zoom-out")
                    i += 2
                elif raw_args[i] == "--crop" and i + 1 < len(raw_args):
                    crop = raw_args[i + 1]
                    i += 2
                elif raw_args[i] == "--settle-ms" and i + 1 < len(raw_args):
                    settle_ms = parse_int(raw_args[i + 1], "--settle-ms")
                    i += 2
                elif raw_args[i] == "--wait" and i + 1 < len(raw_args):
                    wait_seconds = parse_float(raw_args[i + 1], "--wait")
                    i += 2
                elif raw_args[i] == "--dpi" and i + 1 < len(raw_args):
                    dpi = parse_int(raw_args[i + 1], "--dpi")
                    i += 2
                elif raw_args[i] == "--format" and i + 1 < len(raw_args):
                    fmt = raw_args[i + 1]
                    i += 2
                else:
                    i += 1
            cmd_editor_capture(
                raw_args[0],
                raw_args[1],
                backend=backend,
                open_if_needed=open_if_needed,
                page=page,
                cell_range=cell_range,
                slide=slide,
                zoom_reset=zoom_reset,
                zoom_in_steps=zoom_in_steps,
                zoom_out_steps=zoom_out_steps,
                crop=crop,
                settle_ms=settle_ms,
                wait_seconds=wait_seconds,
                dpi=dpi,
                fmt=fmt,
                json_output=json_output,
                print_result=print_result,
                doc_server=doc_server,
            )
        return True

    if command == "backup-list":
        if not raw_args:
            print_result(
                {"success": False, "error": command_usage("backup-list")},
                json_output,
            )
        else:
            file_path = None
            limit = 20
            i = 0
            while i < len(raw_args):
                if raw_args[i] == "--limit" and i + 1 < len(raw_args):
                    limit = parse_int(raw_args[i + 1], "--limit")
                    i += 2
                elif not raw_args[i].startswith("--") and file_path is None:
                    file_path = raw_args[i]
                    i += 1
                else:
                    i += 1
            if not file_path:
                print_result(
                    {"success": False, "error": command_usage("backup-list")},
                    json_output,
                )
            else:
                cmd_backup_list(
                    file_path,
                    limit=limit,
                    json_output=json_output,
                    print_result=print_result,
                    doc_server=doc_server,
                )
        return True

    if command == "backup-prune":
        file_path = None
        keep = 20
        days = None
        i = 0
        while i < len(raw_args):
            if raw_args[i] == "--file" and i + 1 < len(raw_args):
                file_path = raw_args[i + 1]
                i += 2
            elif raw_args[i] == "--keep" and i + 1 < len(raw_args):
                keep = parse_int(raw_args[i + 1], "--keep")
                i += 2
            elif raw_args[i] == "--days" and i + 1 < len(raw_args):
                days = parse_int(raw_args[i + 1], "--days")
                i += 2
            elif not raw_args[i].startswith("--") and file_path is None:
                file_path = raw_args[i]
                i += 1
            else:
                i += 1
        cmd_backup_prune(
            file_path=file_path,
            keep=keep,
            days=days,
            json_output=json_output,
            print_result=print_result,
            doc_server=doc_server,
        )
        return True

    if command == "backup-restore":
        if not raw_args:
            print_result(
                {"success": False, "error": command_usage("backup-restore")},
                json_output,
            )
        else:
            file_path = raw_args[0]
            backup = None
            latest = False
            dry_run = False
            i = 1
            while i < len(raw_args):
                if raw_args[i] == "--backup" and i + 1 < len(raw_args):
                    backup = raw_args[i + 1]
                    i += 2
                elif raw_args[i] == "--latest":
                    latest = True
                    i += 1
                elif raw_args[i] == "--dry-run":
                    dry_run = True
                    i += 1
                else:
                    i += 1
            cmd_backup_restore(
                file_path,
                backup=backup,
                latest=latest,
                dry_run=dry_run,
                json_output=json_output,
                print_result=print_result,
                doc_server=doc_server,
            )
        return True

    if command in {"setup-check", "update-check", "doctor"}:
        live_smoke = "--live-smoke" in raw_args or os.environ.get(
            "ONLYOFFICE_LIVE_SMOKE"
        ) == "1"
        cmd_setup_check(
            json_output=json_output,
            print_result=print_result,
            doc_server=doc_server,
            docx_available=docx_available,
            openpyxl_available=openpyxl_available,
            pptx_available=pptx_available,
            live_smoke=live_smoke,
        )
        return True

    if command == "status":
        cmd_status(
            json_output=json_output,
            print_result=print_result,
            doc_server=doc_server,
            docx_available=docx_available,
            openpyxl_available=openpyxl_available,
            pptx_available=pptx_available,
        )
        return True

    if command in {"help", "--help", "-h"}:
        cmd_help(
            json_output=json_output,
            print_result=print_result,
            docx_available=docx_available,
            openpyxl_available=openpyxl_available,
            pptx_available=pptx_available,
        )
        return True

    return False
