import os
import json
import hashlib
import re
import shutil
import tempfile
import unittest
import xml.etree.ElementTree as ET
from unittest import mock
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

from cli_anything.onlyoffice.utils.docserver import get_client
from cli_anything.onlyoffice.utils.doc_ops import OOXML_NS


class OnlyOfficeDocOpsTests(unittest.TestCase):
    def setUp(self):
        self.client = get_client()
        self.ops = self.client._doc_ops
        self.tmpdir = tempfile.TemporaryDirectory(prefix="oo-doc-ops-test-")
        self.base = self.tmpdir.name

    def tearDown(self):
        self.tmpdir.cleanup()

    def _path(self, name: str) -> str:
        return os.path.join(self.base, name)

    def _sha256(self, path: str) -> str:
        with open(path, "rb") as handle:
            return hashlib.sha256(handle.read()).hexdigest()

    def _make_reference_fixture(self, path: str, *, ns0: bool = False):
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        normal.paragraph_format.space_after = Pt(0)
        doc.add_paragraph("Body before References.")
        doc.add_page_break()
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.add_run("References").bold = True
        ref = doc.add_paragraph(
            "Del Re, A. C., Fluckiger, C., Horvath, A. O., & Wampold, B. E. "
            "(2021). Examining therapist effects in the alliance-outcome "
            "relationship: A multilevel meta-analysis. Journal of Consulting "
            "and Clinical Psychology, 89(5), 371-378. https://doi.org/10.1037/ccp0000637"
        )
        ref.paragraph_format.left_indent = Inches(0.5)
        ref.paragraph_format.first_line_indent = Inches(-0.5)
        ref.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        ref.paragraph_format.space_after = Pt(0)
        doc.save(path)
        if ns0:
            with ZipFile(path, "r") as zin:
                files = {name: zin.read(name) for name in zin.namelist()}
            document_xml = files["word/document.xml"].decode("utf-8")
            document_xml = document_xml.replace(
                f'xmlns:w="{OOXML_NS["w"]}"',
                f'xmlns:ns0="{OOXML_NS["w"]}"',
                1,
            )
            document_xml = re.sub(
                r"<(/?)w:",
                r"<\1ns0:",
                document_xml,
            )
            document_xml = re.sub(
                r"(\s)(?!xmlns:)w:",
                r"\1ns0:",
                document_xml,
            )
            files["word/document.xml"] = document_xml.encode("utf-8")
            with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
                for name, blob in files.items():
                    zout.writestr(name, blob)

    def _add_docx_part(self, path: str, part_name: str, blob: bytes):
        with ZipFile(path, "r") as zin:
            files = {name: zin.read(name) for name in zin.namelist()}
        files[part_name] = blob
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
            for name, data in files.items():
                zout.writestr(name, data)

    def _pdf_block_payload(self, *, good: bool):
        if good:
            heading_page = 1
            heading_top = 66.0
            first_left = 72.0
            cont_left = 108.0
            right = 520.0
            lines_before = []
        else:
            heading_page = 0
            heading_top = 382.0
            first_left = 85.039
            cont_left = 85.039
            right = 548.0
            lines_before = [
                {
                    "line_id": "line_body",
                    "bbox": {
                        "left": 85.039,
                        "top": 326.0,
                        "right": 540.0,
                        "bottom": 342.0,
                    },
                    "text": "Body before References.",
                }
            ]
        return {
            "success": True,
            "total_pages": 2,
            "pages_scanned": 2,
            "pages": [
                {
                    "page_index": heading_page,
                    "page_number": heading_page + 1,
                    "width": 595.3,
                    "height": 841.9,
                    "blocks": [
                        {
                            "type": "text",
                            "lines": lines_before
                            + [
                                {
                                    "line_id": "line_heading",
                                    "bbox": {
                                        "left": first_left,
                                        "top": heading_top,
                                        "right": first_left + 70,
                                        "bottom": heading_top + 18,
                                    },
                                    "text": "References",
                                },
                                {
                                    "line_id": "line_ref_1",
                                    "bbox": {
                                        "left": first_left,
                                        "top": heading_top + 28,
                                        "right": right,
                                        "bottom": heading_top + 44,
                                    },
                                    "text": (
                                        "Del Re, A. C., Fluckiger, C., Horvath, A. O., & Wampold, B. E. "
                                        "(2021). Examining therapist"
                                    ),
                                },
                                {
                                    "line_id": "line_ref_2",
                                    "bbox": {
                                        "left": cont_left,
                                        "top": heading_top + 56,
                                        "right": right,
                                        "bottom": heading_top + 72,
                                    },
                                    "text": (
                                        "effects in the alliance-outcome relationship: A multilevel meta-analysis. "
                                        "Journal of Consulting"
                                    ),
                                },
                            ],
                        }
                    ],
                }
            ],
        }

    def test_doc_ops_sanitize_removes_comments_and_metadata(self):
        path = self._path("sanitize_comments.docx")
        self.client.create_document(path, "Title", "Body")
        self.client.set_metadata(
            path,
            author="Original Author",
            title="Assignment Draft",
            subject="Research Methods",
            keywords="draft,metadata",
            comments="Internal note",
            category="Assignments",
        )
        comment_result = self.client.add_comment(path, "Review note", 0)
        self.assertTrue(comment_result["success"])
        self._add_docx_part(
            path,
            "word/commentsIds.xml",
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w15:commentsIds xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"/>',
        )
        self._add_docx_part(
            path,
            "word/_rels/comments.xml.rels",
            (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                b'<Relationship Id="rId1" '
                b'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" '
                b'Target="media/comment-secret.png"/></Relationships>'
            ),
        )
        self._add_docx_part(path, "word/media/comment-secret.png", b"comment-private")
        with ZipFile(path, "r") as zin:
            content_types = zin.read("[Content_Types].xml").decode("utf-8")
        if 'Extension="png"' not in content_types:
            content_types = content_types.replace(
                "</Types>",
                '<Default Extension="png" ContentType="image/png"/></Types>',
            )
            self._add_docx_part(
                path,
                "[Content_Types].xml",
                content_types.encode("utf-8"),
            )
        self._add_docx_part(
            path,
            "docProps/custom.xml",
            (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
                b'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
                b'<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="InternalReviewer">'
                b"<vt:lpwstr>Jacky</vt:lpwstr></property></Properties>"
            ),
        )

        before = self.ops.inspect_hidden_data(path)
        self.assertTrue(before["success"])
        self.assertTrue(before["comments_part_present"])
        self.assertGreaterEqual(before["comments_count"], 1)
        self.assertIn("word/commentsIds.xml", before["comment_parts"])
        self.assertEqual(before["comment_relationship_part_count"], 1)
        self.assertIn(
            "word/media/comment-secret.png",
            before["comment_related_targets"],
        )
        self.assertEqual(before["custom_document_properties_count"], 1)
        self.assertEqual(before["core_properties"]["author"], "Original Author")

        result = self.ops.sanitize_document(
            path,
            remove_comments=True,
            clear_metadata=True,
            author="benbi",
        )

        self.assertTrue(result["success"])
        after = result["after"]
        self.assertFalse(after["comments_part_present"])
        self.assertEqual(after["comments_count"], 0)
        self.assertEqual(after["comment_reference_count"], 0)
        self.assertEqual(after["comment_relationship_part_count"], 0)
        self.assertEqual(after["comment_related_target_count"], 0)
        self.assertEqual(after["custom_document_properties_count"], 0)
        self.assertEqual(after["core_properties"]["author"], "benbi")
        self.assertEqual(after["core_properties"]["title"], "")
        self.assertEqual(after["core_properties"]["created"], "")
        self.assertEqual(after["core_properties"]["modified"], "")
        self.assertTrue(after["remove_personal_information"])
        self.assertEqual(result["stats"]["comment_relationship_parts_removed"], 1)
        self.assertEqual(result["stats"]["comment_related_targets_removed"], 1)
        with ZipFile(path, "r") as zin:
            names = set(zin.namelist())
        self.assertNotIn("word/_rels/comments.xml.rels", names)
        self.assertNotIn("word/media/comment-secret.png", names)

    def test_doc_ops_preflight_flags_empty_comment_infrastructure(self):
        path = self._path("empty_comment_parts.docx")
        self.client.create_document(path, "Title", "Body")
        self._add_docx_part(
            path,
            "word/commentsIds.xml",
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w15:commentsIds xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"/>',
        )

        result = self.ops.document_preflight(path)

        self.assertTrue(result["success"])
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(statuses["comments"], "fail")

    def test_doc_ops_ooxml_preflight_rejects_duplicate_entries(self):
        path = self._path("duplicate_entries.docx")
        self.client.create_document(path, "Title", "Body")
        with ZipFile(path, "a", compression=ZIP_DEFLATED) as zout:
            with self.assertWarns(UserWarning):
                zout.writestr("word/document.xml", b"<duplicate/>")

        result = self.ops.inspect_hidden_data(path)

        self.assertFalse(result["success"])
        self.assertEqual(result["error_code"], "docx_ooxml_preflight_failed")
        self.assertEqual(result["preflight"]["status"], "fail")
        self.assertIn("word/document.xml", result["preflight"]["duplicate_names"])
        error_codes = {error["code"] for error in result["preflight"]["errors"]}
        self.assertIn("duplicate_part_name", error_codes)

        sanitize = self.ops.sanitize_document(path, clear_metadata=True)
        self.assertFalse(sanitize["success"])
        self.assertEqual(sanitize["error_code"], "docx_ooxml_preflight_failed")

    def test_doc_ops_ooxml_preflight_rejects_unsafe_parent_part_names(self):
        path = self._path("unsafe_part_name.docx")
        self.client.create_document(path, "Title", "Body")
        with ZipFile(path, "a", compression=ZIP_DEFLATED) as zout:
            zout.writestr("../evil.xml", b"<evil/>")

        result = self.ops.inspect_hidden_data(path)

        self.assertFalse(result["success"])
        self.assertEqual(result["error_code"], "docx_ooxml_preflight_failed")
        unsafe = result["preflight"]["unsafe_part_names"]
        self.assertTrue(any(entry["part"] == "../evil.xml" for entry in unsafe))
        error_codes = {error["code"] for error in result["preflight"]["errors"]}
        self.assertIn("unsafe_part_name", error_codes)

    def test_doc_ops_ooxml_preflight_rejects_oversized_xml_parts(self):
        path = self._path("oversized_xml.docx")
        self.client.create_document(path, "Title", "Body")
        self._add_docx_part(path, "word/oversized.xml", b"<x>" + b"a" * 512 + b"</x>")

        with mock.patch.object(self.ops.__class__, "MAX_DOCX_XML_PART_BYTES", 128):
            result = self.ops.inspect_hidden_data(path)

        self.assertFalse(result["success"])
        self.assertEqual(result["error_code"], "docx_ooxml_preflight_failed")
        oversized = result["preflight"]["oversized_xml_parts"]
        self.assertTrue(any(entry["part"] == "word/oversized.xml" for entry in oversized))
        error_codes = {error["code"] for error in result["preflight"]["errors"]}
        self.assertIn("xml_part_too_large", error_codes)

    def test_doc_ops_ooxml_preflight_rejects_excessive_compression_ratio(self):
        path = self._path("compression_ratio.docx")
        self.client.create_document(path, "Title", "Body")

        with mock.patch.object(self.ops.__class__, "MAX_DOCX_COMPRESSION_RATIO", 1.0):
            result = self.ops.inspect_hidden_data(path)

        self.assertFalse(result["success"])
        self.assertEqual(result["error_code"], "docx_ooxml_preflight_failed")
        self.assertTrue(result["preflight"]["compression_ratio_violations"])
        error_codes = {error["code"] for error in result["preflight"]["errors"]}
        self.assertIn("compression_ratio_exceeded", error_codes)

    def test_doc_ops_reports_external_relationship_visibility_and_risk(self):
        path = self._path("external_relationships.docx")
        self.client.create_document(path, "Title", "Body")
        with ZipFile(path, "r") as zin:
            files = {name: zin.read(name) for name in zin.namelist()}
        rel_name = "word/_rels/document.xml.rels"
        root = ET.fromstring(files[rel_name])
        rel_tag = f"{{{OOXML_NS['rel']}}}Relationship"
        ET.SubElement(
            root,
            rel_tag,
            {
                "Id": "rIdExternalHyperlink",
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                "Target": "https://example.test/",
                "TargetMode": "External",
            },
        )
        ET.SubElement(
            root,
            rel_tag,
            {
                "Id": "rIdExternalAltChunk",
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/altChunk",
                "Target": "https://example.test/chunk.html",
                "TargetMode": "External",
            },
        )
        ET.SubElement(
            root,
            rel_tag,
            {
                "Id": "rIdExternalTemplate",
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate",
                "Target": "https://example.test/template.dotm",
                "TargetMode": "External",
            },
        )
        files[rel_name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
            for name, blob in files.items():
                zout.writestr(name, blob)

        hidden = self.ops.inspect_hidden_data(path)

        self.assertTrue(hidden["success"], hidden)
        self.assertEqual(hidden["external_relationship_count"], 3)
        relationships = {entry["id"]: entry for entry in hidden["external_relationships"]}
        hyperlink = relationships["rIdExternalHyperlink"]
        self.assertEqual(hyperlink["source_part"], "word/document.xml")
        self.assertEqual(hyperlink["target_mode"], "External")
        self.assertEqual(hyperlink["risk_category"], "hyperlink")
        self.assertEqual(hyperlink["risk_level"], "warn")
        self.assertFalse(hyperlink["risky"])

        alt_chunk = relationships["rIdExternalAltChunk"]
        template = relationships["rIdExternalTemplate"]
        self.assertEqual(alt_chunk["risk_category"], "altChunk")
        self.assertEqual(template["risk_category"], "template")
        self.assertTrue(alt_chunk["risky"])
        self.assertTrue(template["risky"])
        self.assertEqual(hidden["risky_external_relationship_count"], 2)

        preflight = self.ops.document_preflight(path)
        self.assertTrue(preflight["success"], preflight)
        statuses = {check["name"]: check["status"] for check in preflight["checks"]}
        self.assertEqual(statuses["external_relationships"], "fail")

    def test_doc_ops_accept_revisions_removes_property_change_markup(self):
        path = self._path("property_revision.docx")
        self.client.create_document(path, "Title", "Body")
        with ZipFile(path, "r") as zin:
            files = {name: zin.read(name) for name in zin.namelist()}
        root = ET.fromstring(files["word/document.xml"])
        first_para = root.find(".//w:p", OOXML_NS)
        self.assertIsNotNone(first_para)
        ppr = first_para.find("w:pPr", OOXML_NS)
        if ppr is None:
            ppr = ET.Element(f"{{{OOXML_NS['w']}}}pPr")
            first_para.insert(0, ppr)
        change = ET.Element(f"{{{OOXML_NS['w']}}}pPrChange")
        change.set(f"{{{OOXML_NS['w']}}}id", "1")
        change.set(f"{{{OOXML_NS['w']}}}author", "Reviewer")
        old_props = ET.Element(f"{{{OOXML_NS['w']}}}pPr")
        change.append(old_props)
        ppr.append(change)
        files["word/document.xml"] = ET.tostring(
            root,
            encoding="utf-8",
            xml_declaration=True,
        )
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
            for name, blob in files.items():
                zout.writestr(name, blob)

        before = self.ops.inspect_hidden_data(path)
        self.assertTrue(before["tracked_changes_present"])
        self.assertGreater(before["tracked_changes"].get("pPrChange", 0), 0)

        result = self.ops.sanitize_document(path, accept_revisions=True)

        self.assertTrue(result["success"], result)
        after = result["after"]
        self.assertFalse(after["tracked_changes_present"])
        self.assertEqual(after["tracked_changes"].get("pPrChange", 0), 0)

    def test_doc_ops_sanitize_preserves_wordprocessing_prefixes(self):
        path = self._path("sanitize_prefixes.docx")
        self.client.create_document(path, "Title", "Body")
        self.client.add_comment(path, "Review note", 0)

        result = self.ops.sanitize_document(path, remove_comments=True)

        self.assertTrue(result["success"])
        with ZipFile(path, "r") as zin:
            document_xml = zin.read("word/document.xml").decode("utf-8")
        self.assertIn("<w:document", document_xml)
        self.assertIn('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"', document_xml)
        self.assertNotIn("<ns0:", document_xml)

    def test_doc_ops_sanitize_can_canonicalize_legacy_ns0_wordprocessingml(self):
        path = self._path("legacy_ns0.docx")
        output_path = self._path("legacy_ns0_canonical.docx")
        self._make_reference_fixture(path, ns0=True)

        before = self.ops.get_formatting_info(path, all_paragraphs=True)
        self.assertFalse(before["ooxml"]["uses_canonical_word_prefix"])
        self.assertEqual(before["ooxml"]["noncanonical_part_count"], 1)

        result = self.ops.sanitize_document(
            path,
            output_path=output_path,
            canonicalize_ooxml=True,
        )

        self.assertTrue(result["success"])
        self.assertTrue(result["canonicalize_ooxml"])
        self.assertGreaterEqual(result["stats"]["ooxml_parts_seen"], 1)
        self.assertGreaterEqual(result["stats"]["ooxml_parts_rewritten"], 1)
        after = self.ops.get_formatting_info(output_path, all_paragraphs=True)
        self.assertTrue(after["ooxml"]["uses_canonical_word_prefix"])
        with ZipFile(output_path, "r") as zin:
            document_xml = zin.read("word/document.xml").decode("utf-8")
        self.assertIn("<w:document", document_xml)
        self.assertNotIn("<ns0:", document_xml)

    def test_doc_ops_canonicalize_preserves_mc_ignorable_declarations(self):
        path = self._path("legacy_ns0_ignorable.docx")
        output_path = self._path("legacy_ns0_ignorable_canonical.docx")
        self._make_reference_fixture(path, ns0=True)
        with ZipFile(path, "r") as zin:
            files = {name: zin.read(name) for name in zin.namelist()}
        document_xml = files["word/document.xml"].decode("utf-8")
        document_xml = document_xml.replace(
            'mc:Ignorable="w14 wp14"',
            'mc:Ignorable="w14 wp14 w15"',
            1,
        )
        document_xml = document_xml.replace(
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"',
            (
                'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
                'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"'
            ),
            1,
        )
        document_xml = document_xml.replace(
            "<ns0:body>",
            '<ns0:body><w15:collapsed w15:val="0"/>',
            1,
        )
        files["word/document.xml"] = document_xml.encode("utf-8")
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
            for name, blob in files.items():
                zout.writestr(name, blob)

        with ZipFile(path, "r") as zin:
            before_xml = zin.read("word/document.xml").decode("utf-8")
        self.assertIn('mc:Ignorable="w14 wp14 w15"', before_xml)
        self.assertIn("xmlns:w14=", before_xml)
        self.assertIn("xmlns:wp14=", before_xml)
        self.assertIn("xmlns:w15=", before_xml)

        result = self.ops.sanitize_document(
            path,
            output_path=output_path,
            canonicalize_ooxml=True,
        )

        self.assertTrue(result["success"])
        with ZipFile(output_path, "r") as zin:
            document_xml = zin.read("word/document.xml").decode("utf-8")
        ET.fromstring(document_xml.encode("utf-8"))
        self.assertIn("<w:document", document_xml)
        self.assertNotIn("<ns0:", document_xml)
        self.assertIn('mc:Ignorable="w14 wp14 w15"', document_xml)
        self.assertIn("xmlns:w14=", document_xml)
        self.assertIn("xmlns:wp14=", document_xml)
        self.assertIn("xmlns:w15=", document_xml)
        self.assertIn("<w15:collapsed", document_xml)

    def test_doc_ops_canonicalize_surfaces_skipped_parse_parts(self):
        path = self._path("legacy_ns0_with_bad_part.docx")
        output_path = self._path("legacy_ns0_with_bad_part_out.docx")
        self._make_reference_fixture(path, ns0=True)
        self._add_docx_part(
            path,
            "word/bad.xml",
            (
                b'<?xml version="1.0" encoding="UTF-8"?>'
                b'<ns0:bad xmlns:ns0="'
                + OOXML_NS["w"].encode("utf-8")
                + b'"><ns0:p>'
            ),
        )

        result = self.ops.sanitize_document(
            path,
            output_path=output_path,
            canonicalize_ooxml=True,
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["stats"]["ooxml_parts_skipped"], 1)
        self.assertEqual(result["canonicalization"]["parts_skipped"], 1)
        self.assertEqual(
            result["canonicalization"]["skipped_parts"][0]["part"],
            "word/bad.xml",
        )

        preflight = self.ops.document_preflight(output_path)
        self.assertTrue(preflight["success"])
        self.assertEqual(preflight["overall_status"], "fail")
        self.assertFalse(preflight["submission_ready"])
        self.assertEqual(preflight["ooxml"]["parse_error_part_count"], 1)
        statuses = {check["name"]: check["status"] for check in preflight["checks"]}
        self.assertEqual(statuses["ooxml_word_prefix"], "fail")

    def test_doc_ops_formatting_info_exposes_reference_indent_details(self):
        path = self._path("formatting_refs.docx")
        self._make_reference_fixture(path)

        result = self.ops.get_formatting_info(path, start=2, limit=3)

        self.assertTrue(result["success"])
        self.assertGreaterEqual(result["paragraph_count"], 4)
        ref_info = next(
            paragraph
            for paragraph in result["paragraphs"]
            if paragraph["text_preview"].startswith("Del Re")
        )
        self.assertEqual(ref_info["style_name"], "Normal")
        self.assertEqual(ref_info["alignment"]["readable"], "left")
        self.assertEqual(
            ref_info["indents"]["raw_ooxml_ind_attrs"]["left"],
            "720",
        )
        self.assertEqual(
            ref_info["indents"]["style_resolved"]["hanging"]["pt"],
            36.0,
        )
        self.assertEqual(
            ref_info["line_spacing"]["style_resolved"]["line_spacing"],
            2.0,
        )

    def test_doc_ops_rendered_layout_audit_passes_good_hanging_indent_pdf(self):
        path = self._path("render_audit_good.docx")
        self._make_reference_fixture(path)
        pdf_path = self._path("render_audit_good.pdf")

        with mock.patch.object(
            self.client,
            "pdf_read_blocks",
            return_value=self._pdf_block_payload(good=True),
        ):
            result = self.ops.rendered_layout_audit(
                path, pdf_path=pdf_path, trusted_pdf=True
            )

        self.assertTrue(result["success"])
        self.assertEqual(result["overall_status"], "pass")
        self.assertTrue(result["source_unchanged"])
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(statuses["hanging_indents_rendered"], "pass")
        self.assertEqual(statuses["rendered_margins"], "pass")

    def test_doc_ops_rendered_layout_audit_generic_profile_skips_reference_warnings(self):
        path = self._path("render_audit_generic.docx")
        doc = Document()
        doc.sections[0].left_margin = Inches(1)
        doc.sections[0].right_margin = Inches(1)
        doc.add_paragraph("Plain report body.")
        doc.save(path)
        payload = {
            "success": True,
            "total_pages": 1,
            "pages_scanned": 1,
            "pages": [
                {
                    "page_index": 0,
                    "page_number": 1,
                    "width": 595.3,
                    "height": 841.9,
                    "blocks": [
                        {
                            "type": "text",
                            "lines": [
                                {
                                    "line_id": "line_body",
                                    "bbox": {
                                        "left": 72.0,
                                        "top": 72.0,
                                        "right": 220.0,
                                        "bottom": 88.0,
                                    },
                                    "text": "Plain report body.",
                                }
                            ],
                        }
                    ],
                }
            ],
        }

        def fake_office_to_pdf(file_path, output_path=None):
            self.assertEqual(file_path, path)
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            return {"success": True, "output_file": output_path, "pages": 1}

        with mock.patch.object(self.client, "_office_to_pdf", side_effect=fake_office_to_pdf):
            with mock.patch.object(self.client, "pdf_read_blocks", return_value=payload):
                result = self.ops.rendered_layout_audit(path)

        self.assertTrue(result["success"])
        self.assertEqual(result["profile"], "generic")
        self.assertEqual(result["overall_status"], "pass")
        self.assertTrue(result["submission_ready"])
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertNotIn("references_heading_rendered", statuses)
        self.assertEqual(statuses["pdf_text_rendered"], "pass")
        self.assertEqual(statuses["pdf_provenance"], "pass")
        self.assertEqual(statuses["generic_margin_envelope"], "pass")

    def test_doc_ops_rendered_layout_audit_generic_checks_vertical_margins(self):
        path = self._path("render_audit_generic_vertical.docx")
        doc = Document()
        doc.sections[0].left_margin = Inches(1)
        doc.sections[0].right_margin = Inches(1)
        doc.sections[0].top_margin = Inches(1)
        doc.sections[0].bottom_margin = Inches(1)
        doc.add_paragraph("Plain report body.")
        doc.save(path)
        payload = {
            "success": True,
            "total_pages": 1,
            "pages_scanned": 1,
            "pages": [
                {
                    "page_index": 0,
                    "page_number": 1,
                    "width": 595.3,
                    "height": 841.9,
                    "blocks": [
                        {
                            "type": "text",
                            "lines": [
                                {
                                    "line_id": "line_body",
                                    "bbox": {
                                        "left": 72.0,
                                        "top": 24.0,
                                        "right": 220.0,
                                        "bottom": 40.0,
                                    },
                                    "text": "Plain report body.",
                                }
                            ],
                        }
                    ],
                }
            ],
        }

        def fake_office_to_pdf(file_path, output_path=None):
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            return {"success": True, "output_file": output_path, "pages": 1}

        with mock.patch.object(self.client, "_office_to_pdf", side_effect=fake_office_to_pdf):
            with mock.patch.object(self.client, "pdf_read_blocks", return_value=payload):
                result = self.ops.rendered_layout_audit(path, profile="generic")

        self.assertTrue(result["success"])
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(statuses["generic_margin_envelope"], "fail")

    def test_doc_ops_rendered_layout_audit_checks_all_continuation_lines(self):
        path = self._path("render_audit_bad_third_continuation.docx")
        self._make_reference_fixture(path)
        payload = self._pdf_block_payload(good=True)
        payload["pages"][0]["blocks"][0]["lines"].append(
            {
                "line_id": "line_ref_3",
                "bbox": {
                    "left": 72.0,
                    "top": 150.0,
                    "right": 520.0,
                    "bottom": 166.0,
                },
                "text": "and Clinical Psychology, 89(5), 371-378.",
            }
        )

        with mock.patch.object(self.client, "pdf_read_blocks", return_value=payload):
            result = self.ops.rendered_layout_audit(
                path,
                pdf_path=self._path("render_audit_bad_third.pdf"),
                trusted_pdf=True,
            )

        self.assertTrue(result["success"])
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(statuses["hanging_indents_rendered"], "fail")

    def test_doc_ops_rendered_layout_audit_external_pdf_blocks_readiness(self):
        path = self._path("render_audit_external_pdf.docx")
        doc = Document()
        doc.sections[0].left_margin = Inches(1)
        doc.sections[0].right_margin = Inches(1)
        doc.add_paragraph("Plain report body.")
        doc.save(path)
        pdf_path = self._path("render_audit_external_pdf.pdf")
        payload = {
            "success": True,
            "total_pages": 1,
            "pages_scanned": 1,
            "pages": [
                {
                    "page_index": 0,
                    "page_number": 1,
                    "width": 595.3,
                    "height": 841.9,
                    "blocks": [
                        {
                            "type": "text",
                            "lines": [
                                {
                                    "line_id": "line_body",
                                    "bbox": {
                                        "left": 72.0,
                                        "top": 72.0,
                                        "right": 220.0,
                                        "bottom": 88.0,
                                    },
                                    "text": "Plain report body.",
                                }
                            ],
                        }
                    ],
                }
            ],
        }

        with mock.patch.object(self.client, "pdf_read_blocks", return_value=payload):
            result = self.ops.rendered_layout_audit(path, pdf_path=pdf_path)

        self.assertTrue(result["success"])
        self.assertEqual(result["profile"], "generic")
        self.assertEqual(result["overall_status"], "warn")
        self.assertFalse(result["submission_ready"])
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(statuses["pdf_provenance"], "warn")
        self.assertEqual(statuses["generic_margin_envelope"], "pass")

    def test_doc_ops_rendered_layout_audit_uses_ooxml_start_indent(self):
        path = self._path("render_audit_start_indent.docx")
        self._make_reference_fixture(path)
        with ZipFile(path, "r") as zin:
            files = {name: zin.read(name) for name in zin.namelist()}
        document_xml = files["word/document.xml"].decode("utf-8")
        document_xml = document_xml.replace('w:left="720"', 'w:start="720"', 1)
        files["word/document.xml"] = document_xml.encode("utf-8")
        with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
            for name, blob in files.items():
                zout.writestr(name, blob)
        pdf_path = self._path("render_audit_start_indent.pdf")

        info = self.ops.get_formatting_info(path, all_paragraphs=True)
        self.assertTrue(info["success"])
        ref_info = next(
            paragraph
            for paragraph in info["paragraphs"]
            if paragraph["text_preview"].startswith("Del Re")
        )
        self.assertEqual(
            ref_info["indents"]["style_resolved"]["effective_left"]["pt"],
            36.0,
        )
        self.assertEqual(
            ref_info["indents"]["style_resolved"]["effective_left_source"],
            "direct_ooxml_start",
        )

        with mock.patch.object(
            self.client,
            "pdf_read_blocks",
            return_value=self._pdf_block_payload(good=True),
        ):
            result = self.ops.rendered_layout_audit(
                path, pdf_path=pdf_path, trusted_pdf=True
            )

        self.assertTrue(result["success"])
        self.assertEqual(result["overall_status"], "pass")

    def test_doc_ops_reference_page_break_val_zero_is_not_expected(self):
        path = self._path("render_audit_page_break_zero.docx")
        doc = Document()
        doc.add_paragraph("Body before references on same page.")
        heading = doc.add_paragraph("References")
        ppr = heading._element.get_or_add_pPr()
        page_break = OxmlElement("w:pageBreakBefore")
        page_break.set(qn("w:val"), "0")
        ppr.append(page_break)
        ref = doc.add_paragraph("Smith, J. (2024). Example title. Example Journal.")
        ref.paragraph_format.left_indent = Inches(0.5)
        ref.paragraph_format.first_line_indent = Inches(-0.5)
        doc.save(path)

        expectations = self.ops._reference_layout_expectations(Document(path))

        self.assertFalse(expectations["references_heading_expected_page_break"])
        self.assertTrue(
            expectations["references_heading_page_break"]["present_in_ooxml"]
        )
        self.assertFalse(
            expectations["references_heading_page_break"]["enabled_in_ooxml"]
        )

    def test_doc_ops_reference_page_break_detects_inline_breaks_and_blank_paragraphs(self):
        heading_break_path = self._path("render_audit_heading_inline_break.docx")
        doc = Document()
        doc.add_paragraph("Body before references.")
        heading = doc.add_paragraph()
        run = heading.add_run()
        run.add_break(WD_BREAK.PAGE)
        run.add_text("References")
        ref = doc.add_paragraph("Smith, J. (2024). Example title. Example Journal.")
        ref.paragraph_format.left_indent = Inches(0.5)
        ref.paragraph_format.first_line_indent = Inches(-0.5)
        doc.save(heading_break_path)

        heading_expectations = self.ops._reference_layout_expectations(
            Document(heading_break_path)
        )

        self.assertTrue(heading_expectations["references_heading_expected_page_break"])
        self.assertTrue(
            heading_expectations["references_heading_inline_page_break_before_text"]
        )

        blank_path = self._path("render_audit_blank_after_break.docx")
        doc = Document()
        doc.add_paragraph("Body before references.")
        break_para = doc.add_paragraph()
        break_para.add_run().add_break(WD_BREAK.PAGE)
        doc.add_paragraph("")
        ref_heading = doc.add_paragraph("References")
        ref = doc.add_paragraph("Jones, A. (2025). Example source. Example Journal.")
        ref.paragraph_format.left_indent = Inches(0.5)
        ref.paragraph_format.first_line_indent = Inches(-0.5)
        doc.save(blank_path)

        blank_expectations = self.ops._reference_layout_expectations(Document(blank_path))

        self.assertTrue(blank_expectations["references_heading_expected_page_break"])
        self.assertTrue(blank_expectations["previous_paragraph_inline_page_break"])

    def test_doc_ops_rendered_layout_audit_flags_ns0_layout_mismatch(self):
        path = self._path("render_audit_bad.docx")
        self._make_reference_fixture(path, ns0=True)
        pdf_path = self._path("render_audit_bad.pdf")

        with mock.patch.object(
            self.client,
            "pdf_read_blocks",
            return_value=self._pdf_block_payload(good=False),
        ):
            result = self.ops.rendered_layout_audit(
                path, pdf_path=pdf_path, trusted_pdf=True
            )

        self.assertTrue(result["success"])
        self.assertEqual(result["overall_status"], "fail")
        self.assertTrue(result["source_unchanged"])
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(statuses["ooxml_word_prefix"], "fail")
        self.assertEqual(statuses["references_page_break"], "fail")
        self.assertEqual(statuses["hanging_indents_rendered"], "fail")
        self.assertEqual(statuses["horizontal_alignment"], "fail")
        self.assertEqual(statuses["rendered_margins"], "fail")

    def test_doc_ops_render_readiness_blocks_ns0_even_with_good_pdf_geometry(self):
        path = self._path("render_audit_ns0_good_geometry.docx")
        self._make_reference_fixture(path, ns0=True)
        pdf_path = self._path("render_audit_ns0_good_geometry.pdf")

        with mock.patch.object(
            self.client,
            "pdf_read_blocks",
            return_value=self._pdf_block_payload(good=True),
        ):
            audit = self.ops.rendered_layout_audit(
                path, pdf_path=pdf_path, trusted_pdf=True
            )

        self.assertTrue(audit["success"])
        self.assertEqual(audit["overall_status"], "fail")
        self.assertFalse(audit["submission_ready"])
        statuses = {check["name"]: check["status"] for check in audit["checks"]}
        self.assertEqual(statuses["ooxml_word_prefix"], "fail")
        self.assertEqual(statuses["hanging_indents_rendered"], "pass")
        self.assertEqual(audit["ooxml"]["noncanonical_part_count"], 1)

        def fake_office_to_pdf(file_path, output_path=None):
            self.assertEqual(file_path, path)
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            return {"success": True, "output_file": output_path, "pages": 2}

        with mock.patch.object(self.client, "_office_to_pdf", side_effect=fake_office_to_pdf):
            with mock.patch.object(
                self.client,
                "pdf_read_blocks",
                return_value=self._pdf_block_payload(good=True),
            ):
                preflight = self.ops.document_preflight(path, rendered_layout=True)

        self.assertTrue(preflight["success"])
        self.assertEqual(preflight["overall_status"], "fail")
        self.assertFalse(preflight["submission_ready"])
        statuses = {check["name"]: check["status"] for check in preflight["checks"]}
        self.assertEqual(statuses["ooxml_word_prefix"], "fail")
        self.assertEqual(statuses["rendered_layout"], "fail")

    def test_doc_ops_read_only_layout_workflows_preserve_source_hash(self):
        path = self._path("read_only_hash.docx")
        self.client.create_document(path, "Title", "Body")
        before_hash = self._sha256(path)
        preview_dir = self._path("preview_hash")
        pdf_out = self._path("converted.pdf")

        def fake_doc_to_pdf(file_path, output_path=None, layout_warnings=False):
            self.assertTrue(file_path.endswith(".docx"))
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            return {"success": True, "output_file": output_path, "pages": 1}

        def fake_pdf_page_to_image(file_path, output_dir, pages=None, dpi=150, fmt="png"):
            return {
                "success": True,
                "total_pages": 1,
                "pages_rendered": 1,
                "images": [{"page": 0, "file": os.path.join(output_dir, "page_000.png")}],
            }

        def fake_pdf_read_blocks(
            file_path,
            pages=None,
            include_spans=True,
            include_images=False,
            include_empty=False,
        ):
            return {
                "success": True,
                "pages_scanned": 1,
                "total_pages": 1,
                "pages": [
                    {
                        "page_index": 0,
                        "page_number": 1,
                        "blocks": [
                            {
                                "type": "text",
                                "bbox": {
                                    "left": 10,
                                    "top": 10,
                                    "right": 180,
                                    "bottom": 30,
                                },
                                "lines": [
                                    {
                                        "line_id": "line_1",
                                        "bbox": {
                                            "left": 10,
                                            "top": 10,
                                            "right": 180,
                                            "bottom": 30,
                                        },
                                        "spans": [
                                            {
                                                "span_id": "span_1",
                                                "text": "SLOANE_P_0001 Title",
                                                "bbox": {
                                                    "left": 10,
                                                    "top": 10,
                                                    "right": 180,
                                                    "bottom": 30,
                                                },
                                            }
                                        ],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }

        with mock.patch.object(self.client, "doc_to_pdf", side_effect=fake_doc_to_pdf):
            with mock.patch.object(
                self.client, "pdf_page_to_image", side_effect=fake_pdf_page_to_image
            ):
                with mock.patch.object(
                    self.client, "pdf_read_blocks", side_effect=fake_pdf_read_blocks
                ):
                    self.assertTrue(
                        self.ops.get_formatting_info(path, all_paragraphs=True)["success"]
                    )
                    self.assertTrue(self.ops.document_preflight(path)["success"])
                    self.assertTrue(self.ops.preview_document(path, preview_dir)["success"])
                    self.assertTrue(self.ops.doc_render_map(path)["success"])
        with mock.patch.object(self.client, "_office_to_pdf", side_effect=fake_doc_to_pdf):
            self.assertTrue(self.ops.doc_to_pdf(path, output_path=pdf_out)["success"])

        self.assertEqual(before_hash, self._sha256(path))

    def test_doc_ops_preflight_flags_mixed_fonts_and_metadata(self):
        path = self._path("preflight.docx")
        doc = Document()
        first = doc.add_paragraph()
        first.add_run("Times paragraph")
        second = doc.add_paragraph()
        run = second.add_run("Calibri paragraph")
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        doc.save(path)
        self.client.set_page_layout(path, page_size="A4")
        self.client.set_metadata(path, author="benbi", title="Draft")

        result = self.ops.document_preflight(
            path,
            expected_page_size="A4",
            expected_font_name="Times New Roman",
            expected_font_size=11,
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["overall_status"], "warn")
        check_status = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(check_status["page_size"], "pass")
        self.assertEqual(check_status["metadata"], "warn")
        self.assertEqual(check_status["font_names"], "warn")
        self.assertEqual(check_status["font_sizes"], "warn")
        self.assertEqual(result["font_audit"]["unexpected_font_run_count"], 2)
        self.assertEqual(result["font_audit"]["unexpected_size_run_count"], 2)

    def test_doc_ops_preview_uses_host_doc_pdf_pipeline(self):
        path = self._path("preview.docx")
        self.client.create_document(path, "Title", "Intro")
        output_dir = self._path("previews")

        captured = {}

        def fake_doc_to_pdf(file_path, output_path=None):
            self.assertEqual(file_path, path)
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            captured["pdf_path"] = output_path
            return {
                "success": True,
                "input_file": file_path,
                "output_file": output_path,
                "pages": 2,
            }

        def fake_pdf_page_to_image(file_path, render_dir, pages=None, dpi=150, fmt="png"):
            self.assertEqual(file_path, captured["pdf_path"])
            self.assertTrue(os.path.exists(file_path))
            self.assertEqual(render_dir, output_dir)
            self.assertEqual(pages, "0-1")
            self.assertEqual(dpi, 200)
            self.assertEqual(fmt, "jpg")
            return {
                "success": True,
                "total_pages": 2,
                "pages_rendered": 2,
                "images": [{"page": 0, "file": os.path.join(render_dir, "page_000.jpg")}],
            }

        with mock.patch.object(self.client, "doc_to_pdf", side_effect=fake_doc_to_pdf):
            with mock.patch.object(
                self.client, "pdf_page_to_image", side_effect=fake_pdf_page_to_image
            ):
                result = self.ops.preview_document(
                    path, output_dir, pages="0-1", dpi=200, fmt="jpg"
                )

        self.assertTrue(result["success"])
        self.assertEqual(result["total_pages"], 2)
        self.assertEqual(result["pages_rendered"], 2)
        self.assertEqual(result["format"], "jpg")
        self.assertFalse(os.path.exists(captured["pdf_path"]))

    def test_doc_ops_render_map_uses_host_pdf_block_pipeline(self):
        path = self._path("render_map.docx")
        doc = Document()
        doc.add_paragraph("Executive summary")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "42"
        doc.save(path)

        def fake_doc_to_pdf(file_path, output_path=None):
            self.assertTrue(file_path.endswith(".docx"))
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            return {"success": True, "output_file": output_path, "pages": 1}

        def fake_pdf_read_blocks(file_path, pages=None, include_spans=True, include_images=False, include_empty=False):
            self.assertTrue(file_path.endswith(".pdf"))
            self.assertTrue(os.path.exists(file_path))
            self.assertTrue(include_spans)
            return {
                "success": True,
                "pages_scanned": 1,
                "total_pages": 1,
                "pages": [
                    {
                        "page_index": 0,
                        "page_number": 1,
                        "blocks": [
                            {
                                "block_id": "page_0_block_0",
                                "type": "text",
                                "bbox": {"left": 10, "top": 10, "right": 200, "bottom": 40, "width": 190, "height": 30},
                                "lines": [
                                    {
                                        "line_id": "page_0_block_0_line_0",
                                        "bbox": {"left": 10, "top": 10, "right": 200, "bottom": 25, "width": 190, "height": 15},
                                        "spans": [
                                            {
                                                "span_id": "page_0_block_0_line_0_span_0",
                                                "text": "SLOANE_P_0001 Executive summary",
                                                "bbox": {"left": 10, "top": 10, "right": 200, "bottom": 25, "width": 190, "height": 15},
                                            }
                                        ],
                                    }
                                ],
                            },
                            {
                                "block_id": "page_0_block_1",
                                "type": "text",
                                "bbox": {"left": 10, "top": 50, "right": 200, "bottom": 80, "width": 190, "height": 30},
                                "lines": [
                                    {
                                        "line_id": "page_0_block_1_line_0",
                                        "bbox": {"left": 10, "top": 50, "right": 200, "bottom": 65, "width": 190, "height": 15},
                                        "spans": [
                                            {
                                                "span_id": "page_0_block_1_line_0_span_0",
                                                "text": "SLOANE_T1R1C1 42",
                                                "bbox": {"left": 10, "top": 50, "right": 200, "bottom": 65, "width": 190, "height": 15},
                                            }
                                        ],
                                    }
                                ],
                            },
                        ],
                    }
                ],
            }

        with mock.patch.object(self.client, "doc_to_pdf", side_effect=fake_doc_to_pdf):
            with mock.patch.object(self.client, "pdf_read_blocks", side_effect=fake_pdf_read_blocks):
                result = self.ops.doc_render_map(path)

        self.assertTrue(result["success"])
        self.assertEqual(result["mapped_paragraph_count"], 1)
        self.assertEqual(result["mapped_table_cell_count"], 1)
        self.assertEqual(result["paragraphs"][0]["page_number"], 1)
        self.assertEqual(result["table_cells"][0]["cell_ref"], "T1R1C1")

    def test_doc_ops_basic_crud_replace_format_and_counts(self):
        path = self._path("basic.docx")

        created = self.ops.create_document(path, "Report Title", "First paragraph")
        self.assertTrue(created["success"])

        appended = self.ops.append_to_document(path, "Second paragraph\nThird paragraph")
        self.assertTrue(appended["success"])

        replaced = self.ops.search_replace_document(path, "Second", "Updated Second")
        self.assertTrue(replaced["success"])
        self.assertEqual(replaced["replacements"], 1)

        formatted = self.ops.format_paragraph(
            path,
            1,
            bold=True,
            italic=True,
            font_name="Times New Roman",
            font_size=12,
            alignment="center",
        )
        self.assertTrue(formatted["success"])

        read_back = self.ops.read_document(path)
        self.assertTrue(read_back["success"])
        self.assertEqual(
            read_back["paragraphs"],
            [
                "Report Title",
                "First paragraph",
                "Updated Second paragraph",
                "Third paragraph",
            ],
        )

        counts = self.ops.word_count(path)
        self.assertTrue(counts["success"])
        self.assertGreaterEqual(counts["words"], 8)
        self.assertEqual(counts["paragraphs"], 4)

        doc = Document(path)
        para = doc.paragraphs[1]
        run = para.runs[0]
        self.assertTrue(run.bold)
        self.assertTrue(run.italic)
        self.assertEqual(run.font.name, "Times New Roman")
        self.assertEqual(round(run.font.size.pt), 12)
        self.assertIsNotNone(para.alignment)

    def test_doc_ops_table_search_insert_delete_and_list_helpers(self):
        path = self._path("table_ops.docx")
        self.ops.create_document(path, "Table Report", "Intro")

        table_result = self.ops.add_table(path, "Name,Value", "alpha,1;beta,2")
        self.assertTrue(table_result["success"])
        self.assertEqual(table_result["rows"], 2)

        tables = self.ops.read_tables(path)
        self.assertTrue(tables["success"])
        self.assertEqual(tables["table_count"], 1)
        self.assertEqual(tables["tables"][0]["headers"], ["Name", "Value"])
        self.assertEqual(tables["tables"][0]["rows"][1], ["beta", "2"])

        search = self.ops.search_document(path, "beta")
        self.assertTrue(search["success"])
        self.assertEqual(search["matches"], 0)
        self.assertEqual(search["table_matches"], 1)

        inserted = self.ops.insert_paragraph(path, "Inserted section", 1, "Heading 1")
        self.assertTrue(inserted["success"])
        deleted = self.ops.delete_paragraph(path, 1)
        self.assertTrue(deleted["success"])
        self.assertEqual(deleted["deleted_text"], "Inserted section")

        listed = self.ops.add_list(path, ["Item one", "Item two"], "bullet")
        self.assertTrue(listed["success"])
        self.assertEqual(listed["items_added"], 2)

        page_break = self.ops.add_page_break(path)
        self.assertTrue(page_break["success"])

        styles = self.ops.list_styles(path)
        self.assertTrue(styles["success"])
        self.assertGreater(styles["paragraph_style_count"], 0)

    def test_doc_ops_metadata_layout_and_references_roundtrip(self):
        path = self._path("refs_layout.docx")
        self.ops.create_document(path, "Findings", "Body text")

        metadata = self.ops.set_metadata(
            path,
            author="benbi",
            title="Assignment",
            subject="Methods",
            keywords="health,readiness",
            comments="internal",
            category="Reports",
        )
        self.assertTrue(metadata["success"])

        metadata_read = self.ops.get_metadata(path)
        self.assertTrue(metadata_read["success"])
        self.assertEqual(metadata_read["author"], "benbi")
        self.assertEqual(metadata_read["title"], "Assignment")

        layout = self.ops.set_page_layout(
            path,
            orientation="landscape",
            page_size="A4",
            header_text="Running header",
            page_numbers=True,
        )
        self.assertTrue(layout["success"])
        self.assertEqual(layout["page_size"], "A4")
        self.assertEqual(layout["orientation"], "landscape")

        formatting = self.ops.get_formatting_info(path)
        self.assertTrue(formatting["success"])
        self.assertEqual(formatting["sections"][0]["orientation"], "landscape")

        negative = self.ops.set_page_layout(path, margins={"left": -0.25})
        too_large = self.ops.set_page_layout(path, margins={"left": 7.0, "right": 7.0})
        self.assertFalse(negative["success"])
        self.assertIn("non-negative", negative["error"])
        self.assertFalse(too_large["success"])
        self.assertIn("positive usable page width", too_large["error"])

        ref = {
            "author": "Smith, J.",
            "year": "2024",
            "title": "Work readiness in context",
            "source": "Journal of Health Education",
            "volume": "15",
            "issue": "1",
            "pages": "10-20",
            "doi": "10.1000/example",
            "type": "journal",
        }
        add_ref = self.ops.add_reference(path, json.dumps(ref))
        self.assertTrue(add_ref["success"])
        self.assertEqual(add_ref["action"], "added")

        duplicate_ref = self.ops.add_reference(path, json.dumps(ref))
        self.assertTrue(duplicate_ref["success"])
        self.assertEqual(duplicate_ref["action"], "duplicate_skipped")

        built = self.ops.build_references(path)
        self.assertTrue(built["success"])
        self.assertEqual(built["references_added"], 1)

        read_back = self.ops.read_document(path)
        self.assertTrue(read_back["success"])
        self.assertIn("References", read_back["paragraphs"])
        self.assertIn("Smith, J. (2024). Work readiness in context.", read_back["full_text"])

    def test_doc_ops_add_reference_rewrites_sidecar_atomically_with_backup(self):
        path = self._path("refs_atomic.docx")
        self.client.create_document(path, "Title", "Body")
        first = {
            "author": "Smith, J.",
            "year": "2024",
            "title": "First source",
        }
        second = {
            "author": "Jones, A.",
            "year": "2025",
            "title": "Second source",
        }

        first_result = self.ops.add_reference(path, json.dumps(first))
        self.assertTrue(first_result["success"])
        self.assertIsNone(first_result["backup"])
        second_result = self.ops.add_reference(path, json.dumps(second))

        self.assertTrue(second_result["success"])
        self.assertEqual(second_result["action"], "added")
        self.assertIsNotNone(second_result["backup"])
        self.assertTrue(os.path.exists(second_result["backup"]))
        with open(path + ".refs.json", "r", encoding="utf-8") as handle:
            refs = json.load(handle)
        self.assertEqual([ref["title"] for ref in refs], ["First source", "Second source"])

    def test_doc_ops_citation_audit_matches_apa_like_citations_and_references(self):
        path = self._path("citation_audit_pass.docx")
        doc = Document()
        doc.add_paragraph(
            "Smith (2024) framed the issue, and later evidence supported it "
            "(Jones, 2023; Smith, 2024, p. 12)."
        )
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Table note also cites Brown and Green (2022)."
        doc.add_paragraph("References")
        doc.add_paragraph("Brown, T., & Green, R. (2022). Table evidence. Journal.")
        doc.add_paragraph("Jones, A. (2023). Other evidence. Journal.")
        doc.add_paragraph("Smith, J. (2024). Framing evidence. Journal.")
        doc.save(path)
        before_hash = self._sha256(path)

        result = self.ops.citation_audit(path)

        self.assertTrue(result["success"], result)
        self.assertEqual(result["overall_status"], "pass")
        self.assertEqual(result["counts"]["missing_references"], 0)
        self.assertEqual(result["counts"]["uncited_references"], 0)
        self.assertEqual(result["counts"]["reference_entries"], 3)
        self.assertIn(
            "brown & green|2022",
            {citation["key"] for citation in result["in_text_citations"]},
        )
        self.assertEqual(before_hash, self._sha256(path))
        self.assertFalse(result["network_verification"])

    def test_doc_ops_citation_audit_flags_missing_uncited_malformed_and_sidecar(self):
        path = self._path("citation_audit_fail.docx")
        doc = Document()
        doc.add_paragraph("The draft cites Smith (2024), Jones (2020), and (Taylor, 2025).")
        doc.add_paragraph("References")
        doc.add_paragraph("Smith, J. (2023). Wrong year evidence. Journal.")
        doc.add_paragraph("Unused, A. (2021). Unused evidence. Journal.")
        doc.add_paragraph("Malformed reference without a year.")
        doc.save(path)
        with open(path + ".refs.json", "w", encoding="utf-8") as handle:
            json.dump(
                [
                    {
                        "author": "Smith, J.",
                        "year": "2023",
                        "title": "Wrong year evidence",
                    },
                    {
                        "author": "Sidecar, S.",
                        "year": "2026",
                        "title": "Only in sidecar",
                    },
                ],
                handle,
            )

        result = self.ops.citation_audit(path, include_sidecar=True)

        self.assertTrue(result["success"], result)
        self.assertEqual(result["overall_status"], "fail")
        self.assertGreaterEqual(result["counts"]["missing_references"], 3)
        self.assertGreaterEqual(result["counts"]["uncited_references"], 2)
        self.assertEqual(result["counts"]["malformed_reference_entries"], 1)
        self.assertTrue(result["findings"]["mismatched_entries"])
        self.assertTrue(result["sidecar"]["sidecar_missing_from_docx"])
        self.assertIn(
            "reference_sidecar",
            {check["name"] for check in result["checks"]},
        )

    def test_doc_ops_extract_images_rejects_unsafe_prefix_and_format(self):
        path = self._path("image_extract_safety.docx")
        self.client.create_document(path, "Title", "Body")
        output_dir = self._path("images")

        bad_prefix = self.ops.extract_images_from_docx(
            path, output_dir, prefix="../escape"
        )
        bad_format = self.ops.extract_images_from_docx(
            path, output_dir, fmt="../png"
        )

        self.assertFalse(bad_prefix["success"])
        self.assertIn("Unsafe image prefix", bad_prefix["error"])
        self.assertEqual(bad_prefix["error_code"], "unsafe_output_prefix")
        self.assertFalse(bad_format["success"])
        self.assertEqual(bad_format["error_code"], "unsupported_image_format")

    def test_doc_ops_extract_images_includes_header_images(self):
        path = self._path("header_images.docx")
        image_path = self._path("header_source.png")
        output_dir = self._path("header_images")
        Image.new("RGB", (16, 8), color="green").save(image_path)
        doc = Document()
        section = doc.sections[0]
        header = section.header
        header.paragraphs[0].add_run().add_picture(image_path, width=Inches(1))
        doc.add_paragraph("Body only.")
        doc.save(path)

        result = self.ops.extract_images_from_docx(path, output_dir)

        self.assertTrue(result["success"], result)
        self.assertEqual(result["images_extracted"], 1)
        self.assertTrue(result["images"][0]["source_part"].startswith("word/header"))

    def test_doc_ops_extract_images_ignores_precreated_symlink_temp_file(self):
        path = self._path("symlink_temp.docx")
        image_path = self._path("symlink_source.png")
        output_dir = self._path("symlink_images")
        sentinel = self._path("sentinel.txt")
        Image.new("RGB", (16, 8), color="blue").save(image_path)
        self.client.create_document(path, "Title", "Body")
        add_image = self.client.add_image(path, image_path, width_inches=1.0)
        self.assertTrue(add_image["success"], add_image)
        os.makedirs(output_dir, exist_ok=True)
        with open(sentinel, "wb") as handle:
            handle.write(b"do-not-touch")
        predictable_temp = os.path.join(output_dir, ".image_000.png.tmp")
        try:
            os.symlink(sentinel, predictable_temp)
        except (OSError, NotImplementedError) as exc:
            self.skipTest(f"symlink creation unavailable: {exc}")

        result = self.ops.extract_images_from_docx(
            path,
            output_dir,
            fmt="png",
            prefix="image",
        )

        self.assertTrue(result["success"], result)
        self.assertEqual(result["images_extracted"], 1)
        with open(sentinel, "rb") as handle:
            self.assertEqual(handle.read(), b"do-not-touch")
        self.assertTrue(os.path.islink(predictable_temp))
        self.assertTrue(os.path.exists(result["images"][0]["file"]))

    def test_doc_ops_extract_images_enforces_embedded_image_resource_limits(self):
        path = self._path("bounded_images.docx")
        image_path = self._path("bounded_source.png")
        Image.new("RGB", (12, 8), color="navy").save(image_path)
        self.client.create_document(path, "Title", "Body")
        add_image = self.client.add_image(path, image_path, width_inches=1.0)
        self.assertTrue(add_image["success"], add_image)

        with mock.patch.object(self.ops.__class__, "MAX_EXTRACT_IMAGES", 0):
            count_limited = self.ops.extract_images_from_docx(
                path,
                self._path("count_limited"),
                fmt="png",
            )

        self.assertTrue(count_limited["success"])
        self.assertTrue(count_limited["truncated"])
        self.assertEqual(count_limited["images_extracted"], 0)
        self.assertEqual(count_limited["resource_limits"]["max_images"], 0)
        self.assertTrue(any("Stopped after 0 images" in warning for warning in count_limited["warnings"]))

        with mock.patch.object(self.ops.__class__, "MAX_EXTRACT_IMAGE_COMPRESSED_BYTES", 1):
            byte_limited = self.ops.extract_images_from_docx(
                path,
                self._path("byte_limited"),
                fmt="png",
            )

        self.assertTrue(byte_limited["success"])
        self.assertEqual(byte_limited["images_extracted"], 0)
        self.assertEqual(byte_limited["images_skipped"], 1)
        self.assertEqual(byte_limited["images"][0]["error_code"], "image_compressed_bytes_limit_exceeded")

        with mock.patch.object(self.ops.__class__, "MAX_EXTRACT_IMAGE_PIXELS", 1):
            with mock.patch.object(Image.Image, "save") as image_save:
                pixel_limited = self.ops.extract_images_from_docx(
                    path,
                    self._path("pixel_limited"),
                    fmt="png",
                )

        self.assertTrue(pixel_limited["success"])
        self.assertEqual(pixel_limited["images_extracted"], 0)
        self.assertEqual(pixel_limited["images_skipped"], 1)
        self.assertEqual(pixel_limited["images"][0]["error_code"], "image_pixel_limit_exceeded")
        image_save.assert_not_called()

    def test_doc_ops_normalize_format_updates_whole_document_and_preserves_text(self):
        path = self._path("normalize_source.docx")
        output_path = self._path("normalize_output.docx")
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "Running header"
        title = doc.add_paragraph("Assignment title")
        title.style = doc.styles["Title"]
        title.add_run(" extra")
        doc.add_paragraph("Body paragraph.")
        doc.add_paragraph("References")
        ref = doc.add_paragraph("Smith, J. (2025). Example source. Example Journal.")
        ref.paragraph_format.left_indent = Inches(0)
        ref.paragraph_format.first_line_indent = Inches(0)
        doc.save(path)

        result = self.ops.normalize_document_format(
            path,
            output_path=output_path,
            font_name="Times New Roman",
            body_font_size=11,
            title_font_size=12,
            line_spacing="double",
            paragraph_after=12,
            clear_theme_fonts=True,
            remove_style_borders=True,
            reference_hanging_inches=0.5,
        )

        self.assertTrue(result["success"], result)
        self.assertTrue(result["text_preserved"])
        self.assertEqual(result["font_audit"]["unexpected_font_run_count"], 0)
        self.assertGreater(result["stats"]["runs_updated"], 0)
        self.assertGreaterEqual(result["stats"]["reference_paragraphs_updated"], 1)
        info = self.ops.get_formatting_info(output_path, all_paragraphs=True)
        ref_info = next(
            paragraph
            for paragraph in info["paragraphs"]
            if paragraph["text_preview"].startswith("Smith")
        )
        self.assertEqual(ref_info["indents"]["style_resolved"]["hanging"]["pt"], 36.0)
        self.assertEqual(
            result["font_audit"]["theme_font_attributes"]["theme_font_attribute_count"],
            0,
        )

    def test_doc_ops_font_audit_can_check_rendered_pdf_spans(self):
        path = self._path("font_rendered.docx")
        self._make_reference_fixture(path)
        payload = self._pdf_block_payload(good=True)
        payload["pages"][0]["blocks"][0]["lines"][1]["spans"] = [
            {
                "text": payload["pages"][0]["blocks"][0]["lines"][1]["text"],
                "font": "Times New Roman",
                "size": 12.0,
                "bbox": payload["pages"][0]["blocks"][0]["lines"][1]["bbox"],
            }
        ]
        payload["pages"][0]["blocks"][0]["lines"][2]["spans"] = [
            {
                "text": payload["pages"][0]["blocks"][0]["lines"][2]["text"],
                "font": "Times New Roman",
                "size": 12.0,
                "bbox": payload["pages"][0]["blocks"][0]["lines"][2]["bbox"],
            }
        ]

        with mock.patch.object(self.client, "pdf_read_blocks", return_value=payload):
            result = self.ops.audit_document_fonts(
                path,
                expected_font_name="Times New Roman",
                expected_font_size=12,
                rendered=True,
                pdf_path=self._path("font_rendered.pdf"),
            )

        self.assertTrue(result["success"], result)
        self.assertEqual(result["overall_status"], "pass")
        self.assertEqual(result["rendered"]["overall_status"], "pass")
        self.assertFalse(result["rendered"]["trusted_pdf"])
        self.assertEqual(result["rendered"]["unexpected_font_span_count"], 0)

    def test_doc_ops_rendered_layout_audit_ignores_repeated_page_headers(self):
        path = self._path("render_audit_headers.docx")
        self._make_reference_fixture(path)
        doc = Document(path)
        doc.sections[0].header.paragraphs[0].text = "Running header"
        doc.save(path)
        payload = {
            "success": True,
            "total_pages": 3,
            "pages_scanned": 3,
            "pages": [
                {
                    "page_index": 0,
                    "page_number": 1,
                    "width": 595.3,
                    "height": 841.9,
                    "blocks": [
                        {
                            "type": "text",
                            "lines": [
                                {
                                    "line_id": "body",
                                    "bbox": {"left": 72, "top": 90, "right": 240, "bottom": 106},
                                    "text": "Body before References.",
                                }
                            ],
                        }
                    ],
                },
                {
                    "page_index": 1,
                    "page_number": 2,
                    "width": 595.3,
                    "height": 841.9,
                    "blocks": [
                        {
                            "type": "text",
                            "lines": [
                                {
                                    "line_id": "header_2",
                                    "bbox": {"left": 72, "top": 34, "right": 190, "bottom": 48},
                                    "text": "Running header",
                                },
                                {
                                    "line_id": "heading",
                                    "bbox": {"left": 72, "top": 74, "right": 150, "bottom": 90},
                                    "text": "References",
                                },
                                {
                                    "line_id": "ref_1",
                                    "bbox": {"left": 72, "top": 112, "right": 520, "bottom": 128},
                                    "text": (
                                        "Del Re, A. C., Fluckiger, C., Horvath, A. O., & Wampold, B. E. "
                                        "(2021). Examining therapist"
                                    ),
                                },
                            ],
                        }
                    ],
                },
                {
                    "page_index": 2,
                    "page_number": 3,
                    "width": 595.3,
                    "height": 841.9,
                    "blocks": [
                        {
                            "type": "text",
                            "lines": [
                                {
                                    "line_id": "header_3",
                                    "bbox": {"left": 72, "top": 34, "right": 190, "bottom": 48},
                                    "text": "Running header",
                                },
                                {
                                    "line_id": "ref_2",
                                    "bbox": {"left": 108, "top": 90, "right": 520, "bottom": 106},
                                    "text": (
                                        "effects in the alliance-outcome relationship: A multilevel meta-analysis. "
                                        "Journal of Consulting"
                                    ),
                                },
                            ],
                        }
                    ],
                },
            ],
        }

        with mock.patch.object(self.client, "pdf_read_blocks", return_value=payload):
            result = self.ops.rendered_layout_audit(
                path,
                pdf_path=self._path("render_audit_headers.pdf"),
                trusted_pdf=True,
            )

        self.assertTrue(result["success"], result)
        statuses = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(statuses["references_page_break"], "pass")
        self.assertEqual(statuses["hanging_indents_rendered"], "pass")
        self.assertEqual(statuses["rendered_margins"], "pass")
        artifact_check = next(
            check
            for check in result["checks"]
            if check["name"] == "rendered_header_footer_artifacts"
        )
        self.assertEqual(artifact_check["details"]["ignored_line_count"], 2)

    def test_doc_ops_submission_pack_creates_clean_bundle_manifest(self):
        path = self._path("submission_pack.docx")
        self._make_reference_fixture(path)
        doc = Document(path)
        doc.sections[0].page_width = Inches(8.27)
        doc.sections[0].page_height = Inches(11.69)
        doc.save(path)
        output_dir = self._path("submission_pack_out")
        payload = self._pdf_block_payload(good=True)
        for line in payload["pages"][0]["blocks"][0]["lines"]:
            line["spans"] = [
                {
                    "text": line["text"],
                    "font": "Times New Roman",
                    "size": 12.0,
                    "bbox": line["bbox"],
                }
            ]

        def fake_office_to_pdf(file_path, output_path=None):
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            return {"success": True, "output_file": output_path, "pages": 2}

        def fake_pdf_sanitize(file_path, output_path=None, **kwargs):
            shutil.copy2(file_path, output_path)
            return {
                "success": True,
                "file": output_path,
                "output_file": output_path,
                "clear_metadata": kwargs.get("clear_metadata"),
                "remove_xml_metadata": kwargs.get("remove_xml_metadata"),
            }

        clean_pdf_hidden = {
            "success": True,
            "nonempty_metadata": {},
            "has_xml_metadata": False,
            "annotations_count": 0,
            "embedded_files_count": 0,
            "has_forms": False,
            "pages": 2,
        }

        with mock.patch.object(self.client, "_office_to_pdf", side_effect=fake_office_to_pdf):
            with mock.patch.object(self.client, "pdf_read_blocks", return_value=payload):
                with mock.patch.object(self.client, "pdf_sanitize", side_effect=fake_pdf_sanitize):
                    with mock.patch.object(
                        self.client,
                        "inspect_pdf_hidden_data",
                        return_value=clean_pdf_hidden,
                    ):
                        result = self.ops.submission_pack(
                            path,
                            output_dir,
                            basename="final",
                            expected_page_size="A4",
                            expected_font_name="Times New Roman",
                            expected_font_size=12,
                            render_profile="apa-references",
                        )

        self.assertTrue(result["success"], result)
        self.assertTrue(result["submission_ready"], result["readiness_blockers"])
        self.assertTrue(os.path.exists(result["clean_docx"]))
        self.assertTrue(os.path.exists(result["pdf_file"]))
        self.assertTrue(os.path.exists(result["manifest_file"]))
        self.assertTrue(result["text_preserved"])
        self.assertTrue(result["font_audit"]["rendered"]["trusted_pdf"])
