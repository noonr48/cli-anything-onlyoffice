import io
import json
import os
import subprocess
import tempfile
import unittest
import xml.etree.ElementTree as ET
from zipfile import ZipFile, ZIP_DEFLATED
from contextlib import redirect_stdout
from unittest import mock

import fitz
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from PIL import Image

from cli_anything.onlyoffice.core import cli as cli_module
from cli_anything.onlyoffice.core.command_registry import (
    CLI_SCHEMA_VERSION,
    COMMAND_CATEGORIES,
    HELP_EXAMPLES,
    TOTAL_COMMANDS,
    VERSION,
    command_signature,
    command_usage,
)
from cli_anything.onlyoffice.utils.docserver import get_client


class OnlyOfficeProductionReadinessTests(unittest.TestCase):
    def setUp(self):
        self.client = get_client()
        self.tmpdir = tempfile.TemporaryDirectory(prefix="oo-prod-test-")
        self.base = self.tmpdir.name

    def tearDown(self):
        self.tmpdir.cleanup()

    def _path(self, name: str) -> str:
        return os.path.join(self.base, name)

    def test_xlsx_write_is_non_destructive_by_default(self):
        path = self._path("grades.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "Header"
        keep = wb.create_sheet("KeepSheet")
        keep["A1"] = "KeepMe"
        wb.save(path)

        result = self.client.write_spreadsheet(
            output_path=path,
            headers=["Student", "Score"],
            data=[["Alice", 91], ["Bob", 88]],
            sheet_name="Sheet1",
            overwrite_workbook=False,
        )
        self.assertTrue(result["success"])

        wb2 = load_workbook(path)
        self.assertIn("KeepSheet", wb2.sheetnames)
        self.assertEqual(wb2["KeepSheet"]["A1"].value, "KeepMe")

    def test_xlsx_write_can_explicitly_overwrite(self):
        path = self._path("grades_overwrite.xlsx")
        wb = Workbook()
        wb.active.title = "Legacy"
        wb.create_sheet("KeepSheet")
        wb.save(path)

        result = self.client.write_spreadsheet(
            output_path=path,
            headers=["Student", "Score"],
            data=[["Alice", 91]],
            sheet_name="NewSheet",
            overwrite_workbook=True,
        )
        self.assertTrue(result["success"])

        wb2 = load_workbook(path)
        self.assertEqual(wb2.sheetnames, ["NewSheet"])

    def test_formula_aware_calculation(self):
        path = self._path("formula_calc.xlsx")
        self.client.write_spreadsheet(
            output_path=path,
            headers=["A", "B", "Total"],
            data=[[10, 20, ""], [15, 25, ""]],
            sheet_name="Sheet1",
            overwrite_workbook=True,
        )
        self.assertTrue(
            self.client.add_formula(path, "C2", "=A2+B2", sheet_name="Sheet1")["success"]
        )
        self.assertTrue(
            self.client.add_formula(path, "C3", "=A3+B3", sheet_name="Sheet1")["success"]
        )

        raw = self.client.calculate_column(
            path, "C", "avg", sheet_name="Sheet1", include_formulas=False
        )
        self.assertFalse(raw["success"])

        formula = self.client.calculate_column(
            path, "C", "avg", sheet_name="Sheet1", include_formulas=True
        )
        self.assertTrue(formula["success"])
        self.assertEqual(formula["average"], 35.0)

    def test_chart_validation_rejects_invalid_ranges(self):
        path = self._path("chart_invalid.xlsx")
        self.client.write_spreadsheet(
            output_path=path,
            headers=["Student", "Score"],
            data=[["Alice", 90], ["Bob", 85]],
            sheet_name="Sheet1",
            overwrite_workbook=True,
        )

        result = self.client.create_chart(
            file_path=path,
            chart_type="bar",
            data_range="Z1:Z5",
            categories_range="A2:A3",
            title="Invalid",
        )
        self.assertFalse(result["success"])

    def test_backup_snapshot_is_created_for_mutation(self):
        path = self._path("paper.docx")
        create = self.client.create_document(path, "Title", "Intro")
        self.assertTrue(create["success"])

        append = self.client.append_to_document(path, "New body paragraph")
        self.assertTrue(append["success"])
        self.assertTrue(append.get("backup"))
        self.assertTrue(os.path.exists(append["backup"]))

    def test_doc_add_image_can_anchor_near_paragraph(self):
        path = self._path("anchored.docx")
        image_path = self._path("figure.png")

        doc = Document()
        doc.add_paragraph("Alpha")
        doc.add_paragraph("Beta")
        doc.save(path)

        Image.new("RGB", (40, 20), color="navy").save(image_path)

        result = self.client.add_image(
            path,
            image_path,
            width_inches=1.0,
            caption="Figure 1",
            paragraph_index=0,
            position="after",
        )
        self.assertTrue(result["success"])
        self.assertEqual(result["paragraph_index"], 0)
        self.assertEqual(result["position"], "after")

        updated = Document(path)
        self.assertEqual(updated.paragraphs[0].text, "Alpha")
        self.assertEqual(updated.paragraphs[2].text, "Figure 1")
        self.assertEqual(updated.paragraphs[3].text, "Beta")
        self.assertTrue(updated.paragraphs[1].paragraph_format.keep_with_next)
        self.assertTrue(updated.paragraphs[2].paragraph_format.keep_together)

    def test_preview_document_reuses_pdf_render_pipeline(self):
        path = self._path("preview.docx")
        self.client.create_document(path, "Title", "Intro")
        output_dir = self._path("previews")

        captured = {}

        def fake_doc_to_pdf(file_path, output_path=None):
            self.assertEqual(file_path, path)
            self.assertIsNotNone(output_path)
            self.assertFalse(os.path.exists(output_path))
            with open(output_path, "wb") as f:
                f.write(b"%PDF-1.4 fake")
            captured["pdf_path"] = output_path
            return {
                "success": True,
                "input_file": file_path,
                "output_file": output_path,
                "pages": 2,
            }

        def fake_pdf_page_to_image(file_path, render_dir, pages=None, dpi=150, fmt="png"):
            self.assertEqual(file_path, captured["pdf_path"])
            self.assertTrue(os.path.exists(file_path))
            self.assertEqual(render_dir, output_dir)
            self.assertEqual(pages, "0-1")
            self.assertEqual(dpi, 200)
            self.assertEqual(fmt, "jpg")
            return {
                "success": True,
                "total_pages": 2,
                "pages_rendered": 2,
                "images": [{"page": 0, "file": os.path.join(render_dir, "page_000.jpg")}],
            }

        with mock.patch.object(self.client, "doc_to_pdf", side_effect=fake_doc_to_pdf):
            with mock.patch.object(
                self.client, "pdf_page_to_image", side_effect=fake_pdf_page_to_image
            ):
                result = self.client.preview_document(
                    path, output_dir, pages="0-1", dpi=200, fmt="jpg"
                )

        self.assertTrue(result["success"])
        self.assertEqual(result["total_pages"], 2)
        self.assertEqual(result["pages_rendered"], 2)
        self.assertEqual(result["format"], "jpg")
        self.assertFalse(os.path.exists(captured["pdf_path"]))

    def test_doc_layout_supports_named_page_size(self):
        path = self._path("layout_a4.docx")
        self.client.create_document(path, "Title", "Body")

        result = self.client.set_page_layout(path, page_size="A4")

        self.assertTrue(result["success"])
        self.assertEqual(result["page_size"], "A4")

        doc = Document(path)
        section = doc.sections[0]
        width_mm = round(section.page_width / 36000.0, 1)
        height_mm = round(section.page_height / 36000.0, 1)
        self.assertAlmostEqual(width_mm, 210.0, places=1)
        self.assertAlmostEqual(height_mm, 297.0, places=1)

    def test_doc_layout_preserves_existing_orientation_when_only_size_changes(self):
        path = self._path("layout_preserve_orientation.docx")
        self.client.create_document(path, "Title", "Body")
        self.client.set_page_layout(path, orientation="landscape")

        result = self.client.set_page_layout(path, page_size="A4")

        self.assertTrue(result["success"])
        self.assertEqual(result["orientation"], "landscape")
        doc = Document(path)
        section = doc.sections[0]
        self.assertGreater(section.page_width, section.page_height)

    def test_doc_sanitize_removes_comments_and_clears_metadata(self):
        path = self._path("sanitize_comments.docx")
        self.client.create_document(path, "Title", "Body")
        self.client.set_metadata(
            path,
            author="Original Author",
            title="Assignment Draft",
            subject="Research Methods",
            keywords="draft,metadata",
            comments="Internal note",
            category="Assignments",
        )
        comment_result = self.client.add_comment(path, "Review note", 0)
        self.assertTrue(comment_result["success"])

        before = self.client.inspect_hidden_data(path)
        self.assertTrue(before["success"])
        self.assertTrue(before["comments_part_present"])
        self.assertGreaterEqual(before["comments_count"], 1)
        self.assertEqual(before["core_properties"]["author"], "Original Author")

        result = self.client.sanitize_document(
            path,
            remove_comments=True,
            clear_metadata=True,
            author="benbi",
        )

        self.assertTrue(result["success"])
        after = result["after"]
        self.assertFalse(after["comments_part_present"])
        self.assertEqual(after["comments_count"], 0)
        self.assertEqual(after["comment_reference_count"], 0)
        self.assertEqual(after["core_properties"]["author"], "benbi")
        self.assertEqual(after["core_properties"]["title"], "")
        self.assertTrue(after["remove_personal_information"])

    def test_doc_sanitize_accepts_revisions_and_removes_custom_xml(self):
        path = self._path("sanitize_revisions.docx")
        self.client.create_document(path, "Title", "Keep")

        with ZipFile(path, "r") as zin:
            files = {name: zin.read(name) for name in zin.namelist()}

        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
            "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        }
        ET.register_namespace("w", ns["w"])
        root = ET.fromstring(files["word/document.xml"])
        body = root.find("w:body", ns)
        para = body.find("w:p", ns)
        run = para.find("w:r", ns)
        para.remove(run)
        ins = ET.Element(f"{{{ns['w']}}}ins")
        ins.append(run)
        para.append(ins)
        deleted = ET.Element(f"{{{ns['w']}}}del")
        del_run = ET.SubElement(deleted, f"{{{ns['w']}}}r")
        del_text = ET.SubElement(del_run, f"{{{ns['w']}}}delText")
        del_text.text = "RemoveMe"
        para.append(deleted)
        files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        files["customXml/item1.xml"] = b'<?xml version="1.0" encoding="UTF-8"?><root>secret</root>'
        files["customXml/_rels/item1.xml.rels"] = (
            b'<?xml version="1.0" encoding="UTF-8"?>'
            b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'
        )
        files["customXml/itemProps1.xml"] = (
            b'<?xml version="1.0" encoding="UTF-8"?>'
            b'<ds:datastoreItem xmlns:ds="http://schemas.openxmlformats.org/officeDocument/2006/customXml"></ds:datastoreItem>'
        )
        ct_root = ET.fromstring(files["[Content_Types].xml"])
        override = ET.SubElement(ct_root, f"{{{ns['ct']}}}Override")
        override.set("PartName", "/customXml/item1.xml")
        override.set("ContentType", "application/xml")
        files["[Content_Types].xml"] = ET.tostring(ct_root, encoding="utf-8", xml_declaration=True)

        with ZipFile(path, "w", compression=ZIP_DEFLATED) as zout:
            for name, data in files.items():
                zout.writestr(name, data)

        before = self.client.inspect_hidden_data(path)
        self.assertTrue(before["tracked_changes_present"])
        self.assertGreater(before["custom_xml_part_count"], 0)

        result = self.client.sanitize_document(
            path,
            accept_revisions=True,
            remove_custom_xml=True,
        )
        self.assertTrue(result["success"])

        after = result["after"]
        self.assertFalse(after["tracked_changes_present"])
        self.assertEqual(after["custom_xml_part_count"], 0)
        self.assertEqual(after["custom_xml_support_part_count"], 0)
        doc = Document(path)
        self.assertEqual(doc.paragraphs[1].text, "Keep")

    def test_pdf_sanitize_clears_metadata_and_xmp(self):
        path = self._path("sanitize.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Submission copy")
        doc.set_metadata(
            {
                "author": "Original Author",
                "title": "Draft",
                "subject": "Subject",
                "keywords": "alpha,beta",
                "creator": "OnlyOffice",
                "producer": "OnlyOffice",
            }
        )
        if hasattr(doc, "set_xml_metadata"):
            doc.set_xml_metadata("<x:xmpmeta xmlns:x='adobe:ns:meta/'><rdf:RDF xmlns:rdf='http://www.w3.org/1999/02/22-rdf-syntax-ns#'></rdf:RDF></x:xmpmeta>")
        doc.save(path)
        doc.close()

        result = self.client.pdf_sanitize(
            path,
            clear_metadata=True,
            remove_xml_metadata=True,
            author="benbi",
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["after"]["author"], "benbi")
        self.assertEqual(result["after"]["title"], "")
        self.assertFalse(result["has_xml_metadata"])

    def test_doc_preflight_flags_mixed_fonts_and_metadata(self):
        path = self._path("preflight.docx")
        doc = Document()
        first = doc.add_paragraph()
        first.add_run("Times paragraph")
        second = doc.add_paragraph()
        run = second.add_run("Calibri paragraph")
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        doc.save(path)
        self.client.set_page_layout(path, page_size="A4")
        self.client.set_metadata(path, author="benbi", title="Draft")

        result = self.client.document_preflight(
            path,
            expected_page_size="A4",
            expected_font_name="Times New Roman",
            expected_font_size=11,
        )

        self.assertTrue(result["success"])
        self.assertEqual(result["overall_status"], "warn")
        check_status = {check["name"]: check["status"] for check in result["checks"]}
        self.assertEqual(check_status["page_size"], "pass")
        self.assertEqual(check_status["metadata"], "warn")
        self.assertEqual(check_status["font_names"], "warn")
        self.assertEqual(check_status["font_sizes"], "warn")
        self.assertEqual(result["font_audit"]["unexpected_font_run_count"], 2)
        self.assertEqual(result["font_audit"]["unexpected_size_run_count"], 2)

    def test_pdf_inspect_hidden_data_reports_metadata_and_annotations(self):
        path = self._path("inspect_hidden.pdf")
        doc = fitz.open()
        page = doc.new_page(width=595.0, height=842.0)
        page.insert_text((72, 72), "Annotated")
        page.add_highlight_annot(fitz.Rect(70, 60, 150, 84))
        doc.set_metadata(
            {
                "author": "benbi",
                "title": "Inspection",
                "subject": "PDF hidden data",
                "creator": "OnlyOffice CLI Tests",
            }
        )
        if hasattr(doc, "set_xml_metadata"):
            doc.set_xml_metadata(
                "<x:xmpmeta xmlns:x='adobe:ns:meta/'><rdf:RDF xmlns:rdf='http://www.w3.org/1999/02/22-rdf-syntax-ns#'></rdf:RDF></x:xmpmeta>"
            )
        doc.save(path)
        doc.close()

        result = self.client.inspect_pdf_hidden_data(path)

        self.assertTrue(result["success"])
        self.assertEqual(result["pages"], 1)
        self.assertEqual(result["page_size_labels"], ["A4"])
        self.assertTrue(result["page_size_consistent"])
        self.assertGreaterEqual(result["annotations_count"], 1)
        self.assertEqual(result["nonempty_metadata"]["author"], "benbi")
        self.assertTrue(result["has_xml_metadata"])

    def test_cli_help_exposes_hardened_commands(self):
        proc = subprocess.run(
            ["cli-anything-onlyoffice", "help", "--json"],
            capture_output=True,
            text=True,
            check=True,
        )
        payload = json.loads(proc.stdout)
        self.assertEqual(payload["version"], VERSION)
        self.assertEqual(payload["schema_version"], CLI_SCHEMA_VERSION)
        self.assertEqual(payload["total_commands"], TOTAL_COMMANDS)
        self.assertEqual(payload["command_count"], TOTAL_COMMANDS)
        self.assertEqual(payload["examples"], HELP_EXAMPLES)
        self.assertEqual(payload["categories"]["DOCUMENTS (.docx)"], COMMAND_CATEGORIES["DOCUMENTS (.docx)"])
        self.assertEqual(payload["categories"]["PDF (.pdf)"], COMMAND_CATEGORIES["PDF (.pdf)"])
        self.assertIn("commands", payload)
        self.assertIn("usage", payload)
        self.assertIn("capability_metadata", payload)
        docs = payload["categories"]["DOCUMENTS (.docx)"]
        sheet = payload["categories"]["SPREADSHEETS (.xlsx)"]
        self.assertIn(
            "doc-format <file> <paragraph_index> [--bold] [--italic] [--underline] [--font-name <name>] [--font-size <n>] [--color <hex>] [--align <left|center|right|justify>]",
            docs,
        )
        self.assertIn(
            "doc-preview <file> <output_dir> [--pages <range>] [--dpi <n>] [--format png|jpg]",
            docs,
        )
        self.assertIn("doc-render-map <file>", docs)
        self.assertIn(
            "doc-layout <file> [--size A4|Letter] [--orientation portrait|landscape] [--margin-* <in>] [--header <text>] [--page-numbers]",
            docs,
        )
        self.assertIn("doc-inspect-hidden-data <file>", docs)
        self.assertIn(command_signature("doc-sanitize"), docs)
        self.assertIn(
            "doc-preflight <file> [--expected-page-size <A4|Letter>] [--expected-font <name>] [--expected-font-size <pt>] [--rendered-layout] [--profile auto|generic|apa-references]",
            docs,
        )
        self.assertIn(
            "doc-formatting-info <file> [--all] [--start <n>] [--limit <n>]",
            docs,
        )
        self.assertIn(
            "doc-to-pdf <file> [output_path] [--layout-warnings] [--profile auto|generic|apa-references]",
            docs,
        )
        self.assertIn(
            "doc-render-audit <file> [--pdf <path>] [--tolerance-points <n>] [--profile auto|generic|apa-references]",
            docs,
        )
        self.assertIn(
            "xlsx-calc <file> <column> <operation> [--sheet <name>] [--include-formulas] [--strict-formulas]",
            sheet,
        )
        self.assertIn(
            "xlsx-preview <file> <output_dir> [--pages <range>] [--dpi <n>] [--format png|jpg]",
            sheet,
        )
        validation_signature = next(
            signature for signature in sheet if signature.startswith("xlsx-add-validation ")
        )
        self.assertIn("--no-blank", validation_signature)
        self.assertNotIn("--allow-blank", validation_signature)
        self.assertIn("--no-blank", payload["usage"]["xlsx-add-validation"])
        self.assertNotIn("--allow-blank", payload["usage"]["xlsx-add-validation"])
        self.assertEqual(
            payload["commands"]["xlsx-add-validation"]["usage"],
            command_usage("xlsx-add-validation"),
        )
        pdf = payload["categories"]["PDF (.pdf)"]
        self.assertIn(
            "pdf-read-blocks <file> [--pages <range>] [--no-spans] [--no-images] [--include-empty]",
            pdf,
        )
        self.assertIn(
            "pdf-search-blocks <file> <query> [--pages <range>] [--case-sensitive] [--no-spans]",
            pdf,
        )
        self.assertIn(command_signature("pdf-sanitize"), pdf)
        self.assertIn("pdf-inspect-hidden-data <file>", pdf)
        rdf = payload["categories"]["RDF (Knowledge Graphs)"]
        self.assertIn(command_signature("rdf-remove"), rdf)
        general = payload["categories"]["GENERAL"]
        self.assertIn(
            "editor-session <file> [--open] [--wait <sec>] [--activate]",
            general,
        )
        self.assertIn(
            "editor-capture <file> <output_image> [--backend auto|desktop|rendered] [--open] [--page <n>] [--range <A1:D20>] [--slide <n>] [--zoom-reset] [--zoom-in <n>] [--zoom-out <n>] [--crop x,y,w,h] [--wait <sec>] [--settle-ms <n>] [--dpi <n>] [--format png|jpg]",
            general,
        )

    def test_cli_status_reports_registry_version_and_count(self):
        proc = subprocess.run(
            ["cli-anything-onlyoffice", "status", "--json"],
            capture_output=True,
            text=True,
            check=True,
        )
        payload = json.loads(proc.stdout)
        self.assertTrue(payload["success"])
        self.assertEqual(payload["schema_version"], CLI_SCHEMA_VERSION)
        self.assertEqual(payload["version"], VERSION)
        self.assertEqual(payload["total_commands"], TOTAL_COMMANDS)
        self.assertEqual(payload["command_count"], TOTAL_COMMANDS)
        self.assertEqual(payload["registry"]["schema_version"], CLI_SCHEMA_VERSION)
        self.assertEqual(payload["registry"]["total_commands"], TOTAL_COMMANDS)
        self.assertEqual(payload["registry"]["category_counts"], payload["category_counts"])
        self.assertIn("dependencies", payload)
        self.assertIn("conversion", payload)
        self.assertIn("capability_metadata", payload)
        self.assertEqual(payload["dependencies"]["openpyxl"], payload["openpyxl"])
        self.assertIn("docker", payload["dependencies"])
        self.assertIn("onlyoffice_x2t", payload["dependencies"])
        self.assertIn("office_to_pdf", payload["capabilities"])
        self.assertEqual(
            payload["capability_metadata"]["openpyxl"]["available"],
            payload["openpyxl"],
        )
        self.assertEqual(
            payload["capability_metadata"]["xlsx_create"]["requires"],
            ["openpyxl"],
        )

    def test_cli_doc_usage_error_comes_from_registry(self):
        stdout = io.StringIO()
        with mock.patch("sys.argv", ["cli-anything-onlyoffice", "doc-create", "--json"]):
            with redirect_stdout(stdout):
                cli_module.main()
        payload = json.loads(stdout.getvalue())
        self.assertFalse(payload["success"])
        self.assertEqual(payload["error"], command_usage("doc-create"))

    def test_cli_general_usage_error_comes_from_registry(self):
        stdout = io.StringIO()
        with mock.patch("sys.argv", ["cli-anything-onlyoffice", "open", "--json"]):
            with redirect_stdout(stdout):
                cli_module.main()
        payload = json.loads(stdout.getvalue())
        self.assertFalse(payload["success"])
        self.assertEqual(payload["error"], command_usage("open"))

    def test_agent_style_open_alias_opens_spreadsheet(self):
        path = self._path("alias.xlsx")
        Workbook().save(path)
        stdout = io.StringIO()

        with mock.patch("cli_anything.onlyoffice.core.general_cli.subprocess.Popen") as popen:
            with mock.patch(
                "sys.argv", ["cli-anything-onlyoffice", "document.open", path, "--json"]
            ):
                with redirect_stdout(stdout):
                    cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertTrue(payload["success"])
        self.assertEqual(payload["requested_command"], "document.open")
        self.assertEqual(payload["resolved_command"], "open")
        self.assertEqual(payload["command_namespace"], "document")
        self.assertEqual(payload["detected_type"], "spreadsheet")
        popen.assert_called_once_with(
            ["onlyoffice-desktopeditors", os.path.abspath(path)], start_new_session=True
        )

    def test_agent_style_open_alias_preserves_metadata_on_error(self):
        missing = self._path("missing.xlsx")
        stdout = io.StringIO()

        with mock.patch(
            "sys.argv", ["cli-anything-onlyoffice", "document.open", missing, "--json"]
        ):
            with redirect_stdout(stdout):
                cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertFalse(payload["success"])
        self.assertEqual(payload["requested_command"], "document.open")
        self.assertEqual(payload["resolved_command"], "open")
        self.assertEqual(payload["command_namespace"], "document")
        self.assertIn("File not found", payload["error"])

    def test_agent_style_info_alias_resolves(self):
        path = self._path("alias_info.docx")
        self.client.create_document(path, "Alias", "Body")
        stdout = io.StringIO()

        with mock.patch(
            "sys.argv", ["cli-anything-onlyoffice", "spreadsheet.info", path, "--json"]
        ):
            with redirect_stdout(stdout):
                cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertTrue(payload["success"])
        self.assertEqual(payload["requested_command"], "spreadsheet.info")
        self.assertEqual(payload["resolved_command"], "info")
        self.assertEqual(payload["command_namespace"], "spreadsheet")

    def test_pdf_read_blocks_returns_native_block_span_metadata(self):
        path = self._path("native_blocks.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Alpha Results")
        page.insert_text((72, 120), "Beta Findings")
        doc.save(path)
        doc.close()

        result = self.client.pdf_read_blocks(path, pages="0")

        self.assertTrue(result["success"])
        self.assertEqual(result["pages_scanned"], 1)
        self.assertGreaterEqual(result["text_block_count"], 1)
        first_page = result["pages"][0]
        self.assertEqual(first_page["page_index"], 0)
        text_blocks = [block for block in first_page["blocks"] if block["type"] == "text"]
        self.assertTrue(text_blocks)
        first_block = text_blocks[0]
        self.assertIn("bbox", first_block)
        self.assertTrue(first_block["block_id"].startswith("page_0_block_"))
        self.assertGreaterEqual(first_block["line_count"], 1)
        first_line = first_block["lines"][0]
        self.assertTrue(first_line["line_id"].startswith(first_block["block_id"]))
        self.assertIn("bbox", first_line)
        first_span = first_line["spans"][0]
        self.assertTrue(first_span["span_id"].startswith(first_line["line_id"]))
        self.assertIn("bbox", first_span)
        self.assertTrue(any(token in first_block["text"] for token in ("Alpha", "Results", "Beta", "Findings")))

    def test_doc_render_map_maps_paragraphs_and_table_cells(self):
        path = self._path("render_map.docx")
        doc = Document()
        doc.add_paragraph("Executive summary")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "42"
        doc.save(path)

        def fake_doc_to_pdf(file_path, output_path=None):
            self.assertTrue(file_path.endswith(".docx"))
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as handle:
                handle.write(b"%PDF-1.4 fake")
            return {"success": True, "output_file": output_path, "pages": 1}

        def fake_pdf_read_blocks(file_path, pages=None, include_spans=True, include_images=False, include_empty=False):
            self.assertTrue(file_path.endswith(".pdf"))
            self.assertTrue(os.path.exists(file_path))
            self.assertTrue(include_spans)
            return {
                "success": True,
                "pages_scanned": 1,
                "total_pages": 1,
                "pages": [
                    {
                        "page_index": 0,
                        "page_number": 1,
                        "blocks": [
                            {
                                "block_id": "page_0_block_0",
                                "type": "text",
                                "bbox": {"left": 10, "top": 10, "right": 200, "bottom": 40, "width": 190, "height": 30},
                                "lines": [
                                    {
                                        "line_id": "page_0_block_0_line_0",
                                        "bbox": {"left": 10, "top": 10, "right": 200, "bottom": 25, "width": 190, "height": 15},
                                        "spans": [
                                            {
                                                "span_id": "page_0_block_0_line_0_span_0",
                                                "text": "SLOANE_P_0001 Executive summary",
                                                "bbox": {"left": 10, "top": 10, "right": 200, "bottom": 25, "width": 190, "height": 15},
                                            }
                                        ],
                                    }
                                ],
                            },
                            {
                                "block_id": "page_0_block_1",
                                "type": "text",
                                "bbox": {"left": 10, "top": 50, "right": 200, "bottom": 80, "width": 190, "height": 30},
                                "lines": [
                                    {
                                        "line_id": "page_0_block_1_line_0",
                                        "bbox": {"left": 10, "top": 50, "right": 200, "bottom": 65, "width": 190, "height": 15},
                                        "spans": [
                                            {
                                                "span_id": "page_0_block_1_line_0_span_0",
                                                "text": "SLOANE_T1R1C1 42",
                                                "bbox": {"left": 10, "top": 50, "right": 200, "bottom": 65, "width": 190, "height": 15},
                                            }
                                        ],
                                    }
                                ],
                            },
                        ],
                    }
                ],
            }

        with mock.patch.object(self.client, "doc_to_pdf", side_effect=fake_doc_to_pdf):
            with mock.patch.object(self.client, "pdf_read_blocks", side_effect=fake_pdf_read_blocks):
                result = self.client.doc_render_map(path)

        self.assertTrue(result["success"])
        self.assertEqual(result["mapped_paragraph_count"], 1)
        self.assertEqual(result["mapped_table_cell_count"], 1)
        self.assertEqual(result["paragraphs"][0]["page_number"], 1)
        self.assertEqual(result["paragraphs"][0]["paragraph_index"], 1)
        self.assertEqual(result["table_cells"][0]["cell_ref"], "T1R1C1")
        self.assertEqual(result["table_cells"][0]["page_number"], 1)
        self.assertFalse(result["unresolved_anchor_ids"])

    def test_cli_pdf_search_blocks_dispatches_to_docserver(self):
        path = self._path("dispatch.pdf")
        with open(path, "wb") as handle:
            handle.write(b"%PDF-1.4")
        stdout = io.StringIO()

        with mock.patch.object(
            cli_module.doc_server,
            "pdf_search_blocks",
            return_value={
                "success": True,
                "file": path,
                "query": "Results",
                "match_count": 1,
                "matches": [{"page_index": 0, "block_id": "page_0_block_0"}],
            },
        ) as search_blocks:
            with mock.patch(
                "sys.argv",
                [
                    "cli-anything-onlyoffice",
                    "pdf-search-blocks",
                    path,
                    "Results",
                    "--pages",
                    "0-1",
                    "--case-sensitive",
                    "--json",
                ],
            ):
                with redirect_stdout(stdout):
                    cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertTrue(payload["success"])
        self.assertEqual(payload["match_count"], 1)
        search_blocks.assert_called_once_with(
            path,
            "Results",
            pages="0-1",
            case_sensitive=True,
            include_spans=True,
        )

    def test_cli_doc_sanitize_dispatches_to_docserver(self):
        path = self._path("dispatch.docx")
        self.client.create_document(path, "Dispatch", "Body")
        stdout = io.StringIO()

        with mock.patch.object(
            cli_module.doc_server,
            "sanitize_document",
            return_value={"success": True, "file": path, "after": {}},
        ) as sanitize_document:
            with mock.patch(
                "sys.argv",
                [
                    "cli-anything-onlyoffice",
                    "doc-sanitize",
                    path,
                    "--remove-comments",
                    "--clear-metadata",
                    "--author",
                    "benbi",
                    "--json",
                ],
            ):
                with redirect_stdout(stdout):
                    cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertTrue(payload["success"])
        sanitize_document.assert_called_once_with(
            path,
            output_path=None,
            remove_comments=True,
            accept_revisions=False,
            clear_metadata=True,
            remove_custom_xml=False,
            set_remove_personal_information=False,
            canonicalize_ooxml=False,
            author="benbi",
            title=None,
            subject=None,
            keywords=None,
        )

    def test_cli_doc_preflight_dispatches_to_docserver(self):
        path = self._path("dispatch_preflight.docx")
        self.client.create_document(path, "Dispatch", "Body")
        stdout = io.StringIO()

        with mock.patch.object(
            cli_module.doc_server,
            "document_preflight",
            return_value={"success": True, "file": path, "overall_status": "pass", "checks": []},
        ) as document_preflight:
            with mock.patch(
                "sys.argv",
                [
                    "cli-anything-onlyoffice",
                    "doc-preflight",
                    path,
                    "--expected-page-size",
                    "A4",
                    "--expected-font",
                    "Times New Roman",
                    "--expected-font-size",
                    "12",
                    "--rendered-layout",
                    "--profile",
                    "apa-references",
                    "--json",
                ],
            ):
                with redirect_stdout(stdout):
                    cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertTrue(payload["success"])
        document_preflight.assert_called_once_with(
            path,
            expected_page_size="A4",
            expected_font_name="Times New Roman",
            expected_font_size=12.0,
            rendered_layout=True,
            render_profile="apa-references",
        )

    def test_cli_pdf_sanitize_dispatches_to_docserver(self):
        path = self._path("dispatch_meta.pdf")
        with open(path, "wb") as handle:
            handle.write(b"%PDF-1.4")
        stdout = io.StringIO()

        with mock.patch.object(
            cli_module.doc_server,
            "pdf_sanitize",
            return_value={"success": True, "file": path, "after": {}},
        ) as pdf_sanitize:
            with mock.patch(
                "sys.argv",
                [
                    "cli-anything-onlyoffice",
                    "pdf-sanitize",
                    path,
                    "--clear-metadata",
                    "--remove-xml-metadata",
                    "--author",
                    "benbi",
                    "--json",
                ],
            ):
                with redirect_stdout(stdout):
                    cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertTrue(payload["success"])
        pdf_sanitize.assert_called_once_with(
            path,
            output_path=None,
            clear_metadata=True,
            remove_xml_metadata=True,
            author="benbi",
            title=None,
            subject=None,
            keywords=None,
            creator=None,
            producer=None,
        )

    def test_cli_pdf_inspect_hidden_data_dispatches_to_docserver(self):
        path = self._path("dispatch_hidden.pdf")
        with open(path, "wb") as handle:
            handle.write(b"%PDF-1.4")
        stdout = io.StringIO()

        with mock.patch.object(
            cli_module.doc_server,
            "inspect_pdf_hidden_data",
            return_value={"success": True, "file": path, "pages": 1},
        ) as inspect_pdf_hidden_data:
            with mock.patch(
                "sys.argv",
                [
                    "cli-anything-onlyoffice",
                    "pdf-inspect-hidden-data",
                    path,
                    "--json",
                ],
            ):
                with redirect_stdout(stdout):
                    cli_module.main()

        payload = json.loads(stdout.getvalue())
        self.assertTrue(payload["success"])
        inspect_pdf_hidden_data.assert_called_once_with(path)

    def test_preview_spreadsheet_reuses_pdf_render_pipeline(self):
        path = self._path("preview.xlsx")
        Workbook().save(path)
        output_dir = self._path("sheet_previews")

        captured = {}

        def fake_xlsx_to_pdf(file_path, output_path=None):
            self.assertEqual(file_path, path)
            self.assertIsNotNone(output_path)
            with open(output_path, "wb") as f:
                f.write(b"%PDF-1.4 fake")
            captured["pdf_path"] = output_path
            return {
                "success": True,
                "input_file": file_path,
                "output_file": output_path,
                "pages": 1,
            }

        def fake_pdf_page_to_image(file_path, render_dir, pages=None, dpi=150, fmt="png"):
            self.assertEqual(file_path, captured["pdf_path"])
            self.assertTrue(os.path.exists(file_path))
            self.assertEqual(render_dir, output_dir)
            self.assertEqual(pages, "0")
            self.assertEqual(dpi, 180)
            self.assertEqual(fmt, "png")
            return {
                "success": True,
                "total_pages": 1,
                "pages_rendered": 1,
                "images": [{"page": 0, "file": os.path.join(render_dir, "page_000.png")}],
            }

        with mock.patch.object(
            self.client, "spreadsheet_to_pdf", side_effect=fake_xlsx_to_pdf
        ):
            with mock.patch.object(
                self.client, "pdf_page_to_image", side_effect=fake_pdf_page_to_image
            ):
                result = self.client.preview_spreadsheet(
                    path, output_dir, pages="0", dpi=180, fmt="png"
                )

        self.assertTrue(result["success"])
        self.assertEqual(result["total_pages"], 1)
        self.assertEqual(result["pages_rendered"], 1)
        self.assertEqual(result["format"], "png")
        self.assertFalse(os.path.exists(captured["pdf_path"]))

    def test_capture_editor_view_uses_rendered_backend(self):
        path = self._path("rendered_capture.xlsx")
        Workbook().save(path)
        output = self._path("capture.png")
        preview_dir = self._path("rendered_backend")
        os.makedirs(preview_dir, exist_ok=True)
        preview_image = os.path.join(preview_dir, "page_000.png")
        Image.new("RGB", (200, 100), color="green").save(preview_image)

        with mock.patch.object(
            self.client,
            "preview_spreadsheet",
            return_value={
                "success": True,
                "images": [{"page": 0, "file": preview_image}],
                "total_pages": 1,
                "pages_rendered": 1,
            },
        ) as preview:
            result = self.client.capture_editor_view(
                path, output, backend="rendered", page=0, crop="10,10,50,40"
            )

        self.assertTrue(result["success"])
        self.assertEqual(result["backend"], "rendered")
        self.assertFalse(result["exact_viewport"])
        self.assertEqual((result["width"], result["height"]), (50, 40))
        self.assertTrue(os.path.exists(output))
        preview.assert_called_once()

    def test_capture_editor_view_refuses_unverified_spectacle_focus(self):
        path = self._path("desktop_capture.xlsx")
        Workbook().save(path)
        output = self._path("desktop_capture.png")

        with mock.patch.object(
            self.client,
            "editor_session",
            return_value={
                "success": True,
                "backend": "desktop",
                "file": path,
                "type": "spreadsheet",
                "window_id": 4242,
                "geometry": {"x": 0, "y": 0, "width": 1280, "height": 720},
            },
        ):
            with mock.patch.object(
                self.client, "_desktop_apply_viewport", return_value=[]
            ):
                with mock.patch.object(
                    self.client,
                    "_desktop_capture_tools",
                    return_value={"available": True, "capture_tool": "spectacle"},
                ):
                    with mock.patch.object(
                        self.client,
                        "_desktop_ensure_active_window",
                        return_value=False,
                    ) as ensure_active:
                        result = self.client.capture_editor_view(
                            path, output, backend="desktop"
                        )

        self.assertFalse(result["success"])
        self.assertIn("target OnlyOffice window is active", result["error"])
        ensure_active.assert_called_once_with(4242)
        self.assertFalse(os.path.exists(output))

    def test_editor_session_missing_window_explains_next_steps(self):
        path = self._path("missing_window.docx")
        self.client.create_document(path, "Title", "Body")

        with mock.patch.object(
            self.client,
            "_desktop_capture_tools",
            return_value={
                "available": True,
                "xdotool": "/usr/bin/xdotool",
                "xprop": "/usr/bin/xprop",
                "onlyoffice-desktopeditors": "/usr/bin/onlyoffice-desktopeditors",
                "capture_tool": "spectacle",
            },
        ):
            with mock.patch.object(
                self.client, "_desktop_find_editor_window", return_value={}
            ):
                result = self.client.editor_session(path, open_if_needed=False)

        self.assertFalse(result["success"])
        self.assertIn("Desktop capture needs a live editor window", result["error"])
        self.assertIn("--backend rendered", result["error"])

    def test_strict_schema_validation_rejects_row_width_mismatch(self):
        path = self._path("schema_strict.xlsx")
        result = self.client.write_spreadsheet(
            output_path=path,
            headers=["A", "B", "C"],
            data=[[1, 2], [3, 4, 5]],
            sheet_name="Sheet1",
            overwrite_workbook=True,
            coerce_rows=False,
        )
        self.assertFalse(result["success"])
        self.assertIn("expected 3", result["error"])

    def test_coerce_rows_allows_width_normalization(self):
        path = self._path("schema_coerce.xlsx")
        result = self.client.write_spreadsheet(
            output_path=path,
            headers=["A", "B", "C"],
            data=[[1, 2], [3, 4, 5, 6]],
            sheet_name="Sheet1",
            overwrite_workbook=True,
            coerce_rows=True,
        )
        self.assertTrue(result["success"])
        wb = load_workbook(path)
        ws = wb["Sheet1"]
        self.assertIn(ws["C2"].value, ("", None))
        self.assertEqual(ws["C3"].value, 5)

    def test_backup_list_restore_and_prune(self):
        path = self._path("restore_target.docx")
        self.client.create_document(path, "Title", "Version 1")
        self.client.append_to_document(path, "Version 2")
        self.client.append_to_document(path, "Version 3")

        listing = self.client.list_backups(path, limit=10)
        self.assertTrue(listing["success"])
        self.assertGreaterEqual(listing["count"], 1)

        before_restore = self.client.read_document(path)
        self.assertIn("Version 3", before_restore["full_text"])

        restore = self.client.restore_backup(
            path, backup=listing["backups"][-1]["name"]
        )
        self.assertTrue(restore["success"])
        after_restore = self.client.read_document(path)
        self.assertNotIn("Version 3", after_restore["full_text"])

        pruned = self.client.prune_backups(file_path=path, keep=1)
        self.assertTrue(pruned["success"])
        listing2 = self.client.list_backups(path, limit=10)
        self.assertLessEqual(listing2["count"], 1)


if __name__ == "__main__":
    unittest.main()
