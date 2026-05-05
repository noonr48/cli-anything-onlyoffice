#!/usr/bin/env python3
"""DOCX submission/runtime operations for the OnlyOffice CLI."""

from __future__ import annotations

import json
import hashlib
import io
import math
import os
import posixpath
import re
import shutil
import tempfile
import warnings
import xml.etree.ElementTree as ET
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from zipfile import BadZipFile, ZipFile

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


OOXML_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16du": "http://schemas.microsoft.com/office/word/2023/wordml/word16du",
    "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
    "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
    "xsi": "http://www.w3.org/2001/XMLSchema-instance",
    "ct": "http://schemas.openxmlformats.org/package/2006/content-types",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    "ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties",
    "vt": "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes",
    "ds": "http://schemas.openxmlformats.org/officeDocument/2006/customXml",
}


def register_ooxml_namespaces() -> None:
    """Register stable OOXML prefixes before ElementTree serialisation.

    The OOXML spec permits arbitrary prefixes, but OnlyOffice x2t 9.3 has been
    observed to ignore some WordprocessingML layout properties after document.xml
    is round-tripped with generic ns0/ns1 prefixes. Keep canonical prefixes for
    DOCX story/settings parts we rewrite.
    """
    for prefix, uri in OOXML_NS.items():
        ET.register_namespace(prefix, uri)


def serialize_ooxml_element(
    root: ET.Element,
    *,
    encoding: str = "utf-8",
    xml_declaration: bool = True,
) -> bytes:
    """Serialise an OOXML ElementTree element with stable namespace prefixes."""
    register_ooxml_namespaces()
    serialized = ET.tostring(root, encoding=encoding, xml_declaration=xml_declaration)
    return repair_mc_ignorable_declarations(serialized)


def repair_mc_ignorable_declarations(xml_blob: bytes) -> bytes:
    """Restore known namespace declarations referenced only by mc:Ignorable."""
    try:
        text = xml_blob.decode("utf-8")
    except UnicodeDecodeError:
        return xml_blob

    ignorable_tokens = set()
    for match in re.finditer(
        r"\b(?:mc:)?Ignorable\s*=\s*([\"'])(.*?)\1",
        text,
        flags=re.DOTALL,
    ):
        ignorable_tokens.update(
            token
            for token in re.split(r"\s+", match.group(2).strip())
            if token and token in OOXML_NS
        )
    missing = [
        prefix
        for prefix in sorted(ignorable_tokens)
        if not re.search(rf"\bxmlns:{re.escape(prefix)}\s*=", text)
    ]
    if not missing:
        return xml_blob

    root_match = re.search(r"<(?![?!/])([A-Za-z_][\w.\-]*:)?[A-Za-z_][\w.\-]*\b", text)
    if not root_match:
        return xml_blob
    insert_at = root_match.end()
    declarations = "".join(
        f' xmlns:{prefix}="{OOXML_NS[prefix]}"' for prefix in missing
    )
    return (text[:insert_at] + declarations + text[insert_at:]).encode("utf-8")


class DocumentOperations:
    """Encapsulate DOCX submission, sanitization, and rendering workflows."""

    MAX_EXTRACT_IMAGES = 1_000
    MAX_EXTRACT_IMAGE_COMPRESSED_BYTES = 50 * 1024 * 1024
    MAX_EXTRACT_IMAGE_PIXELS = 64_000_000
    MAX_DOCX_ZIP_ENTRIES = 10_000
    MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES = 512 * 1024 * 1024
    MAX_DOCX_XML_PART_BYTES = 32 * 1024 * 1024
    MAX_DOCX_COMPRESSION_RATIO = 100.0
    MAX_OUTPUT_PREFIX_LENGTH = 80
    IMAGE_OUTPUT_FORMATS = {
        "png": ("png", "PNG"),
        "jpg": ("jpg", "JPEG"),
        "jpeg": ("jpg", "JPEG"),
    }
    OOXML_XML_PART_SUFFIXES = (".xml", ".rels")
    SAFE_PREFIX_RE = re.compile(r"^[A-Za-z0-9_. -]+$")
    VECTOR_IMAGE_CONTENT_TYPES = {
        "image/x-emf": "emf",
        "image/x-wmf": "wmf",
    }
    TRACKED_REVISION_NODES = {
        "cellDel",
        "cellIns",
        "cellMerge",
        "customXmlDelRangeEnd",
        "customXmlDelRangeStart",
        "customXmlInsRangeEnd",
        "customXmlInsRangeStart",
        "customXmlMoveFromRangeEnd",
        "customXmlMoveFromRangeStart",
        "customXmlMoveToRangeEnd",
        "customXmlMoveToRangeStart",
        "del",
        "delInstrText",
        "delText",
        "ins",
        "moveFrom",
        "moveFromRangeEnd",
        "moveFromRangeStart",
        "moveTo",
        "moveToRangeEnd",
        "moveToRangeStart",
        "numberingChange",
        "pPrChange",
        "rPrChange",
        "sectPrChange",
        "tblGridChange",
        "tblPrChange",
        "tblPrExChange",
        "tcPrChange",
        "trPrChange",
    }
    REVISION_PROPERTY_CHANGE_NODES = {
        "numberingChange",
        "pPrChange",
        "rPrChange",
        "sectPrChange",
        "tblGridChange",
        "tblPrChange",
        "tblPrExChange",
        "tcPrChange",
        "trPrChange",
    }
    EXTERNAL_REL_ALT_CHUNK = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/altChunk"
    )
    EXTERNAL_REL_ATTACHED_TEMPLATE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/attachedTemplate"
    )
    EXTERNAL_REL_HYPERLINK = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
    )
    EXTERNAL_REL_IMAGE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
    )
    EXTERNAL_REL_OLE_OBJECT = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
    )
    EXTERNAL_REL_PACKAGE = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/package"
    )

    def __init__(self, host: Any):
        self.host = host

    @staticmethod
    def _sha256_file(file_path: str) -> str:
        digest = hashlib.sha256()
        with open(file_path, "rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    @staticmethod
    def _normalize_text(text: str) -> str:
        return re.sub(r"\s+", " ", str(text or "")).strip()

    @staticmethod
    def _length_points(value: Any) -> Optional[float]:
        if value is None:
            return None
        if hasattr(value, "pt"):
            return float(value.pt)
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    @classmethod
    def _length_payload(cls, value: Any) -> Optional[Dict[str, Any]]:
        points = cls._length_points(value)
        if points is None:
            return None
        return cls._points_payload(points)

    @staticmethod
    def _points_payload(points: float) -> Dict[str, Any]:
        return {
            "pt": round(points, 3),
            "in": round(points / 72.0, 4),
            "twips": int(round(points * 20)),
        }

    @staticmethod
    def _twips_payload(value: Any) -> Optional[Dict[str, Any]]:
        if value is None:
            return None
        try:
            twips = int(value)
        except (TypeError, ValueError):
            return None
        points = twips / 20.0
        return {
            "twips": twips,
            "pt": round(points, 3),
            "in": round(points / 72.0, 4),
        }

    @staticmethod
    def _local_attrs(element: Any) -> Dict[str, str]:
        if element is None:
            return {}
        return {
            str(key).rsplit("}", 1)[-1]: value
            for key, value in element.attrib.items()
        }

    @classmethod
    def _docx_zip_resource_limits(cls) -> Dict[str, Any]:
        return {
            "max_entries": cls.MAX_DOCX_ZIP_ENTRIES,
            "max_total_uncompressed_bytes": cls.MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES,
            "max_xml_part_bytes": cls.MAX_DOCX_XML_PART_BYTES,
            "max_compression_ratio": cls.MAX_DOCX_COMPRESSION_RATIO,
        }

    @staticmethod
    def _unsafe_ooxml_part_name_reason(part_name: str) -> Optional[str]:
        name = str(part_name or "")
        if not name:
            return "part name is empty"
        if "\x00" in name:
            return "part name contains a NUL byte"
        if name.startswith(("/", "\\")):
            return "part name must be relative, not absolute"
        if re.match(r"^[A-Za-z]:", name):
            return "part name must not use a Windows drive prefix"
        if "\\" in name:
            return "part name must use forward slashes"
        parts = name.split("/")
        if any(part == ".." for part in parts):
            return "part name must not contain '..' segments"
        return None

    @classmethod
    def _docx_zip_preflight(cls, file_path: str) -> Dict[str, Any]:
        abs_path = str(Path(file_path).resolve())
        limits = cls._docx_zip_resource_limits()
        preflight: Dict[str, Any] = {
            "success": True,
            "status": "pass",
            "file": abs_path,
            "limits": limits,
            "entry_count": 0,
            "total_uncompressed_bytes": 0,
            "total_compressed_bytes": 0,
            "xml_part_count": 0,
            "max_observed_compression_ratio": 0.0,
            "duplicate_names": [],
            "unsafe_part_names": [],
            "oversized_xml_parts": [],
            "compression_ratio_violations": [],
            "errors": [],
        }

        def add_error(code: str, message: str, **details: Any) -> None:
            entry = {"code": code, "message": message}
            entry.update(details)
            preflight["errors"].append(entry)

        try:
            with ZipFile(abs_path, "r") as zin:
                infos = zin.infolist()
        except BadZipFile as exc:
            add_error(
                "invalid_zip",
                f"File is not a valid ZIP/DOCX package: {exc}",
            )
            preflight["success"] = False
            preflight["status"] = "fail"
            return preflight
        except Exception as exc:
            add_error(
                "zip_open_failed",
                f"Could not open DOCX package for preflight: {exc}",
            )
            preflight["success"] = False
            preflight["status"] = "fail"
            return preflight

        seen = set()
        duplicate_names = []
        unsafe_part_names = []
        oversized_xml_parts = []
        compression_ratio_violations = []
        total_uncompressed = 0
        total_compressed = 0
        max_observed_ratio = 0.0
        xml_part_count = 0

        for info in infos:
            name = info.filename
            if name in seen and name not in duplicate_names:
                duplicate_names.append(name)
            seen.add(name)

            reason = cls._unsafe_ooxml_part_name_reason(name)
            if reason:
                unsafe_part_names.append({"part": name, "reason": reason})

            file_size = int(getattr(info, "file_size", 0) or 0)
            compressed_size = int(getattr(info, "compress_size", 0) or 0)
            total_uncompressed += file_size
            total_compressed += compressed_size

            lower_name = name.lower()
            if lower_name.endswith(cls.OOXML_XML_PART_SUFFIXES):
                xml_part_count += 1
                if file_size > cls.MAX_DOCX_XML_PART_BYTES:
                    oversized_xml_parts.append(
                        {
                            "part": name,
                            "size_bytes": file_size,
                            "max_size_bytes": cls.MAX_DOCX_XML_PART_BYTES,
                        }
                    )

            if file_size > 0:
                ratio = math.inf if compressed_size <= 0 else file_size / compressed_size
                if math.isfinite(ratio):
                    max_observed_ratio = max(max_observed_ratio, ratio)
                if ratio > cls.MAX_DOCX_COMPRESSION_RATIO:
                    compression_ratio_violations.append(
                        {
                            "part": name,
                            "compressed_size_bytes": compressed_size,
                            "uncompressed_size_bytes": file_size,
                            "compression_ratio": (
                                "inf" if not math.isfinite(ratio) else round(ratio, 3)
                            ),
                            "max_compression_ratio": cls.MAX_DOCX_COMPRESSION_RATIO,
                        }
                    )

        preflight.update(
            {
                "entry_count": len(infos),
                "total_uncompressed_bytes": total_uncompressed,
                "total_compressed_bytes": total_compressed,
                "xml_part_count": xml_part_count,
                "max_observed_compression_ratio": round(max_observed_ratio, 3),
                "duplicate_names": sorted(duplicate_names),
                "unsafe_part_names": unsafe_part_names,
                "oversized_xml_parts": oversized_xml_parts,
                "compression_ratio_violations": compression_ratio_violations,
            }
        )

        if len(infos) > cls.MAX_DOCX_ZIP_ENTRIES:
            add_error(
                "too_many_zip_entries",
                (
                    f"DOCX package has {len(infos)} entries; "
                    f"max is {cls.MAX_DOCX_ZIP_ENTRIES}."
                ),
                entry_count=len(infos),
                max_entries=cls.MAX_DOCX_ZIP_ENTRIES,
            )
        if total_uncompressed > cls.MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES:
            add_error(
                "total_uncompressed_size_exceeded",
                (
                    f"DOCX package expands to {total_uncompressed} bytes; "
                    f"max is {cls.MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES}."
                ),
                total_uncompressed_bytes=total_uncompressed,
                max_total_uncompressed_bytes=cls.MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES,
            )
        if duplicate_names:
            add_error(
                "duplicate_part_name",
                "DOCX package contains duplicate part names.",
                part_names=sorted(duplicate_names),
            )
        if unsafe_part_names:
            add_error(
                "unsafe_part_name",
                "DOCX package contains unsafe part names.",
                unsafe_part_names=unsafe_part_names,
            )
        if oversized_xml_parts:
            add_error(
                "xml_part_too_large",
                "DOCX package contains XML parts exceeding the configured size cap.",
                oversized_xml_parts=oversized_xml_parts,
            )
        if compression_ratio_violations:
            add_error(
                "compression_ratio_exceeded",
                "DOCX package contains entries exceeding the compression-ratio cap.",
                compression_ratio_violations=compression_ratio_violations,
            )

        if preflight["errors"]:
            preflight["success"] = False
            preflight["status"] = "fail"
        return preflight

    @classmethod
    def _docx_preflight_error_result(
        cls,
        file_path: str,
        preflight: Dict[str, Any],
    ) -> Dict[str, Any]:
        messages = [
            str(error.get("message", error.get("code", "preflight failed")))
            for error in preflight.get("errors", [])
        ]
        suffix = f": {'; '.join(messages[:3])}" if messages else ""
        return {
            "success": False,
            "file": str(Path(file_path).resolve()),
            "error": f"DOCX OOXML package preflight failed{suffix}",
            "error_code": "docx_ooxml_preflight_failed",
            "preflight": preflight,
        }

    @classmethod
    def _read_docx_zip_parts(
        cls,
        file_path: str,
    ) -> Tuple[Optional[Dict[str, bytes]], Dict[str, Any], Optional[Dict[str, Any]]]:
        preflight = cls._docx_zip_preflight(file_path)
        if not preflight.get("success"):
            return None, preflight, cls._docx_preflight_error_result(file_path, preflight)
        with ZipFile(str(Path(file_path).resolve()), "r") as zin:
            files = {info.filename: zin.read(info) for info in zin.infolist()}
        return files, preflight, None

    @staticmethod
    def _docx_relationship_source_part(rel_part_name: str) -> Optional[str]:
        if not rel_part_name.endswith(".rels"):
            return None
        rel_dir = posixpath.dirname(rel_part_name)
        rel_filename = posixpath.basename(rel_part_name)
        source_filename = rel_filename[: -len(".rels")]
        if rel_dir == "_rels":
            return source_filename
        if not rel_dir.endswith("/_rels"):
            return None
        source_dir = rel_dir[: -len("/_rels")]
        return posixpath.normpath(posixpath.join(source_dir, source_filename))

    @classmethod
    def _docx_relationship_source_part_for_report(
        cls,
        rel_part_name: str,
    ) -> Optional[str]:
        source_part = cls._docx_relationship_source_part(rel_part_name)
        if source_part:
            return source_part
        if rel_part_name == "_rels/.rels":
            return "/"
        return source_part

    @classmethod
    def _external_relationship_risk(cls, rel_type: str) -> Dict[str, Any]:
        rel_type_text = str(rel_type or "")
        rel_type_lower = rel_type_text.lower()
        local = rel_type_lower.rsplit("/", 1)[-1]
        if rel_type_text == cls.EXTERNAL_REL_ALT_CHUNK or local == "altchunk":
            return {
                "risk_level": "fail",
                "risk_category": "altChunk",
                "message": "External altChunk relationships can pull remote content into the document.",
            }
        if (
            rel_type_text == cls.EXTERNAL_REL_ATTACHED_TEMPLATE
            or local in {"attachedtemplate", "template"}
        ):
            return {
                "risk_level": "fail",
                "risk_category": "template",
                "message": "External template relationships can attach remote template content or macros.",
            }
        if (
            rel_type_text in {cls.EXTERNAL_REL_OLE_OBJECT, cls.EXTERNAL_REL_PACKAGE}
            or local in {"oleobject", "package"}
            or "ole" in local
        ):
            return {
                "risk_level": "fail",
                "risk_category": "ole",
                "message": "External OLE/package relationships can reference active or embedded remote content.",
            }
        if rel_type_text == cls.EXTERNAL_REL_IMAGE or local == "image":
            return {
                "risk_level": "fail",
                "risk_category": "external_image",
                "message": "External image relationships can load remote tracking or document content.",
            }
        if rel_type_text == cls.EXTERNAL_REL_HYPERLINK or local == "hyperlink":
            return {
                "risk_level": "warn",
                "risk_category": "hyperlink",
                "message": "External hyperlinks are visible links and require review.",
            }
        return {
            "risk_level": "warn",
            "risk_category": "external_relationship",
            "message": "External relationship requires review.",
        }

    @classmethod
    def _docx_external_relationship_report(
        cls,
        rel_blobs: Dict[str, bytes],
    ) -> Dict[str, Any]:
        report: Dict[str, Any] = {
            "relationship_parts": sorted(rel_blobs),
            "relationship_part_count": len(rel_blobs),
            "external_relationships": [],
            "external_relationship_count": 0,
            "risky_external_relationships": [],
            "risky_external_relationship_count": 0,
            "warning_external_relationships": [],
            "warning_external_relationship_count": 0,
            "external_hyperlink_relationships": [],
            "external_hyperlink_relationship_count": 0,
            "relationship_parse_errors": [],
            "relationship_parse_error_count": 0,
        }
        for rel_part_name, blob in sorted(rel_blobs.items()):
            try:
                root = ET.fromstring(blob)
            except ET.ParseError as exc:
                report["relationship_parse_errors"].append(
                    {"part": rel_part_name, "reason": str(exc)}
                )
                continue
            for child in list(root):
                if cls._xml_local_name(child.tag) != "Relationship":
                    continue
                target_mode = str(child.get("TargetMode", "") or "")
                if target_mode.lower() != "external":
                    continue
                rel_type = str(child.get("Type", "") or "")
                risk = cls._external_relationship_risk(rel_type)
                entry = {
                    "relationship_part": rel_part_name,
                    "source_part": cls._docx_relationship_source_part_for_report(
                        rel_part_name
                    ),
                    "id": child.get("Id", ""),
                    "relationship_id": child.get("Id", ""),
                    "type": rel_type,
                    "target": child.get("Target", ""),
                    "target_mode": target_mode,
                    "risk_level": risk["risk_level"],
                    "risk_category": risk["risk_category"],
                    "risk_message": risk["message"],
                    "risky": risk["risk_level"] == "fail",
                    "warning": risk["risk_level"] == "warn",
                }
                report["external_relationships"].append(entry)
                if entry["risky"]:
                    report["risky_external_relationships"].append(entry)
                if entry["warning"]:
                    report["warning_external_relationships"].append(entry)
                if entry["risk_category"] == "hyperlink":
                    report["external_hyperlink_relationships"].append(entry)

        report["external_relationships"].sort(
            key=lambda item: (
                item.get("relationship_part") or "",
                item.get("id") or "",
                item.get("target") or "",
            )
        )
        report["risky_external_relationships"].sort(
            key=lambda item: (
                item.get("relationship_part") or "",
                item.get("id") or "",
                item.get("target") or "",
            )
        )
        report["warning_external_relationships"].sort(
            key=lambda item: (
                item.get("relationship_part") or "",
                item.get("id") or "",
                item.get("target") or "",
            )
        )
        report["external_hyperlink_relationships"].sort(
            key=lambda item: (
                item.get("relationship_part") or "",
                item.get("id") or "",
                item.get("target") or "",
            )
        )
        report["external_relationship_count"] = len(report["external_relationships"])
        report["risky_external_relationship_count"] = len(
            report["risky_external_relationships"]
        )
        report["warning_external_relationship_count"] = len(
            report["warning_external_relationships"]
        )
        report["external_hyperlink_relationship_count"] = len(
            report["external_hyperlink_relationships"]
        )
        report["relationship_parse_error_count"] = len(
            report["relationship_parse_errors"]
        )
        return report

    @classmethod
    def _resolve_docx_relationship_target(
        cls, rel_part_name: str, target: str, target_mode: str = ""
    ) -> Optional[str]:
        target = str(target or "").strip()
        if not target:
            return None
        if str(target_mode or "").lower() == "external":
            return None
        if re.match(r"^[A-Za-z][A-Za-z0-9+.-]*:", target):
            return None
        if target.startswith("/"):
            resolved = posixpath.normpath(target.lstrip("/"))
        else:
            source_part = cls._docx_relationship_source_part(rel_part_name)
            if not source_part:
                return None
            resolved = posixpath.normpath(
                posixpath.join(posixpath.dirname(source_part), target)
            )
        if resolved in {"", "."} or resolved.startswith("../") or "/../" in resolved:
            return None
        return resolved

    @classmethod
    def _docx_relationship_internal_targets(
        cls, rel_part_name: str, blob: bytes
    ) -> List[str]:
        rel_tag = f"{{{OOXML_NS['rel']}}}Relationship"
        try:
            root = ET.fromstring(blob)
        except ET.ParseError:
            return []
        targets = []
        for child in list(root):
            if child.tag != rel_tag:
                continue
            resolved = cls._resolve_docx_relationship_target(
                rel_part_name,
                child.get("Target", ""),
                child.get("TargetMode", ""),
            )
            if resolved:
                targets.append(resolved)
        return sorted(set(targets))

    @classmethod
    def _docx_all_relationship_internal_targets(
        cls, files: Dict[str, bytes]
    ) -> List[str]:
        targets = []
        for name, blob in files.items():
            if name.endswith(".rels"):
                targets.extend(cls._docx_relationship_internal_targets(name, blob))
        return sorted(set(targets))

    @staticmethod
    def _is_comment_resource_part(part_name: str) -> bool:
        lowered = str(part_name or "").lower()
        return lowered.startswith(
            (
                "word/media/",
                "word/embeddings/",
                "word/activeX/",
                "word/charts/",
                "word/diagrams/",
            )
        )

    @staticmethod
    def _xml_local_name(tag: Any) -> str:
        return str(tag or "").rsplit("}", 1)[-1]

    @classmethod
    def _tracked_revision_counts(cls, root: ET.Element) -> Dict[str, int]:
        counts: Dict[str, int] = {}
        for elem in root.iter():
            local = cls._xml_local_name(elem.tag)
            if local in cls.TRACKED_REVISION_NODES:
                counts[local] = counts.get(local, 0) + 1
        return counts

    @classmethod
    def _drop_revision_property_changes(
        cls,
        elem: ET.Element,
        stats: Dict[str, int],
    ) -> None:
        kept = []
        for child in list(elem):
            cls._drop_revision_property_changes(child, stats)
            local = cls._xml_local_name(child.tag)
            if local in cls.REVISION_PROPERTY_CHANGE_NODES:
                stats[f"{local}_removed"] = stats.get(f"{local}_removed", 0) + 1
                continue
            kept.append(child)
        elem[:] = kept

    @classmethod
    def _normalize_image_format(cls, fmt: str) -> Tuple[Optional[Tuple[str, str]], Optional[Dict[str, Any]]]:
        fmt_key = str(fmt or "").strip().lower().lstrip(".")
        normalized = cls.IMAGE_OUTPUT_FORMATS.get(fmt_key)
        if normalized:
            return normalized, None
        return None, {
            "success": False,
            "error": "Unsupported image format. Use png or jpg.",
            "error_code": "unsupported_image_format",
            "allowed_formats": ["png", "jpg"],
        }

    @classmethod
    def _validate_image_output_prefix(
        cls, prefix: str
    ) -> Tuple[Optional[str], Dict[str, Any], Optional[Dict[str, Any]]]:
        prefix_text = str(prefix or "").strip()
        preflight = {
            "operation": "doc_extract_images",
            "prefix": prefix,
            "max_prefix_length": cls.MAX_OUTPUT_PREFIX_LENGTH,
        }
        if not prefix_text:
            preflight["status"] = "fail"
            preflight["reason"] = "prefix must not be empty"
        elif len(prefix_text) > cls.MAX_OUTPUT_PREFIX_LENGTH:
            preflight["status"] = "fail"
            preflight["reason"] = (
                f"prefix is {len(prefix_text)} characters; max is {cls.MAX_OUTPUT_PREFIX_LENGTH}"
            )
        elif (
            Path(prefix_text).is_absolute()
            or "/" in prefix_text
            or "\\" in prefix_text
            or ":" in prefix_text
            or prefix_text in {".", ".."}
            or any(part == ".." for part in Path(prefix_text).parts)
        ):
            preflight["status"] = "fail"
            preflight["reason"] = "prefix must be a filename prefix, not a path"
        elif not cls.SAFE_PREFIX_RE.match(prefix_text):
            preflight["status"] = "fail"
            preflight["reason"] = (
                "prefix may contain only letters, numbers, spaces, dots, hyphens, and underscores"
            )
        else:
            preflight["status"] = "pass"
            return prefix_text, preflight, None

        return None, preflight, {
            "success": False,
            "error": f"Unsafe image prefix: {preflight['reason']}.",
            "error_code": "unsafe_output_prefix",
            "preflight": preflight,
        }

    @staticmethod
    def _prepare_output_dir(output_dir: str) -> Path:
        output_root = Path(output_dir).expanduser()
        output_root.mkdir(parents=True, exist_ok=True)
        return output_root.resolve()

    @staticmethod
    def _safe_child_path(output_root: Path, filename: str) -> Path:
        candidate = (output_root / filename).resolve()
        try:
            common = os.path.commonpath([str(output_root), str(candidate)])
        except ValueError:
            common = ""
        if common != str(output_root):
            raise ValueError("Refusing to write outside output_dir")
        return candidate

    @classmethod
    def _image_extract_resource_limits(cls) -> Dict[str, int]:
        return {
            "max_images": cls.MAX_EXTRACT_IMAGES,
            "max_compressed_image_bytes": cls.MAX_EXTRACT_IMAGE_COMPRESSED_BYTES,
            "max_decoded_image_pixels": cls.MAX_EXTRACT_IMAGE_PIXELS,
        }

    @classmethod
    def _load_bounded_image(cls, PILImage: Any, image_bytes: bytes) -> Tuple[Any, Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
        compressed_size = len(image_bytes)
        if compressed_size > cls.MAX_EXTRACT_IMAGE_COMPRESSED_BYTES:
            return None, None, {
                "skipped": True,
                "error": (
                    f"Embedded image is {compressed_size} compressed bytes, exceeding "
                    f"the safe limit {cls.MAX_EXTRACT_IMAGE_COMPRESSED_BYTES}"
                ),
                "error_code": "image_compressed_bytes_limit_exceeded",
                "compressed_size_bytes": compressed_size,
                "max_compressed_size_bytes": cls.MAX_EXTRACT_IMAGE_COMPRESSED_BYTES,
            }

        img = None
        try:
            stream = io.BytesIO(image_bytes)
            bomb_warning = getattr(PILImage, "DecompressionBombWarning", None)
            with warnings.catch_warnings():
                if bomb_warning is not None:
                    warnings.simplefilter("error", bomb_warning)
                img = PILImage.open(stream)
                width, height = [int(value) for value in img.size]
                pixels = width * height
                if pixels > cls.MAX_EXTRACT_IMAGE_PIXELS:
                    img.close()
                    return None, None, {
                        "skipped": True,
                        "error": (
                            f"Embedded image decodes to {pixels} pixels, exceeding "
                            f"the safe limit {cls.MAX_EXTRACT_IMAGE_PIXELS}"
                        ),
                        "error_code": "image_pixel_limit_exceeded",
                        "width": width,
                        "height": height,
                        "pixels": pixels,
                        "max_decoded_image_pixels": cls.MAX_EXTRACT_IMAGE_PIXELS,
                        "compressed_size_bytes": compressed_size,
                    }
                img.load()
            return img, {
                "width": width,
                "height": height,
                "pixels": pixels,
                "compressed_size_bytes": compressed_size,
            }, None
        except Exception as exc:
            if img is not None:
                img.close()
            return None, None, {
                "skipped": True,
                "error": f"Could not safely decode embedded image: {exc}",
                "error_code": "unsafe_image_decode",
                "compressed_size_bytes": compressed_size,
            }

    @staticmethod
    def _atomic_write_with_mkstemp(out_path: Path, writer: Any) -> None:
        tmp_name = None
        fd = None
        try:
            fd, tmp_name = tempfile.mkstemp(
                prefix=f".{out_path.name}.",
                suffix=".tmp",
                dir=str(out_path.parent),
            )
            with os.fdopen(fd, "wb") as handle:
                fd = None
                writer(handle)
                handle.flush()
                os.fsync(handle.fileno())
            os.replace(tmp_name, str(out_path))
            tmp_name = None
        finally:
            if fd is not None:
                try:
                    os.close(fd)
                except OSError:
                    pass
            if tmp_name is not None:
                try:
                    os.unlink(tmp_name)
                except FileNotFoundError:
                    pass

    @classmethod
    def _atomic_write_bytes(cls, out_path: Path, data: bytes) -> None:
        cls._atomic_write_with_mkstemp(out_path, lambda handle: handle.write(data))

    @classmethod
    def _save_bounded_image(
        cls,
        img: Any,
        out_path: Path,
        fmt_ext: str,
        pillow_format: str,
    ) -> None:
        def write_image(handle: Any) -> None:
            if fmt_ext == "jpg":
                converted = img.convert("RGB")
                try:
                    converted.save(handle, "JPEG", quality=90)
                finally:
                    converted.close()
            else:
                img.save(handle, pillow_format)

        cls._atomic_write_with_mkstemp(out_path, write_image)

    @staticmethod
    def _docx_related_image_parts(doc: Any) -> List[Dict[str, Any]]:
        image_entries = []
        seen_partnames = set()
        for source_part in getattr(doc.part.package, "parts", []):
            source_partname = str(getattr(source_part, "partname", "")).lstrip("/")
            for rel in getattr(source_part, "rels", {}).values():
                if "image" not in str(getattr(rel, "reltype", "")):
                    continue
                image_part = getattr(rel, "target_part", None)
                if image_part is None:
                    continue
                partname = str(getattr(image_part, "partname", "")).lstrip("/")
                if partname in seen_partnames:
                    continue
                seen_partnames.add(partname)
                image_entries.append(
                    {
                        "source_part": source_partname,
                        "partname": partname,
                        "part": image_part,
                    }
                )
        return image_entries

    @staticmethod
    def _enum_raw(value: Any) -> Optional[int]:
        if value is None:
            return None
        try:
            return int(value)
        except (TypeError, ValueError):
            return None

    @staticmethod
    def _enum_name(value: Any) -> Optional[str]:
        if value is None:
            return None
        text = str(value)
        return text.split(".", 1)[-1].lower() if "." in text else text.lower()

    @staticmethod
    def _readable_alignment(raw_value: Any, ooxml_value: Optional[str] = None) -> Optional[str]:
        if ooxml_value:
            return {
                "both": "justify",
                "center": "center",
                "distribute": "distributed",
                "end": "right",
                "left": "left",
                "right": "right",
                "start": "left",
            }.get(str(ooxml_value), str(ooxml_value))
        raw = DocumentOperations._enum_raw(raw_value)
        return {
            0: "left",
            1: "center",
            2: "right",
            3: "justify",
            4: "distributed",
            5: "justify_medium",
            7: "justify_high",
            8: "justify_low",
            9: "thai_distribute",
        }.get(raw)

    def _style_chain(self, para: Any, doc: Any) -> List[Any]:
        chain = []
        seen = set()
        style = getattr(para, "style", None)
        while style is not None and id(style) not in seen:
            seen.add(id(style))
            chain.append(style)
            style = getattr(style, "base_style", None)
        try:
            normal = doc.styles["Normal"]
        except Exception:
            normal = None
        if normal is not None and id(normal) not in seen:
            chain.append(normal)
        return chain

    def _resolve_paragraph_format_value(
        self, para: Any, doc: Any, attr: str
    ) -> Tuple[Any, Optional[str]]:
        direct_value = getattr(para.paragraph_format, attr)
        if direct_value is not None:
            return direct_value, "direct"
        for style in self._style_chain(para, doc):
            value = getattr(style.paragraph_format, attr)
            if value is not None:
                return value, getattr(style, "name", "style")
        return None, None

    def _resolve_alignment_value(self, para: Any, doc: Any) -> Tuple[Any, str]:
        if para.alignment is not None:
            return para.alignment, "direct"
        for style in self._style_chain(para, doc):
            value = getattr(style.paragraph_format, "alignment", None)
            if value is not None:
                return value, getattr(style, "name", "style")
        return WD_ALIGN_PARAGRAPH.LEFT if DOCX_AVAILABLE else 0, "default"

    def _tab_stops_from_ppr(self, ppr: Any) -> List[Dict[str, Any]]:
        if ppr is None:
            return []
        tabs = ppr.find(qn("w:tabs"))
        if tabs is None:
            return []
        stops = []
        for tab in tabs.findall(qn("w:tab")):
            attrs = self._local_attrs(tab)
            entry: Dict[str, Any] = {"raw_attrs": attrs}
            if "pos" in attrs:
                entry["position"] = self._twips_payload(attrs.get("pos"))
            if "val" in attrs:
                entry["alignment"] = attrs.get("val")
            if "leader" in attrs:
                entry["leader"] = attrs.get("leader")
            stops.append(entry)
        return stops

    def _style_tab_stops(self, para: Any, doc: Any) -> List[Dict[str, Any]]:
        for style in self._style_chain(para, doc):
            style_el = getattr(style, "_element", None)
            ppr = getattr(style_el, "pPr", None)
            if ppr is None and style_el is not None:
                ppr = style_el.find(qn("w:pPr"))
            stops = self._tab_stops_from_ppr(ppr)
            if stops:
                return stops
        return []

    def _inline_page_breaks(self, para: Any) -> List[Dict[str, Any]]:
        breaks = []
        for run_index, run in enumerate(para.runs):
            for br in run._element.findall(qn("w:br")):
                attrs = self._local_attrs(br)
                breaks.append(
                    {
                        "run_index": run_index,
                        "type": attrs.get("type", "textWrapping"),
                        "raw_attrs": attrs,
                    }
                )
        return breaks

    def _inline_page_break_before_text(self, para: Any) -> bool:
        """Return True when a page break appears before visible paragraph text."""
        seen_text = False
        for run in para.runs:
            for child in list(run._element):
                if child.tag == qn("w:t"):
                    if (child.text or "").strip():
                        seen_text = True
                elif child.tag == qn("w:br"):
                    attrs = self._local_attrs(child)
                    if attrs.get("type", "textWrapping") == "page" and not seen_text:
                        return True
        return False

    def _paragraph_section_indices(self, doc: Any) -> List[int]:
        indices = []
        section_index = 0
        for para in doc.paragraphs:
            indices.append(section_index)
            ppr = para._element.pPr
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                section_index += 1
        return indices

    def _paragraph_indent_info(self, para: Any, doc: Any) -> Dict[str, Any]:
        ppr = para._element.pPr
        ind = ppr.find(qn("w:ind")) if ppr is not None else None
        raw_attrs = self._local_attrs(ind)
        raw_payload = {
            key: self._twips_payload(value)
            for key, value in raw_attrs.items()
            if key in {"left", "right", "firstLine", "hanging", "start", "end"}
        }

        direct_left = para.paragraph_format.left_indent
        direct_right = para.paragraph_format.right_indent
        direct_first = para.paragraph_format.first_line_indent

        resolved_left, left_source = self._resolve_paragraph_format_value(
            para, doc, "left_indent"
        )
        resolved_right, right_source = self._resolve_paragraph_format_value(
            para, doc, "right_indent"
        )
        resolved_first, first_source = self._resolve_paragraph_format_value(
            para, doc, "first_line_indent"
        )

        direct_first_pt = self._length_points(direct_first)
        resolved_first_pt = self._length_points(resolved_first)
        resolved_left_payload = self._length_payload(resolved_left)
        resolved_right_payload = self._length_payload(resolved_right)
        effective_left_payload = resolved_left_payload
        effective_left_source = left_source if resolved_left_payload is not None else None
        if effective_left_payload is None:
            for key in ("start", "left"):
                if raw_payload.get(key) is not None:
                    effective_left_payload = raw_payload[key]
                    effective_left_source = f"direct_ooxml_{key}"
                    break
        effective_right_payload = resolved_right_payload
        effective_right_source = right_source if resolved_right_payload is not None else None
        if effective_right_payload is None:
            for key in ("end", "right"):
                if raw_payload.get(key) is not None:
                    effective_right_payload = raw_payload[key]
                    effective_right_source = f"direct_ooxml_{key}"
                    break

        return {
            "raw_ooxml_ind_attrs": raw_attrs,
            "raw_ooxml_ind": raw_payload,
            "direct": {
                "left": self._length_payload(direct_left),
                "right": self._length_payload(direct_right),
                "first_line": self._length_payload(direct_first),
                "hanging": self._points_payload(abs(direct_first_pt))
                if direct_first_pt is not None and direct_first_pt < 0
                else None,
            },
            "style_resolved": {
                "left": resolved_left_payload,
                "left_source": left_source,
                "right": resolved_right_payload,
                "right_source": right_source,
                "effective_left": effective_left_payload,
                "effective_left_source": effective_left_source,
                "effective_right": effective_right_payload,
                "effective_right_source": effective_right_source,
                "first_line": self._length_payload(resolved_first),
                "first_line_source": first_source,
                "hanging": self._points_payload(abs(resolved_first_pt))
                if resolved_first_pt is not None and resolved_first_pt < 0
                else self._twips_payload(raw_attrs.get("hanging")),
                "hanging_source": first_source
                if resolved_first_pt is not None and resolved_first_pt < 0
                else ("direct_ooxml" if raw_attrs.get("hanging") else None),
            },
        }

    @staticmethod
    def _ooxml_on_off_enabled(attrs: Dict[str, Any]) -> bool:
        if not attrs:
            return False
        val = attrs.get("val")
        if val is None:
            return True
        return str(val).strip().lower() not in {"0", "false", "off", "no"}

    def _paragraph_line_spacing_info(self, para: Any, doc: Any) -> Dict[str, Any]:
        ppr = para._element.pPr
        spacing = ppr.find(qn("w:spacing")) if ppr is not None else None
        raw_attrs = self._local_attrs(spacing)
        resolved_spacing, spacing_source = self._resolve_paragraph_format_value(
            para, doc, "line_spacing"
        )
        resolved_rule, rule_source = self._resolve_paragraph_format_value(
            para, doc, "line_spacing_rule"
        )
        return {
            "raw_ooxml_spacing_attrs": raw_attrs,
            "direct": {
                "line_spacing": self._length_payload(
                    para.paragraph_format.line_spacing
                )
                if hasattr(para.paragraph_format.line_spacing, "pt")
                else para.paragraph_format.line_spacing,
                "line_spacing_rule_raw": self._enum_raw(
                    para.paragraph_format.line_spacing_rule
                ),
                "line_spacing_rule": self._enum_name(
                    para.paragraph_format.line_spacing_rule
                ),
            },
            "style_resolved": {
                "line_spacing": self._length_payload(resolved_spacing)
                if hasattr(resolved_spacing, "pt")
                else resolved_spacing,
                "line_spacing_source": spacing_source,
                "line_spacing_rule_raw": self._enum_raw(resolved_rule),
                "line_spacing_rule": self._enum_name(resolved_rule),
                "line_spacing_rule_source": rule_source,
            },
        }

    def _paragraph_page_break_before_info(self, para: Any, doc: Any) -> Dict[str, Any]:
        ppr = para._element.pPr
        raw = ppr.find(qn("w:pageBreakBefore")) if ppr is not None else None
        resolved, source = self._resolve_paragraph_format_value(
            para, doc, "page_break_before"
        )
        return {
            "direct": para.paragraph_format.page_break_before,
            "style_resolved": resolved,
            "style_resolved_source": source,
            "raw_ooxml": self._local_attrs(raw),
            "present_in_ooxml": raw is not None,
            "enabled_in_ooxml": self._ooxml_on_off_enabled(self._local_attrs(raw)),
        }

    def _paragraph_format_payload(
        self, para: Any, doc: Any, index: int, section_index: int
    ) -> Dict[str, Any]:
        ppr = para._element.pPr
        jc = ppr.find(qn("w:jc")) if ppr is not None else None
        jc_attrs = self._local_attrs(jc)
        resolved_alignment, resolved_alignment_source = self._resolve_alignment_value(
            para, doc
        )
        direct_tabs = self._tab_stops_from_ppr(ppr)
        style_tabs = self._style_tab_stops(para, doc)
        inline_breaks = self._inline_page_breaks(para)
        payload = {
            "index": index,
            "section_index": section_index,
            "style_name": para.style.name if para.style is not None else None,
            "alignment": {
                "raw": self._enum_raw(para.alignment),
                "readable": self._readable_alignment(
                    para.alignment if para.alignment is not None else resolved_alignment,
                    jc_attrs.get("val"),
                ),
                "raw_ooxml": jc_attrs.get("val"),
                "style_resolved_raw": self._enum_raw(resolved_alignment),
                "style_resolved_readable": self._readable_alignment(resolved_alignment),
                "style_resolved_source": resolved_alignment_source,
            },
            "text_preview": para.text[:120],
            "text_length": len(para.text or ""),
            "indents": self._paragraph_indent_info(para, doc),
            "tab_stops": {
                "direct": direct_tabs,
                "style_resolved": direct_tabs or style_tabs,
                "style_source": None if direct_tabs or not style_tabs else "style",
            },
            "line_spacing": self._paragraph_line_spacing_info(para, doc),
            "page_break_before": self._paragraph_page_break_before_info(para, doc),
            "inline_page_breaks": inline_breaks,
            "inline_page_break_count": len(
                [br for br in inline_breaks if br.get("type") == "page"]
            ),
        }
        if para.runs:
            run = para.runs[0]
            payload["font"] = {
                "name": run.font.name,
                "size": run.font.size.pt if run.font.size else None,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            }
        return payload

    @staticmethod
    def _xml_declared_encoding(blob: bytes) -> Optional[str]:
        match = re.match(
            rb"\s*<\?xml\b[^>]*\bencoding\s*=\s*([\"'])([^\"']+)\1",
            blob[:256],
            flags=re.IGNORECASE,
        )
        if not match:
            return None
        return match.group(2).decode("ascii", errors="replace")

    @staticmethod
    def _decode_ooxml_xml(blob: bytes) -> Tuple[Optional[str], Optional[str]]:
        encoding = DocumentOperations._xml_declared_encoding(blob)
        if encoding and encoding.lower().replace("_", "-") not in {"utf-8", "utf8"}:
            return None, f"unsupported XML encoding: {encoding}"
        try:
            return blob.decode("utf-8"), None
        except UnicodeDecodeError as exc:
            return None, f"UTF-8 decode failed: {exc}"

    @staticmethod
    def _namespace_prefixes_for_uri(text: str, uri: str) -> List[Optional[str]]:
        prefixes: List[Optional[str]] = []
        pattern = re.compile(
            r"\bxmlns(?::([A-Za-z_][\w.\-]*))?\s*=\s*([\"'])(.*?)\2",
            flags=re.DOTALL,
        )
        for match in pattern.finditer(text):
            if match.group(3) == uri:
                prefixes.append(match.group(1))
        return prefixes

    @staticmethod
    def _root_tag_prefix(text: str) -> Tuple[Optional[str], Optional[str]]:
        for match in re.finditer(
            r"<(?![?!/])([A-Za-z_][\w.\-]*)(?::([A-Za-z_][\w.\-]*))?\b",
            text,
        ):
            if match.group(2):
                return match.group(1), match.group(2)
            return None, match.group(1)
        return None, None

    @staticmethod
    def _attribute_value_mentions_prefix(text: str, prefix: str) -> bool:
        token = re.escape(prefix)
        value_pattern = re.compile(r"=\s*([\"'])(.*?)\1", flags=re.DOTALL)
        for match in value_pattern.finditer(text):
            value = match.group(2)
            if re.search(rf"(?<![\w.\-]){token}(?=[:\s,;]|$)", value):
                return True
        return False

    @staticmethod
    def _replace_xml_name_prefixes_in_tag(tag: str, prefixes: List[str]) -> str:
        if not prefixes:
            return tag
        ordered = sorted(prefixes, key=len, reverse=True)
        name_chars = set(
            "abcdefghijklmnopqrstuvwxyz"
            "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
            "0123456789_.-"
        )
        out: List[str] = []
        quote: Optional[str] = None
        i = 0
        while i < len(tag):
            char = tag[i]
            if quote:
                out.append(char)
                if char == quote:
                    quote = None
                i += 1
                continue
            if char in {"'", '"'}:
                quote = char
                out.append(char)
                i += 1
                continue
            replaced = False
            for prefix in ordered:
                probe = f"{prefix}:"
                if not tag.startswith(probe, i):
                    continue
                previous = tag[i - 1] if i > 0 else ""
                if previous == ":" or previous in name_chars:
                    continue
                out.append("w:")
                i += len(probe)
                replaced = True
                break
            if replaced:
                continue
            out.append(char)
            i += 1
        return "".join(out)

    def _inspect_wordprocessingml_part_prefixes(
        self,
        part_name: str,
        blob: bytes,
    ) -> Dict[str, Any]:
        text, decode_error = self._decode_ooxml_xml(blob)
        info: Dict[str, Any] = {
            "part": part_name,
            "word_prefixes": [],
            "default_namespace_is_word": False,
            "root_prefix": None,
            "root_local_name": None,
            "root_namespace": None,
            "root_uses_canonical_word_prefix": None,
            "uses_canonical_word_prefix": False,
            "noncanonical_word_prefixes": [],
            "parse_error": decode_error,
            "layout_risk": True,
        }
        if text is None:
            return info

        prefixes = self._namespace_prefixes_for_uri(text, OOXML_NS["w"])
        word_prefixes = sorted(prefix for prefix in prefixes if prefix)
        default_is_word = any(prefix is None for prefix in prefixes)
        root_prefix, root_local = self._root_tag_prefix(text)
        info.update(
            {
                "word_prefixes": word_prefixes,
                "default_namespace_is_word": default_is_word,
                "root_prefix": root_prefix,
                "root_local_name": root_local,
            }
        )

        try:
            root = ET.fromstring(blob)
        except ET.ParseError as exc:
            info["parse_error"] = str(exc)
            return info

        root_namespace = root.tag[1:].split("}", 1)[0] if root.tag.startswith("{") else None
        root_is_word = root_namespace == OOXML_NS["w"]
        noncanonical = [
            prefix for prefix in word_prefixes if prefix != "w"
        ]
        if default_is_word:
            noncanonical.append("<default>")
        root_uses_canonical = root_prefix == "w" if root_is_word else True
        uses_canonical = (
            "w" in word_prefixes
            and not noncanonical
            and root_uses_canonical
        )
        info.update(
            {
                "root_namespace": root_namespace,
                "root_uses_canonical_word_prefix": root_uses_canonical,
                "uses_canonical_word_prefix": uses_canonical,
                "noncanonical_word_prefixes": noncanonical,
                "parse_error": None,
                "layout_risk": not uses_canonical,
            }
        )
        return info

    def _docx_ooxml_prefix_report(self, file_path: str) -> Dict[str, Any]:
        report = {
            "document_xml_word_prefixes": [],
            "document_xml_root_prefix": None,
            "document_xml_uses_canonical_word_prefix": None,
            "uses_canonical_word_prefix": None,
            "wordprocessingml_parts": [],
            "wordprocessingml_part_count": 0,
            "noncanonical_parts": [],
            "noncanonical_part_count": 0,
            "parse_error_parts": [],
            "parse_error_part_count": 0,
            "layout_risk_parts": [],
            "layout_risk_part_count": 0,
            "has_layout_risk": False,
            "submission_blocking": False,
            "zip_preflight": None,
            "warnings": [],
        }
        preflight = self._docx_zip_preflight(file_path)
        report["zip_preflight"] = preflight
        if not preflight.get("success"):
            report["warnings"].append("DOCX OOXML package preflight failed.")
            report["submission_blocking"] = True
            return report
        try:
            files = {}
            with ZipFile(file_path, "r") as zin:
                for name in zin.namelist():
                    if not name.endswith(".xml"):
                        continue
                    blob = zin.read(name)
                    if OOXML_NS["w"].encode("utf-8") in blob:
                        files[name] = blob
        except Exception as exc:
            report["warnings"].append(f"Could not inspect DOCX OOXML parts: {exc}")
            report["submission_blocking"] = True
            return report

        for part_name, blob in sorted(files.items()):
            part_info = self._inspect_wordprocessingml_part_prefixes(part_name, blob)
            report["wordprocessingml_parts"].append(part_info)
            if part_name == "word/document.xml":
                report["document_xml_word_prefixes"] = part_info["word_prefixes"]
                report["document_xml_root_prefix"] = part_info["root_prefix"]
                report["document_xml_uses_canonical_word_prefix"] = part_info[
                    "uses_canonical_word_prefix"
                ]
            if part_info.get("parse_error"):
                report["parse_error_parts"].append(
                    {
                        "part": part_name,
                        "reason": part_info["parse_error"],
                    }
                )
            if part_info.get("noncanonical_word_prefixes"):
                report["noncanonical_parts"].append(part_info)
            if part_info.get("layout_risk") or part_info.get("parse_error"):
                report["layout_risk_parts"].append(part_info)

        report["wordprocessingml_part_count"] = len(report["wordprocessingml_parts"])
        report["noncanonical_part_count"] = len(report["noncanonical_parts"])
        report["parse_error_part_count"] = len(report["parse_error_parts"])
        report["layout_risk_part_count"] = len(report["layout_risk_parts"])
        report["has_layout_risk"] = bool(report["layout_risk_parts"])
        report["submission_blocking"] = report["has_layout_risk"]
        report["uses_canonical_word_prefix"] = (
            report["document_xml_uses_canonical_word_prefix"] is True
            and not report["submission_blocking"]
        )

        if report["noncanonical_parts"]:
            report["warnings"].append(
                "DOCX contains WordprocessingML parts with non-canonical prefixes. "
                "OnlyOffice x2t 9.3 has been observed to ignore page breaks, "
                "indents, and margins in ns0/ns1-round-tripped DOCX parts."
            )
        if report["parse_error_parts"]:
            report["warnings"].append(
                "Some WordprocessingML XML parts could not be parsed and were not "
                "safe to inspect or repair."
            )
        if report["document_xml_uses_canonical_word_prefix"] is not True:
            report["warnings"].append(
                "word/document.xml does not use the canonical w: prefix for "
                "WordprocessingML."
            )
        return report

    def _safe_repair_wordprocessingml_prefixes(
        self,
        part_name: str,
        blob: bytes,
    ) -> Tuple[Optional[bytes], Dict[str, Any]]:
        text, decode_error = self._decode_ooxml_xml(blob)
        if text is None:
            return None, {"part": part_name, "reason": decode_error or "decode failed"}

        before = self._inspect_wordprocessingml_part_prefixes(part_name, blob)
        if before.get("parse_error"):
            return None, {"part": part_name, "reason": before["parse_error"]}
        noncanonical = [
            prefix
            for prefix in before.get("noncanonical_word_prefixes", [])
            if prefix != "<default>"
        ]
        if before.get("default_namespace_is_word"):
            return None, {
                "part": part_name,
                "reason": "default WordprocessingML namespace repair is not safe",
            }
        if not noncanonical:
            return blob, {"part": part_name, "changed": False, "before": before}

        declared_for_w = self._namespace_prefixes_for_uri(text, OOXML_NS["w"])
        declared_prefixes = [prefix for prefix in declared_for_w if prefix]
        if "w" not in declared_prefixes:
            if re.search(r"\bxmlns:w\s*=", text):
                return None, {
                    "part": part_name,
                    "reason": "prefix 'w' is already declared for another namespace",
                }
            first_prefix = noncanonical[0]
            declaration = re.compile(
                rf"\bxmlns:{re.escape(first_prefix)}\s*=\s*([\"'])"
                rf"{re.escape(OOXML_NS['w'])}\1",
                flags=re.DOTALL,
            )
            text, replacements = declaration.subn(
                lambda match: f"xmlns:w={match.group(1)}{OOXML_NS['w']}{match.group(1)}",
                text,
                count=1,
            )
            if replacements != 1:
                return None, {
                    "part": part_name,
                    "reason": f"could not replace xmlns:{first_prefix} declaration",
                }
        for prefix in noncanonical:
            if self._attribute_value_mentions_prefix(text, prefix):
                return None, {
                    "part": part_name,
                    "reason": (
                        f"prefix '{prefix}' is referenced inside an attribute value; "
                        "safe prefix repair skipped"
                    ),
                }

        def rewrite_tag(match: re.Match) -> str:
            tag = match.group(0)
            if tag.startswith(("<?", "<!", "<!--")):
                return tag
            return self._replace_xml_name_prefixes_in_tag(tag, noncanonical)

        repaired_text = re.sub(r"<[^<>]+>", rewrite_tag, text)
        for prefix in noncanonical:
            if prefix == "w":
                continue
            declaration = re.compile(
                rf"\s+xmlns:{re.escape(prefix)}\s*=\s*([\"'])"
                rf"{re.escape(OOXML_NS['w'])}\1",
                flags=re.DOTALL,
            )
            repaired_text = declaration.sub("", repaired_text)

        repaired = repaired_text.encode("utf-8")
        try:
            ET.fromstring(repaired)
        except ET.ParseError as exc:
            return None, {
                "part": part_name,
                "reason": f"repaired XML did not parse: {exc}",
            }
        after = self._inspect_wordprocessingml_part_prefixes(part_name, repaired)
        if after.get("parse_error") or after.get("noncanonical_word_prefixes"):
            return None, {
                "part": part_name,
                "reason": "safe prefix repair did not produce canonical WordprocessingML prefixes",
                "after": after,
            }
        return repaired, {
            "part": part_name,
            "changed": repaired != blob,
            "before": before,
            "after": after,
        }

    def _canonicalize_wordprocessingml_parts(
        self, files: Dict[str, bytes]
    ) -> Dict[str, Any]:
        """Repair non-canonical WordprocessingML prefixes without ET round-tripping."""
        word_ns = OOXML_NS["w"].encode("utf-8")
        rewritten_parts = []
        unchanged_parts = []
        skipped_parts = []

        for part_name, blob in list(files.items()):
            if not part_name.endswith(".xml") or word_ns not in blob:
                continue
            repaired, info = self._safe_repair_wordprocessingml_prefixes(part_name, blob)
            if repaired is None:
                skipped_parts.append(info)
                continue
            files[part_name] = repaired
            if info.get("changed"):
                rewritten_parts.append(part_name)
            else:
                unchanged_parts.append(part_name)

        return {
            "parts_seen": len(rewritten_parts) + len(unchanged_parts) + len(skipped_parts),
            "parts_rewritten": len(rewritten_parts),
            "parts_unchanged": len(unchanged_parts),
            "parts_skipped": len(skipped_parts),
            "rewritten_parts": rewritten_parts,
            "unchanged_parts": unchanged_parts,
            "skipped_parts": skipped_parts,
            "critical_skipped": [
                item for item in skipped_parts if item.get("part") == "word/document.xml"
            ],
            "warnings": [
                f"Skipped {item.get('part')}: {item.get('reason')}"
                for item in skipped_parts
            ],
        }

    def create_document(
        self, output_path: str, title: str = "", content: str = ""
    ) -> Dict[str, Any]:
        """Create a new .docx document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(output_path):
                backup = self.host._snapshot_backup(output_path)
                doc = Document()
                section = doc.sections[0]
                section.page_width = Mm(210)
                section.page_height = Mm(297)
                section.top_margin = Pt(72)
                section.bottom_margin = Pt(72)
                section.left_margin = Pt(72)
                section.right_margin = Pt(72)
                normal = doc.styles["Normal"]
                normal.font.name = "Calibri"
                normal.font.size = Pt(11)
                normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                normal.paragraph_format.space_after = Pt(0)
                if title:
                    heading = doc.add_heading(title, 0)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if content:
                    for paragraph in content.split("\n"):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph)
                self.host._safe_save(doc, output_path)
            return {
                "success": True,
                "file": output_path,
                "title": title,
                "size": Path(output_path).stat().st_size,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def read_document(self, file_path: str) -> Dict[str, Any]:
        """Read and extract all text from a .docx document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return {
                "success": True,
                "file": file_path,
                "paragraphs": paragraphs,
                "paragraph_count": len(paragraphs),
                "full_text": "\n\n".join(paragraphs),
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def append_to_document(self, file_path: str, content: str) -> Dict[str, Any]:
        """Append content to a .docx document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                for paragraph in content.split("\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "appended_length": len(content),
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def search_replace_document(
        self, file_path: str, search_text: str, replace_text: str
    ) -> Dict[str, Any]:
        """Find and replace text in a .docx document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                replacements = 0
                for para in doc.paragraphs:
                    replacements += self.host._replace_across_runs(
                        para, search_text, replace_text
                    )
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "replacements": replacements,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def format_paragraph(
        self,
        file_path: str,
        paragraph_index: int,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_name: str = None,
        font_size: int = None,
        color: str = None,
        alignment: str = None,
    ) -> Dict[str, Any]:
        """Apply formatting to a specific paragraph."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                if paragraph_index >= len(doc.paragraphs):
                    return {
                        "success": False,
                        "error": f"Paragraph index {paragraph_index} out of range",
                    }
                para = doc.paragraphs[paragraph_index]
                run = para.runs[0] if para.runs else para.add_run("")
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if underline:
                    run.underline = True
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if color:
                    run.font.color.rgb = RGBColor(*self.host._hex_to_rgb(color))
                if alignment:
                    para.alignment = self.host._get_alignment(alignment)
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "paragraph": paragraph_index,
                "formatted": True,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def highlight_text(
        self, file_path: str, search_text: str, color: str = "yellow"
    ) -> Dict[str, Any]:
        """Highlight all occurrences of text in a document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                highlights = 0
                for para in doc.paragraphs:
                    for run in para.runs:
                        if search_text in run.text:
                            rPr = run._element.get_or_add_rPr()
                            highlight = rPr.find(qn("w:highlight"))
                            if highlight is None:
                                highlight = OxmlElement("w:highlight")
                                rPr.append(highlight)
                            valid_colors = {
                                "yellow",
                                "green",
                                "cyan",
                                "magenta",
                                "blue",
                                "red",
                                "darkBlue",
                                "darkCyan",
                                "darkGreen",
                                "darkMagenta",
                                "darkRed",
                                "darkYellow",
                                "darkGray",
                                "lightGray",
                                "black",
                                "white",
                                "none",
                            }
                            color_val = (
                                color if color and color.lower() in valid_colors else "yellow"
                            )
                            highlight.set(qn("w:val"), color_val)
                            highlights += 1
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "highlights": highlights,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_table(self, file_path: str, headers_csv: str, data_csv: str) -> Dict[str, Any]:
        """Add a formatted table to a Word document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            headers = [h.strip() for h in headers_csv.split(",")]
            rows = [
                [c.strip() for c in row.split(",")]
                for row in data_csv.split(";")
                if row.strip()
            ]
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                for ri, row in enumerate(rows, 1):
                    for ci, val in enumerate(row):
                        if ci < len(headers):
                            table.rows[ri].cells[ci].text = val
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "rows": len(rows),
                "columns": len(headers),
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_comment(
        self,
        file_path: str,
        comment_text: str,
        paragraph_index: int = 0,
        author: str = "SLOANE Agent",
    ) -> Dict[str, Any]:
        """Add a real OOXML comment annotation to a specific paragraph."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            from lxml import etree as LET
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI

            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                if paragraph_index >= len(doc.paragraphs):
                    return {
                        "success": False,
                        "error": f"Paragraph index {paragraph_index} out of range (max {len(doc.paragraphs) - 1})",
                    }

                w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                comments_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
                comments_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

                part = doc.part
                comments_el = None
                comments_part = None
                for rel in part.rels.values():
                    if "comments" in rel.reltype:
                        comments_part = rel.target_part
                        break

                if comments_part is not None:
                    comments_el = LET.fromstring(comments_part.blob)
                else:
                    comments_el = LET.Element(
                        f"{{{w_ns}}}comments", nsmap={"w": w_ns, "r": r_ns}
                    )
                    blob = LET.tostring(
                        comments_el,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )
                    comments_part = Part(
                        PackURI("/word/comments.xml"), comments_ct, blob, part.package
                    )
                    part.relate_to(comments_part, comments_rel)

                existing_ids = []
                for c in comments_el.iter(f"{{{w_ns}}}comment"):
                    cid = c.get(f"{{{w_ns}}}id")
                    if cid is not None:
                        try:
                            existing_ids.append(int(cid))
                        except ValueError:
                            pass
                comment_id = str(max(existing_ids, default=-1) + 1)

                now_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
                comment_elem = LET.SubElement(comments_el, f"{{{w_ns}}}comment")
                comment_elem.set(f"{{{w_ns}}}id", comment_id)
                comment_elem.set(f"{{{w_ns}}}author", author)
                comment_elem.set(f"{{{w_ns}}}date", now_str)
                cp = LET.SubElement(comment_elem, f"{{{w_ns}}}p")
                cr = LET.SubElement(cp, f"{{{w_ns}}}r")
                ct = LET.SubElement(cr, f"{{{w_ns}}}t")
                ct.text = comment_text
                comments_part._blob = LET.tostring(
                    comments_el,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )

                para_el = doc.paragraphs[paragraph_index]._element
                range_start = OxmlElement("w:commentRangeStart")
                range_start.set(qn("w:id"), comment_id)
                range_end = OxmlElement("w:commentRangeEnd")
                range_end.set(qn("w:id"), comment_id)
                ref_run = OxmlElement("w:r")
                ref_rpr = OxmlElement("w:rPr")
                ref_style = OxmlElement("w:rStyle")
                ref_style.set(qn("w:val"), "CommentReference")
                ref_rpr.append(ref_style)
                ref_run.append(ref_rpr)
                comment_ref = OxmlElement("w:commentReference")
                comment_ref.set(qn("w:id"), comment_id)
                ref_run.append(comment_ref)
                para_el.insert(0, range_start)
                para_el.append(range_end)
                para_el.append(ref_run)

                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "comment_id": int(comment_id),
                "comment_text": comment_text,
                "paragraph_index": paragraph_index,
                "author": author,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_reference(self, file_path: str, ref_json: str) -> Dict[str, Any]:
        """Add a reference to the sidecar .refs.json file for later formatting."""
        try:
            ref = json.loads(ref_json)
            required = ["author", "year", "title"]
            missing = [k for k in required if not ref.get(k)]
            if missing:
                return {"success": False, "error": f"Missing required fields: {missing}"}

            refs_path = file_path + ".refs.json"
            refs_target = Path(refs_path)
            refs_target.parent.mkdir(parents=True, exist_ok=True)
            with self.host._file_lock(str(refs_target)):
                refs = []
                if refs_target.exists():
                    with open(refs_target, "r", encoding="utf-8") as f:
                        refs = json.load(f)
                    if not isinstance(refs, list):
                        return {
                            "success": False,
                            "error": f"References sidecar must contain a JSON list: {refs_path}",
                        }

                sig = (
                    ref["author"].strip().lower(),
                    str(ref["year"]).strip(),
                    ref["title"].strip().lower(),
                )
                for existing in refs:
                    esig = (
                        str(existing.get("author", "")).strip().lower(),
                        str(existing.get("year", "")).strip(),
                        str(existing.get("title", "")).strip().lower(),
                    )
                    if esig == sig:
                        return {
                            "success": True,
                            "file": refs_path,
                            "action": "duplicate_skipped",
                            "total_refs": len(refs),
                            "note": f"Reference already exists: {ref['author']} ({ref['year']})",
                        }

                backup = (
                    self.host._snapshot_backup(str(refs_target))
                    if refs_target.exists()
                    else ""
                )
                ref.setdefault("type", "journal")
                refs.append(ref)
                fd, tmp_path = tempfile.mkstemp(
                    prefix=f".{refs_target.name}.",
                    suffix=".tmp",
                    dir=str(refs_target.parent),
                    text=True,
                )
                try:
                    with os.fdopen(fd, "w", encoding="utf-8") as f:
                        json.dump(refs, f, indent=2)
                        f.write("\n")
                        f.flush()
                        os.fsync(f.fileno())
                    os.replace(tmp_path, str(refs_target))
                    self.host._fsync_directory(refs_target.parent)
                finally:
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)

            return {
                "success": True,
                "file": refs_path,
                "action": "added",
                "total_refs": len(refs),
                "backup": backup or None,
                "in_text_citation": self.host._apa_in_text(ref["author"], ref["year"]),
            }
        except json.JSONDecodeError as e:
            return {"success": False, "error": f"Invalid JSON: {e}"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    def build_references(self, file_path: str) -> Dict[str, Any]:
        """Build an APA 7th References section from the sidecar .refs.json file."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        refs_path = file_path + ".refs.json"
        if not os.path.exists(refs_path):
            return {
                "success": False,
                "error": f"No references file found at {refs_path}. Use doc-add-reference first.",
            }
        try:
            with open(refs_path, "r") as f:
                refs = json.load(f)
            if not refs:
                return {"success": False, "error": "References file is empty"}

            refs.sort(key=lambda r: r.get("author", "").split(",")[0].strip().lower())

            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)

                ref_start_idx = None
                for i, para in enumerate(doc.paragraphs):
                    if para.text.strip() == "References" and any(
                        r.bold for r in para.runs if r.bold
                    ):
                        ref_start_idx = i
                        break
                if ref_start_idx is not None:
                    body = doc.element.body
                    paras_to_remove = list(doc.paragraphs[ref_start_idx:])
                    for p in paras_to_remove:
                        body.remove(p._element)
                    if ref_start_idx > 0:
                        prev = (
                            doc.paragraphs[ref_start_idx - 1]
                            if ref_start_idx - 1 < len(doc.paragraphs)
                            else None
                        )
                        if prev and not prev.text.strip():
                            for run_el in prev._element.findall(qn("w:r")):
                                for br in run_el.findall(qn("w:br")):
                                    if br.get(qn("w:type")) == "page":
                                        body.remove(prev._element)
                                        break

                doc.add_page_break()
                heading = doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = heading.add_run("References")
                run.bold = True

                for ref in refs:
                    spans = self.host._format_apa7_reference(ref)
                    para = doc.add_paragraph()
                    for span in spans:
                        run = para.add_run(span["text"])
                        if span.get("italic"):
                            run.italic = True
                    pf = para.paragraph_format
                    pf.first_line_indent = -Inches(0.5)
                    pf.left_indent = Inches(0.5)
                    pf.space_after = Pt(0)
                    pf.space_before = Pt(0)
                    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

                self.host._safe_save(doc, file_path)

            return {
                "success": True,
                "file": file_path,
                "references_added": len(refs),
                "refs_file": refs_path,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_image(
        self,
        file_path: str,
        image_path: str,
        width_inches: float = 5.5,
        caption: str = None,
        paragraph_index: int = None,
        position: str = "after",
    ) -> Dict[str, Any]:
        """Add an image to a Word document with optional caption."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        if not os.path.exists(image_path):
            return {"success": False, "error": f"Image not found: {image_path}"}
        position = position.lower()
        if position not in {"before", "after"}:
            return {"success": False, "error": "position must be 'before' or 'after'"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                total = len(doc.paragraphs)
                if paragraph_index is not None and (
                    paragraph_index < 0 or paragraph_index >= total
                ):
                    return {
                        "success": False,
                        "error": f"Paragraph index {paragraph_index} out of range (0..{total - 1})",
                    }
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                run = para.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                cap_para = None
                if caption:
                    para.paragraph_format.keep_with_next = True
                    cap_para = doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_para.paragraph_format.space_before = Pt(0)
                    cap_para.paragraph_format.space_after = Pt(0)
                    cap_para.paragraph_format.keep_together = True
                    cap_run = cap_para.add_run(caption)
                    cap_run.italic = True
                    cap_run.font.size = Pt(10)
                if paragraph_index is not None:
                    body = doc.element.body
                    ref = doc.paragraphs[paragraph_index]._element
                    elements = [para._element]
                    if cap_para is not None:
                        elements.append(cap_para._element)
                    for element in elements:
                        body.remove(element)
                    insert_at = body.index(ref)
                    if position == "after":
                        insert_at += 1
                    for offset, element in enumerate(elements):
                        body.insert(insert_at + offset, element)
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "image": image_path,
                "width_inches": width_inches,
                "caption": caption,
                "paragraph_index": paragraph_index,
                "position": "append" if paragraph_index is None else position,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def set_paragraph_style(
        self, file_path: str, paragraph_index: int, style_name: str
    ) -> Dict[str, Any]:
        """Set a paragraph's Word style."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                if paragraph_index >= len(doc.paragraphs):
                    return {
                        "success": False,
                        "error": f"Paragraph index {paragraph_index} out of range (max {len(doc.paragraphs) - 1})",
                    }
                para = doc.paragraphs[paragraph_index]
                available_styles = [s.name for s in doc.styles if s.type == 1]
                if style_name not in available_styles:
                    return {
                        "success": False,
                        "error": f"Style '{style_name}' not found. Available: {available_styles[:20]}",
                    }
                para.style = doc.styles[style_name]
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "paragraph_index": paragraph_index,
                "style_applied": style_name,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def set_page_layout(
        self,
        file_path: str,
        orientation: str = None,
        margins: Dict[str, float] = None,
        header_text: str = None,
        page_numbers: bool = False,
        page_size: str = None,
    ) -> Dict[str, Any]:
        """Set page layout (page size, orientation, margins, header, page numbers)."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            from docx.enum.section import WD_ORIENT

            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                resolved_page_size = self.host._resolve_page_size(page_size)
                normalized_orientation = (
                    str(orientation).strip().lower() if orientation is not None else None
                )
                if normalized_orientation is not None and normalized_orientation not in {
                    "portrait",
                    "landscape",
                }:
                    return {
                        "success": False,
                        "error": "orientation must be 'portrait' or 'landscape'",
                    }

                for section in doc.sections:
                    effective_orientation = normalized_orientation or self.host._current_orientation(
                        section
                    )
                    if resolved_page_size:
                        _, width, height = resolved_page_size
                        section.page_width = width
                        section.page_height = height

                    if effective_orientation == "landscape":
                        section.orientation = WD_ORIENT.LANDSCAPE
                        if section.page_width < section.page_height:
                            section.page_width, section.page_height = (
                                section.page_height,
                                section.page_width,
                            )
                    else:
                        section.orientation = WD_ORIENT.PORTRAIT
                        if section.page_width > section.page_height:
                            section.page_width, section.page_height = (
                                section.page_height,
                                section.page_width,
                            )

                    if margins:
                        candidate_margins = {
                            "top": section.top_margin.inches,
                            "bottom": section.bottom_margin.inches,
                            "left": section.left_margin.inches,
                            "right": section.right_margin.inches,
                        }
                        for side, value in margins.items():
                            if side in candidate_margins:
                                try:
                                    numeric = float(value)
                                except (TypeError, ValueError):
                                    return {
                                        "success": False,
                                        "error": f"Invalid {side} margin: expected number.",
                                    }
                                if not math.isfinite(numeric) or numeric < 0:
                                    return {
                                        "success": False,
                                        "error": (
                                            f"Invalid {side} margin: expected a finite "
                                            "non-negative value in inches."
                                        ),
                                    }
                                candidate_margins[side] = numeric
                        if (
                            candidate_margins["left"] + candidate_margins["right"]
                            >= section.page_width.inches
                        ):
                            return {
                                "success": False,
                                "error": (
                                    "Invalid margins: left + right must leave positive "
                                    "usable page width."
                                ),
                            }
                        if (
                            candidate_margins["top"] + candidate_margins["bottom"]
                            >= section.page_height.inches
                        ):
                            return {
                                "success": False,
                                "error": (
                                    "Invalid margins: top + bottom must leave positive "
                                    "usable page height."
                                ),
                            }
                        for side, value in margins.items():
                            if hasattr(section, f"{side}_margin"):
                                setattr(section, f"{side}_margin", Inches(value))

                    if header_text is not None:
                        header = section.header
                        header.is_linked_to_previous = False
                        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                        hp.text = ""
                        run = hp.add_run(header_text)
                        run.font.size = Pt(10)
                        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    if page_numbers:
                        footer = section.footer
                        footer.is_linked_to_previous = False
                        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                        fp.text = ""
                        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = fp.add_run()
                        run.font.size = Pt(10)
                        fld_char_begin = OxmlElement("w:fldChar")
                        fld_char_begin.set(qn("w:fldCharType"), "begin")
                        run._element.append(fld_char_begin)
                        instr_run = fp.add_run()
                        instr_text = OxmlElement("w:instrText")
                        instr_text.set(qn("xml:space"), "preserve")
                        instr_text.text = " PAGE "
                        instr_run._element.append(instr_text)
                        fld_char_end_run = fp.add_run()
                        fld_char_end = OxmlElement("w:fldChar")
                        fld_char_end.set(qn("w:fldCharType"), "end")
                        fld_char_end_run._element.append(fld_char_end)

                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "layout_updated": True,
                "page_size": resolved_page_size[0] if resolved_page_size else None,
                "orientation": normalized_orientation
                or self.host._current_orientation(doc.sections[0]),
                "header_text": header_text,
                "page_numbers": page_numbers,
                "section_count": len(doc.sections),
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def insert_paragraph(
        self, file_path: str, text: str, index: int, style: str = None
    ) -> Dict[str, Any]:
        """Insert a paragraph at a specific index."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            from docx.text.paragraph import Paragraph

            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                total = len(doc.paragraphs)
                if index < 0 or index > total:
                    return {"success": False, "error": f"Index {index} out of range (0..{total})"}
                body = doc.element.body
                new_para = OxmlElement("w:p")
                run = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = text
                t.set(qn("xml:space"), "preserve")
                run.append(t)
                new_para.append(run)
                if index < total:
                    ref = doc.paragraphs[index]._element
                    body.insert(body.index(ref), new_para)
                else:
                    body.append(new_para)
                if style:
                    p = Paragraph(new_para, doc)
                    available = [s.name for s in doc.styles if s.type == 1]
                    if style in available:
                        p.style = doc.styles[style]
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "inserted_at": index,
                "text": text[:100],
                "style": style,
                "total_paragraphs": total + 1,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def delete_paragraph(self, file_path: str, index: int) -> Dict[str, Any]:
        """Delete a paragraph by index."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                total = len(doc.paragraphs)
                if index < 0 or index >= total:
                    return {
                        "success": False,
                        "error": f"Index {index} out of range (0..{total - 1})",
                    }
                para = doc.paragraphs[index]
                deleted_text = para.text[:200]
                doc.element.body.remove(para._element)
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "deleted_index": index,
                "deleted_text": deleted_text,
                "remaining_paragraphs": total - 1,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def search_document(
        self, file_path: str, search_text: str, case_sensitive: bool = False
    ) -> Dict[str, Any]:
        """Search for text in a document and return all locations."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            results = []
            query = search_text if case_sensitive else search_text.lower()
            for i, para in enumerate(doc.paragraphs):
                text = para.text
                target = text if case_sensitive else text.lower()
                start = 0
                while True:
                    idx = target.find(query, start)
                    if idx == -1:
                        break
                    results.append(
                        {
                            "paragraph_index": i,
                            "char_offset": idx,
                            "context": text[max(0, idx - 30) : idx + len(search_text) + 30],
                            "style": para.style.name if para.style else None,
                        }
                    )
                    start = idx + len(search_text)
            table_results = []
            for ti, table in enumerate(doc.tables):
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        text = cell.text
                        target = text if case_sensitive else text.lower()
                        if query in target:
                            table_results.append(
                                {"table_index": ti, "row": ri, "col": ci, "text": text[:200]}
                            )
            return {
                "success": True,
                "file": file_path,
                "search_text": search_text,
                "matches": len(results),
                "table_matches": len(table_results),
                "results": results,
                "table_results": table_results,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def read_tables(self, file_path: str) -> Dict[str, Any]:
        """Read all tables from a document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            tables = []
            for ti, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    rows_data.append([cell.text for cell in row.cells])
                headers = rows_data[0] if rows_data else []
                data = rows_data[1:] if len(rows_data) > 1 else []
                tables.append(
                    {
                        "table_index": ti,
                        "headers": headers,
                        "rows": data,
                        "row_count": len(data),
                        "col_count": len(headers),
                    }
                )
            return {
                "success": True,
                "file": file_path,
                "table_count": len(tables),
                "tables": tables,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_hyperlink(
        self, file_path: str, text: str, url: str, paragraph_index: int = -1
    ) -> Dict[str, Any]:
        """Add a hyperlink to the document. paragraph_index=-1 appends a new paragraph."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                if paragraph_index == -1:
                    para = doc.add_paragraph()
                else:
                    if paragraph_index >= len(doc.paragraphs):
                        return {
                            "success": False,
                            "error": f"Paragraph index {paragraph_index} out of range",
                        }
                    para = doc.paragraphs[paragraph_index]
                part = doc.part
                r_id = part.relate_to(
                    url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rStyle = OxmlElement("w:rStyle")
                rStyle.set(qn("w:val"), "Hyperlink")
                rPr.append(rStyle)
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "0563C1")
                rPr.append(color)
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "single")
                rPr.append(u)
                run.append(rPr)
                t = OxmlElement("w:t")
                t.text = text
                t.set(qn("xml:space"), "preserve")
                run.append(t)
                hyperlink.append(run)
                para._element.append(hyperlink)
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "text": text,
                "url": url,
                "paragraph_index": paragraph_index,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_page_break(self, file_path: str) -> Dict[str, Any]:
        """Add a page break at the end of the document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                doc.add_page_break()
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "action": "page_break_added",
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def list_styles(self, file_path: str) -> Dict[str, Any]:
        """List all available paragraph styles in the document."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            para_styles = []
            char_styles = []
            for s in doc.styles:
                entry = {"name": s.name, "builtin": s.builtin}
                if s.type == 1:
                    para_styles.append(entry)
                elif s.type == 2:
                    char_styles.append(entry)
            return {
                "success": True,
                "file": file_path,
                "paragraph_styles": para_styles,
                "character_styles": char_styles,
                "paragraph_style_count": len(para_styles),
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def set_metadata(
        self,
        file_path: str,
        author: str = None,
        title: str = None,
        subject: str = None,
        keywords: str = None,
        comments: str = None,
        category: str = None,
    ) -> Dict[str, Any]:
        """Set document core properties (metadata)."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                cp = doc.core_properties
                if author is not None:
                    cp.author = author
                if title is not None:
                    cp.title = title
                if subject is not None:
                    cp.subject = subject
                if keywords is not None:
                    cp.keywords = keywords
                if comments is not None:
                    cp.comments = comments
                if category is not None:
                    cp.category = category
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "metadata_set": {
                    k: v
                    for k, v in {
                        "author": author,
                        "title": title,
                        "subject": subject,
                        "keywords": keywords,
                        "comments": comments,
                        "category": category,
                    }.items()
                    if v is not None
                },
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Read document core properties (metadata)."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            cp = doc.core_properties
            return {
                "success": True,
                "file": file_path,
                "author": cp.author,
                "title": cp.title,
                "subject": cp.subject,
                "keywords": cp.keywords,
                "comments": cp.comments,
                "category": cp.category,
                "created": cp.created.isoformat() if cp.created else None,
                "modified": cp.modified.isoformat() if cp.modified else None,
                "last_modified_by": cp.last_modified_by,
                "revision": cp.revision,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def add_list(
        self, file_path: str, items: List[str], list_type: str = "bullet"
    ) -> Dict[str, Any]:
        """Add a bulleted or numbered list."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            style_name = "List Bullet" if list_type == "bullet" else "List Number"
            with self.host._file_lock(file_path):
                backup = self.host._snapshot_backup(file_path)
                doc = Document(file_path)
                available = [s.name for s in doc.styles if s.type == 1]
                if style_name not in available:
                    style_name = "Normal"
                for item in items:
                    doc.add_paragraph(item, style=style_name)
                self.host._safe_save(doc, file_path)
            return {
                "success": True,
                "file": file_path,
                "list_type": list_type,
                "items_added": len(items),
                "style_used": style_name,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def word_count(self, file_path: str) -> Dict[str, Any]:
        """Get word, character, paragraph, and page estimate counts."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            words = len(full_text.split())
            chars = len(full_text)
            chars_no_spaces = len(full_text.replace(" ", "").replace("\n", ""))
            paras = len([p for p in doc.paragraphs if p.text.strip()])
            pages_est = max(1, round(words / 250))
            return {
                "success": True,
                "file": file_path,
                "words": words,
                "characters": chars,
                "characters_no_spaces": chars_no_spaces,
                "paragraphs": paras,
                "pages_estimate": pages_est,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def get_formatting_info(
        self,
        file_path: str,
        *,
        start: int = 0,
        limit: Optional[int] = 10,
        all_paragraphs: bool = False,
    ) -> Dict[str, Any]:
        """Get detailed paragraph and section formatting information."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            start = max(0, int(start or 0))
            if limit is not None:
                limit = max(0, int(limit))
            doc = Document(file_path)
            section_indices = self._paragraph_section_indices(doc)
            paragraph_count = len(doc.paragraphs)
            end = paragraph_count if all_paragraphs or limit is None else start + limit
            selected = list(enumerate(doc.paragraphs))[start:end]
            info = {
                "success": True,
                "file": file_path,
                "paragraph_count": paragraph_count,
                "start": start,
                "limit": None if all_paragraphs else limit,
                "returned_paragraph_count": len(selected),
                "truncated": end < paragraph_count,
                "sections": [],
                "paragraphs": [],
                "ooxml": self._docx_ooxml_prefix_report(file_path),
            }
            for section in doc.sections:
                info["sections"].append(
                    {
                        "orientation": "landscape"
                        if section.page_width > section.page_height
                        else "portrait",
                        "margins": {
                            "top": round(section.top_margin.inches, 2),
                            "bottom": round(section.bottom_margin.inches, 2),
                            "left": round(section.left_margin.inches, 2),
                            "right": round(section.right_margin.inches, 2),
                        },
                        "page_width_in": round(section.page_width.inches, 3),
                        "page_height_in": round(section.page_height.inches, 3),
                    }
                )
            for i, para in selected:
                section_index = section_indices[i] if i < len(section_indices) else 0
                info["paragraphs"].append(
                    self._paragraph_format_payload(para, doc, i, section_index)
                )
            return info
        except Exception as e:
            return {"success": False, "error": str(e)}

    def inspect_hidden_data(self, file_path: str) -> Dict[str, Any]:
        """Inspect hidden DOCX data relevant to submission prep."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            abs_path = str(Path(file_path).resolve())
            preflight = self._docx_zip_preflight(abs_path)
            if not preflight.get("success"):
                return self._docx_preflight_error_result(abs_path, preflight)
            with ZipFile(abs_path, "r") as zin:
                names = zin.namelist()
                story_parts = self.host._docx_story_xml_parts(names)
                relationship_report = self._docx_external_relationship_report(
                    {
                        name: zin.read(name)
                        for name in names
                        if name.lower().endswith(".rels")
                    }
                )

                comment_reference_count = 0
                tracked_counts = {
                    "insertions": 0,
                    "deletions": 0,
                    "move_to": 0,
                    "move_from": 0,
                    "move_to_range_markers": 0,
                    "move_from_range_markers": 0,
                }
                for part_name in story_parts:
                    root = ET.fromstring(zin.read(part_name))
                    comment_reference_count += len(
                        root.findall(".//w:commentReference", OOXML_NS)
                    )
                    comment_reference_count += len(
                        root.findall(".//w:commentRangeStart", OOXML_NS)
                    )
                    tracked_counts["insertions"] += len(root.findall(".//w:ins", OOXML_NS))
                    tracked_counts["deletions"] += len(root.findall(".//w:del", OOXML_NS))
                    tracked_counts["move_to"] += len(root.findall(".//w:moveTo", OOXML_NS))
                    tracked_counts["move_from"] += len(
                        root.findall(".//w:moveFrom", OOXML_NS)
                    )
                    tracked_counts["move_to_range_markers"] += len(
                        root.findall(".//w:moveToRangeStart", OOXML_NS)
                    ) + len(root.findall(".//w:moveToRangeEnd", OOXML_NS))
                    tracked_counts["move_from_range_markers"] += len(
                        root.findall(".//w:moveFromRangeStart", OOXML_NS)
                    ) + len(root.findall(".//w:moveFromRangeEnd", OOXML_NS))
                    legacy_counted = {
                        "del",
                        "ins",
                        "moveFrom",
                        "moveFromRangeEnd",
                        "moveFromRangeStart",
                        "moveTo",
                        "moveToRangeEnd",
                        "moveToRangeStart",
                    }
                    for local, count in self._tracked_revision_counts(root).items():
                        if local not in legacy_counted:
                            tracked_counts[local] = tracked_counts.get(local, 0) + count

                comment_part_names = sorted(
                    name
                    for name in names
                    if (
                        name.lower().startswith("word/comments")
                        and name.lower().endswith(".xml")
                    )
                    or name.lower() == "word/people.xml"
                )
                comment_relationship_part_names = sorted(
                    name
                    for name in names
                    if (
                        name.lower().startswith("word/_rels/comments")
                        and name.lower().endswith(".xml.rels")
                    )
                    or name.lower() == "word/_rels/people.xml.rels"
                )
                comment_related_targets = []
                for rel_part in comment_relationship_part_names:
                    comment_related_targets.extend(
                        self._docx_relationship_internal_targets(
                            rel_part, zin.read(rel_part)
                        )
                    )
                comment_related_targets = sorted(set(comment_related_targets))
                comments_count = 0
                for comment_part in comment_part_names:
                    try:
                        root = ET.fromstring(zin.read(comment_part))
                    except ET.ParseError:
                        continue
                    comments_count += len(root.findall(".//w:comment", OOXML_NS))

                settings_flag = False
                if "word/settings.xml" in names:
                    root = ET.fromstring(zin.read("word/settings.xml"))
                    flag = root.find(".//w:removePersonalInformation", OOXML_NS)
                    settings_flag = bool(
                        flag is not None
                        and str(flag.get(f"{{{OOXML_NS['w']}}}val", "true")).lower()
                        != "false"
                    )

                core_props = {}
                if "docProps/core.xml" in names:
                    root = ET.fromstring(zin.read("docProps/core.xml"))
                    core_props = {
                        "author": (
                            root.findtext("dc:creator", default="", namespaces=OOXML_NS)
                            or ""
                        ),
                        "title": (
                            root.findtext("dc:title", default="", namespaces=OOXML_NS)
                            or ""
                        ),
                        "subject": (
                            root.findtext("dc:subject", default="", namespaces=OOXML_NS)
                            or ""
                        ),
                        "keywords": (
                            root.findtext("cp:keywords", default="", namespaces=OOXML_NS)
                            or ""
                        ),
                        "description": (
                            root.findtext(
                                "dc:description", default="", namespaces=OOXML_NS
                            )
                            or ""
                        ),
                        "category": (
                            root.findtext("cp:category", default="", namespaces=OOXML_NS)
                            or ""
                        ),
                        "last_modified_by": (
                            root.findtext(
                                "cp:lastModifiedBy", default="", namespaces=OOXML_NS
                            )
                            or ""
                        ),
                        "created": (
                            root.findtext(
                                "dcterms:created", default="", namespaces=OOXML_NS
                            )
                            or ""
                        ),
                        "modified": (
                            root.findtext(
                                "dcterms:modified", default="", namespaces=OOXML_NS
                            )
                            or ""
                        ),
                    }

                app_props = {}
                if "docProps/app.xml" in names:
                    root = ET.fromstring(zin.read("docProps/app.xml"))
                    app_props = {
                        "application": root.findtext(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Application",
                            default="",
                        ),
                        "company": root.findtext(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Company",
                            default="",
                        ),
                        "manager": root.findtext(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Manager",
                            default="",
                        ),
                        "hyperlink_base": root.findtext(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}HyperlinkBase",
                            default="",
                        ),
                    }

                custom_doc_properties = []
                if "docProps/custom.xml" in names:
                    root = ET.fromstring(zin.read("docProps/custom.xml"))
                    for prop in list(root):
                        if prop.tag.rsplit("}", 1)[-1] != "property":
                            continue
                        entry = dict(self._local_attrs(prop))
                        value_node = next(iter(list(prop)), None)
                        entry["value"] = value_node.text if value_node is not None else ""
                        entry["value_type"] = (
                            value_node.tag.rsplit("}", 1)[-1]
                            if value_node is not None
                            else None
                        )
                        custom_doc_properties.append(entry)

                custom_xml = self.host._custom_xml_summary(names)

            doc = Document(abs_path)
            first_section = doc.sections[0] if doc.sections else None
            page_size = None
            orientation = None
            if first_section is not None:
                page_size = self.host._label_page_size(
                    first_section.page_width, first_section.page_height
                )
                orientation = (
                    "landscape"
                    if first_section.page_width > first_section.page_height
                    else "portrait"
                )

            tracked_total = sum(tracked_counts.values())
            return {
                "success": True,
                "file": abs_path,
                "ooxml_preflight": preflight,
                "page_size": page_size,
                "orientation": orientation,
                "section_count": len(doc.sections),
                "comments_part_present": bool(
                    comment_part_names
                    or comment_relationship_part_names
                    or comment_related_targets
                ),
                "comments_extended_part_present": "word/commentsExtended.xml" in names,
                "people_part_present": "word/people.xml" in names,
                "comment_parts": comment_part_names,
                "comment_part_count": len(comment_part_names),
                "comment_relationship_parts": comment_relationship_part_names,
                "comment_relationship_part_count": len(comment_relationship_part_names),
                "comment_related_targets": comment_related_targets,
                "comment_related_target_count": len(comment_related_targets),
                "comments_count": comments_count,
                "comment_reference_count": comment_reference_count,
                "tracked_change_count": tracked_total,
                "tracked_changes_present": tracked_total > 0,
                "tracked_changes": tracked_counts,
                "remove_personal_information": settings_flag,
                "custom_xml_item_count": len(custom_xml["payload_parts"]),
                "custom_xml_part_count": len(custom_xml["payload_parts"]),
                "custom_xml_parts_total_count": len(custom_xml["all_parts"]),
                "custom_xml_support_part_count": len(custom_xml["support_parts"]),
                "custom_xml_parts": custom_xml["payload_parts"],
                "custom_xml_support_parts": custom_xml["support_parts"],
                "custom_document_properties_part_present": "docProps/custom.xml" in names,
                "custom_document_properties_count": len(custom_doc_properties),
                "custom_document_properties": custom_doc_properties,
                "core_properties": core_props,
                "app_properties": app_props,
                "external_relationships": relationship_report[
                    "external_relationships"
                ],
                "external_relationship_count": relationship_report[
                    "external_relationship_count"
                ],
                "risky_external_relationships": relationship_report[
                    "risky_external_relationships"
                ],
                "risky_external_relationship_count": relationship_report[
                    "risky_external_relationship_count"
                ],
                "warning_external_relationships": relationship_report[
                    "warning_external_relationships"
                ],
                "warning_external_relationship_count": relationship_report[
                    "warning_external_relationship_count"
                ],
                "external_hyperlink_relationships": relationship_report[
                    "external_hyperlink_relationships"
                ],
                "external_hyperlink_relationship_count": relationship_report[
                    "external_hyperlink_relationship_count"
                ],
                "relationship_parts": relationship_report["relationship_parts"],
                "relationship_part_count": relationship_report[
                    "relationship_part_count"
                ],
                "relationship_parse_errors": relationship_report[
                    "relationship_parse_errors"
                ],
                "relationship_parse_error_count": relationship_report[
                    "relationship_parse_error_count"
                ],
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def audit_document_fonts(
        self,
        file_path: str,
        expected_font_name: str = None,
        expected_font_size: float = None,
    ) -> Dict[str, Any]:
        """Audit font consistency across visible text runs."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            runs_payload = self.host._collect_text_runs(doc)
            font_name_counts = Counter(
                run["font_name"] or "<unspecified>" for run in runs_payload
            )
            font_size_counts = Counter(
                run["font_size"] if run["font_size"] is not None else "<unspecified>"
                for run in runs_payload
            )
            dominant_font = (
                font_name_counts.most_common(1)[0][0] if font_name_counts else None
            )
            dominant_size = (
                font_size_counts.most_common(1)[0][0] if font_size_counts else None
            )

            unexpected_font_runs = []
            if expected_font_name:
                target = self.host._normalized_font_name(expected_font_name)
                unexpected_font_runs = [
                    run for run in runs_payload if run["font_name_normalized"] != target
                ]

            unexpected_size_runs = []
            if expected_font_size is not None:
                target_size = round(float(expected_font_size), 2)
                unexpected_size_runs = [
                    run for run in runs_payload if run["font_size"] != target_size
                ]

            mixed_fonts = len(font_name_counts) > 1
            mixed_sizes = len(font_size_counts) > 1
            return {
                "success": True,
                "file": file_path,
                "text_run_count": len(runs_payload),
                "font_name_counts": dict(font_name_counts),
                "font_size_counts": dict(font_size_counts),
                "dominant_font_name": dominant_font,
                "dominant_font_size": dominant_size,
                "mixed_fonts": mixed_fonts,
                "mixed_sizes": mixed_sizes,
                "expected_font_name": expected_font_name,
                "expected_font_size": expected_font_size,
                "unexpected_font_run_count": len(unexpected_font_runs),
                "unexpected_size_run_count": len(unexpected_size_runs),
                "unexpected_font_examples": unexpected_font_runs[:10],
                "unexpected_size_examples": unexpected_size_runs[:10],
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def audit_document_images(self, file_path: str) -> Dict[str, Any]:
        """Audit inline image sizing against available page content area."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            doc = Document(file_path)
            sections = self.host._document_sections_summary(doc)
            max_usable_width = max(
                (s["usable_width_in"] for s in sections), default=0.0
            )
            max_usable_height = max(
                (s["usable_height_in"] for s in sections), default=0.0
            )

            oversized = []
            near_limit = []
            images = []
            for idx, shape in enumerate(doc.inline_shapes):
                width_in = round(shape.width.inches, 3)
                height_in = round(shape.height.inches, 3)
                width_ratio = (
                    round(width_in / max_usable_width, 3) if max_usable_width else None
                )
                height_ratio = (
                    round(height_in / max_usable_height, 3)
                    if max_usable_height
                    else None
                )
                image_info = {
                    "index": idx,
                    "width_in": width_in,
                    "height_in": height_in,
                    "width_ratio_to_max_content": width_ratio,
                    "height_ratio_to_max_content": height_ratio,
                }
                if max_usable_width and width_in > max_usable_width + 0.02:
                    oversized.append(image_info)
                elif max_usable_width and width_ratio is not None and width_ratio >= 0.95:
                    near_limit.append(image_info)
                images.append(image_info)

            return {
                "success": True,
                "file": file_path,
                "inline_image_count": len(images),
                "section_max_usable_width_in": max_usable_width,
                "section_max_usable_height_in": max_usable_height,
                "images": images,
                "oversized_image_count": len(oversized),
                "oversized_images": oversized[:10],
                "near_limit_image_count": len(near_limit),
                "near_limit_images": near_limit[:10],
                "note": "Only inline images are audited. Floating/anchored images are not inspected by python-docx.",
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def document_preflight(
        self,
        file_path: str,
        expected_page_size: str = None,
        expected_font_name: str = None,
        expected_font_size: float = None,
        rendered_layout: bool = False,
        render_profile: str = "auto",
    ) -> Dict[str, Any]:
        """Run a submission-oriented DOCX preflight."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            abs_path = str(Path(file_path).resolve())
            hidden = self.inspect_hidden_data(abs_path)
            if not hidden.get("success"):
                return hidden
            fonts = self.audit_document_fonts(
                abs_path,
                expected_font_name=expected_font_name,
                expected_font_size=expected_font_size,
            )
            if not fonts.get("success"):
                return fonts
            images = self.audit_document_images(abs_path)
            if not images.get("success"):
                return images
            ooxml = self._docx_ooxml_prefix_report(abs_path)

            sections = self.host._document_sections_summary(Document(abs_path))
            section_page_sizes = sorted({s["page_size"] for s in sections})
            page_size_consistent = len(section_page_sizes) <= 1

            checks = []

            def add_check(
                name: str, status: str, message: str, details: Dict[str, Any] = None
            ):
                checks.append(
                    {
                        "name": name,
                        "status": status,
                        "message": message,
                        "details": details or {},
                    }
                )

            if expected_page_size:
                matches_expected = (
                    page_size_consistent
                    and len(section_page_sizes) == 1
                    and section_page_sizes[0].lower() == expected_page_size.lower()
                )
                add_check(
                    "page_size",
                    "pass" if matches_expected else "warn",
                    (
                        f"All sections match expected page size {expected_page_size}."
                        if matches_expected
                        else f"Detected page sizes {section_page_sizes}; expected {expected_page_size}."
                    ),
                    {"detected": section_page_sizes, "expected": expected_page_size},
                )
            else:
                add_check(
                    "page_size",
                    "pass" if page_size_consistent else "warn",
                    (
                        f"All sections use {section_page_sizes[0]}."
                        if page_size_consistent and section_page_sizes
                        else f"Mixed section page sizes detected: {section_page_sizes}."
                    ),
                    {"detected": section_page_sizes},
                )

            package_preflight = hidden.get("ooxml_preflight", {})
            add_check(
                "ooxml_package",
                "pass" if package_preflight.get("status") == "pass" else "fail",
                (
                    "DOCX OOXML package passed ZIP safety preflight."
                    if package_preflight.get("status") == "pass"
                    else "DOCX OOXML package failed ZIP safety preflight."
                ),
                package_preflight,
            )

            external_count = hidden.get("external_relationship_count", 0)
            risky_external_count = hidden.get("risky_external_relationship_count", 0)
            external_status = (
                "fail"
                if risky_external_count
                else "warn"
                if external_count
                else "pass"
            )
            add_check(
                "external_relationships",
                external_status,
                (
                    "No external relationships detected."
                    if external_count == 0
                    else (
                        f"Detected {risky_external_count} risky external relationship(s) "
                        f"among {external_count} external relationship(s)."
                        if risky_external_count
                        else f"Detected {external_count} external relationship(s) requiring review."
                    )
                ),
                {
                    "external_relationships": hidden.get("external_relationships", []),
                    "risky_external_relationships": hidden.get(
                        "risky_external_relationships", []
                    ),
                    "warning_external_relationships": hidden.get(
                        "warning_external_relationships", []
                    ),
                    "relationship_parse_errors": hidden.get(
                        "relationship_parse_errors", []
                    ),
                },
            )

            add_check(
                "comments",
                "pass"
                if not hidden.get("comments_part_present")
                and hidden["comments_count"] == 0
                and hidden["comment_reference_count"] == 0
                else "fail",
                (
                    "No comment parts or references detected."
                    if not hidden.get("comments_part_present")
                    and hidden["comments_count"] == 0
                    and hidden["comment_reference_count"] == 0
                    else (
                        "Detected comment package parts, relationships, related targets, "
                        f"{hidden['comments_count']} comments, and "
                        f"{hidden['comment_reference_count']} comment references."
                    )
                ),
                {
                    "comments_part_present": hidden.get("comments_part_present"),
                    "comment_part_count": hidden.get("comment_part_count"),
                    "comment_relationship_part_count": hidden.get(
                        "comment_relationship_part_count"
                    ),
                    "comment_related_target_count": hidden.get(
                        "comment_related_target_count"
                    ),
                    "comments_count": hidden["comments_count"],
                    "comment_reference_count": hidden["comment_reference_count"],
                },
            )

            add_check(
                "tracked_changes",
                "pass" if not hidden["tracked_changes_present"] else "fail",
                (
                    "No tracked revision markup detected."
                    if not hidden["tracked_changes_present"]
                    else f"Detected {hidden['tracked_change_count']} tracked-change elements."
                ),
                hidden["tracked_changes"],
            )

            metadata_values = hidden["core_properties"].copy()
            metadata_values.update(hidden["app_properties"])
            for prop in hidden.get("custom_document_properties", []):
                metadata_values[f"custom:{prop.get('name', '<unnamed>')}"] = prop.get(
                    "value", ""
                )
            nonempty_metadata = {
                k: v for k, v in metadata_values.items() if str(v or "").strip()
            }
            add_check(
                "metadata",
                "pass" if not nonempty_metadata else "warn",
                (
                    "No non-empty hidden core/app metadata fields detected."
                    if not nonempty_metadata
                    else f"Detected non-empty metadata fields: {sorted(nonempty_metadata)}."
                ),
                nonempty_metadata,
            )

            ooxml_ready = not ooxml.get("submission_blocking")
            if ooxml_ready:
                ooxml_message = "All WordprocessingML parts use canonical w: prefixes."
            elif ooxml.get("parse_error_part_count"):
                ooxml_message = (
                    f"{ooxml['parse_error_part_count']} WordprocessingML part(s) "
                    "could not be parsed for OOXML readiness."
                )
            else:
                ooxml_message = (
                    f"{ooxml['noncanonical_part_count']} WordprocessingML part(s) "
                    "use non-canonical prefixes that OnlyOffice x2t may not render reliably."
                )
            add_check(
                "ooxml_word_prefix",
                "pass" if ooxml_ready else "fail",
                ooxml_message,
                ooxml,
            )

            add_check(
                "custom_xml",
                "pass" if hidden["custom_xml_item_count"] == 0 else "warn",
                (
                    "No custom XML payload items detected."
                    if hidden["custom_xml_item_count"] == 0
                    else f"Detected {hidden['custom_xml_item_count']} custom XML payload item(s)."
                ),
                {
                    "custom_xml_items": hidden["custom_xml_parts"],
                    "custom_xml_support_parts": hidden["custom_xml_support_parts"],
                },
            )

            add_check(
                "remove_personal_information_flag",
                "pass" if hidden["remove_personal_information"] else "warn",
                (
                    "removePersonalInformation flag is enabled."
                    if hidden["remove_personal_information"]
                    else "removePersonalInformation flag is not enabled."
                ),
            )

            font_status = "pass"
            font_message = "Document font usage looks consistent."
            if expected_font_name and fonts["unexpected_font_run_count"] > 0:
                font_status = "warn"
                font_message = (
                    f"Detected {fonts['unexpected_font_run_count']} run(s) not matching expected font '{expected_font_name}'."
                )
            elif fonts["mixed_fonts"]:
                font_status = "warn"
                font_message = (
                    f"Multiple font families detected; dominant font is '{fonts['dominant_font_name']}'."
                )
            add_check(
                "font_names",
                font_status,
                font_message,
                {
                    "font_name_counts": fonts["font_name_counts"],
                    "dominant_font_name": fonts["dominant_font_name"],
                    "unexpected_examples": fonts["unexpected_font_examples"],
                },
            )

            size_status = "pass"
            size_message = "Document font sizes look consistent."
            if expected_font_size is not None and fonts["unexpected_size_run_count"] > 0:
                size_status = "warn"
                size_message = (
                    f"Detected {fonts['unexpected_size_run_count']} run(s) not matching expected font size {expected_font_size}."
                )
            elif fonts["mixed_sizes"]:
                size_status = "warn"
                size_message = (
                    f"Multiple font sizes detected; dominant size is {fonts['dominant_font_size']}."
                )
            add_check(
                "font_sizes",
                size_status,
                size_message,
                {
                    "font_size_counts": fonts["font_size_counts"],
                    "dominant_font_size": fonts["dominant_font_size"],
                    "unexpected_examples": fonts["unexpected_size_examples"],
                },
            )

            image_status = "pass"
            image_message = "Inline images fit within the maximum content area."
            if images["oversized_image_count"] > 0:
                image_status = "warn"
                image_message = (
                    f"Detected {images['oversized_image_count']} inline image(s) wider than the available content width."
                )
            elif images["near_limit_image_count"] > 0:
                image_status = "warn"
                image_message = (
                    f"Detected {images['near_limit_image_count']} inline image(s) at or near the available content width."
                )
            add_check(
                "images",
                image_status,
                image_message,
                {
                    "oversized_images": images["oversized_images"],
                    "near_limit_images": images["near_limit_images"],
                    "note": images["note"],
                },
            )

            rendered_audit = None
            if rendered_layout:
                rendered_audit = self.rendered_layout_audit(
                    abs_path,
                    profile=render_profile,
                    trusted_pdf=True,
                )
                if not rendered_audit.get("success"):
                    return rendered_audit
                add_check(
                    "rendered_layout",
                    rendered_audit["overall_status"],
                    (
                        "Rendered PDF layout matches DOCX reference/margin expectations."
                        if rendered_audit["overall_status"] == "pass"
                        else "Rendered PDF layout does not match DOCX reference/margin expectations."
                    ),
                    {
                        "overall_status": rendered_audit["overall_status"],
                        "source_unchanged": rendered_audit.get("source_unchanged"),
                        "checks": rendered_audit.get("checks", []),
                    },
                )

            overall = "pass"
            if any(check["status"] == "fail" for check in checks):
                overall = "fail"
            elif any(check["status"] == "warn" for check in checks):
                overall = "warn"
            readiness_blockers = [
                {
                    "name": check["name"],
                    "status": check["status"],
                    "message": check["message"],
                    "details": check.get("details", {}),
                }
                for check in checks
                if check["status"] in {"fail", "warn"}
            ]
            if not rendered_layout:
                readiness_blockers.append(
                    {
                        "name": "rendered_layout_not_run",
                        "status": "blocked",
                        "message": (
                            "Submission-ready status requires a rendered layout audit; "
                            "conversion or OOXML inspection alone is insufficient."
                        ),
                        "details": {"required_flag": "--rendered-layout"},
                    }
                )
            rendered_ready = (
                rendered_layout
                and rendered_audit is not None
                and rendered_audit.get("overall_status") == "pass"
            )

            return {
                "success": True,
                "file": abs_path,
                "overall_status": overall,
                "submission_ready": overall == "pass" and rendered_ready,
                "readiness_blockers": readiness_blockers,
                "checks": checks,
                "sections": sections,
                "hidden_data": hidden,
                "font_audit": fonts,
                "image_audit": images,
                "ooxml": ooxml,
                "rendered_layout_audit": rendered_audit,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def _reference_layout_expectations(self, doc: Any) -> Dict[str, Any]:
        section_indices = self._paragraph_section_indices(doc)
        heading_index = None
        for idx, para in enumerate(doc.paragraphs):
            if self._normalize_text(para.text).lower() == "references":
                heading_index = idx
                break

        if heading_index is None:
            return {
                "references_heading_index": None,
                "references_heading_section_index": None,
                "references_heading_expected_page_break": False,
                "reference_paragraphs": [],
            }

        heading = doc.paragraphs[heading_index]
        heading_break = self._paragraph_page_break_before_info(heading, doc)
        heading_inline_page_break = self._inline_page_break_before_text(heading)
        previous_inline_page_break = False
        scan_index = heading_index - 1
        while scan_index >= 0:
            prev = doc.paragraphs[scan_index]
            previous_inline_page_break = any(
                br.get("type") == "page" for br in self._inline_page_breaks(prev)
            )
            if previous_inline_page_break or self._normalize_text(prev.text):
                break
            scan_index -= 1

        refs = []
        for idx, para in enumerate(doc.paragraphs[heading_index + 1 :], start=heading_index + 1):
            text = self._normalize_text(para.text)
            if not text:
                continue
            indent = self._paragraph_indent_info(para, doc)
            resolved = indent["style_resolved"]
            left_pt = (
                (resolved.get("effective_left") or resolved.get("left") or {}).get("pt")
                or 0.0
            )
            hanging_pt = (resolved.get("hanging") or {}).get("pt") or 0.0
            first_line_pt = (resolved.get("first_line") or {}).get("pt") or 0.0
            right_pt = (
                (resolved.get("effective_right") or resolved.get("right") or {}).get("pt")
                or 0.0
            )
            first_offset_pt = left_pt + first_line_pt
            if hanging_pt:
                first_offset_pt = left_pt - hanging_pt
            refs.append(
                {
                    "paragraph_index": idx,
                    "section_index": section_indices[idx] if idx < len(section_indices) else 0,
                    "text_preview": text[:180],
                    "text": text,
                    "expected_left_indent_pt": round(left_pt, 3),
                    "expected_right_indent_pt": round(right_pt, 3),
                    "expected_hanging_indent_pt": round(hanging_pt, 3),
                    "expected_first_line_offset_pt": round(first_offset_pt, 3),
                    "expected_continuation_offset_pt": round(left_pt, 3),
                    "indents": indent,
                }
            )

        return {
            "references_heading_index": heading_index,
            "references_heading_section_index": section_indices[heading_index]
            if heading_index < len(section_indices)
            else 0,
            "references_heading_expected_page_break": bool(
                previous_inline_page_break
                or heading_inline_page_break
                or heading_break.get("style_resolved") is True
                or heading_break.get("direct") is True
                or heading_break.get("enabled_in_ooxml")
            ),
            "references_heading_page_break": heading_break,
            "references_heading_inline_page_break_before_text": heading_inline_page_break,
            "previous_paragraph_inline_page_break": previous_inline_page_break,
            "reference_paragraphs": refs,
        }

    @staticmethod
    def _section_margin_points(
        sections: List[Dict[str, Any]], section_index: Optional[int]
    ) -> Dict[str, float]:
        if sections:
            try:
                idx = int(section_index) if section_index is not None else 0
            except (TypeError, ValueError):
                idx = 0
            idx = min(max(idx, 0), len(sections) - 1)
            margins_in = sections[idx].get("margins_in", {})
        else:
            idx = 0
            margins_in = {}
        return {
            "section_index": idx,
            "left": float(margins_in.get("left", 1.0)) * 72.0,
            "right": float(margins_in.get("right", 1.0)) * 72.0,
            "top": float(margins_in.get("top", 1.0)) * 72.0,
            "bottom": float(margins_in.get("bottom", 1.0)) * 72.0,
        }

    def _audit_generic_rendered_margins(
        self,
        lines: List[Dict[str, Any]],
        sections: List[Dict[str, Any]],
        tolerance_points: float,
    ) -> Dict[str, Any]:
        """Audit generic rendered text against a single-section margin envelope."""
        if not lines:
            return {
                "status": "warn",
                "message": "No rendered text lines were available for generic margin checks.",
                "details": {"lines_checked": 0, "violations": []},
            }
        if len(sections) != 1:
            return {
                "status": "warn",
                "message": (
                    "Generic rendered audit cannot map text lines to section-specific "
                    f"margins when the DOCX has {len(sections)} sections."
                ),
                "details": {
                    "section_count": len(sections),
                    "lines_checked": 0,
                    "requires": "doc-render-map anchors or a profile-specific audit",
                },
            }

        margins = self._section_margin_points(sections, 0)
        violations = []
        checked = 0
        for line in lines:
            page_width = float(line.get("page_width") or 0)
            page_height = float(line.get("page_height") or 0)
            if page_width <= 0 or page_height <= 0:
                violations.append(
                    {
                        "page_number": line.get("page_number"),
                        "text": line.get("text", "")[:120],
                        "reason": "missing_page_geometry",
                    }
                )
                continue
            checked += 1
            expected_left = margins["left"]
            expected_right = page_width - margins["right"]
            expected_top = margins["top"]
            expected_bottom = page_height - margins["bottom"]
            if line["left"] < expected_left - tolerance_points:
                violations.append(
                    {
                        "page_number": line.get("page_number"),
                        "text": line.get("text", "")[:120],
                        "left": round(line["left"], 3),
                        "expected_left": round(expected_left, 3),
                        "under_by": round(expected_left - line["left"], 3),
                    }
                )
            if line["right"] > expected_right + tolerance_points:
                violations.append(
                    {
                        "page_number": line.get("page_number"),
                        "text": line.get("text", "")[:120],
                        "right": round(line["right"], 3),
                        "expected_right": round(expected_right, 3),
                        "over_by": round(line["right"] - expected_right, 3),
                    }
                )
            if line["top"] < expected_top - tolerance_points:
                violations.append(
                    {
                        "page_number": line.get("page_number"),
                        "text": line.get("text", "")[:120],
                        "top": round(line["top"], 3),
                        "expected_top": round(expected_top, 3),
                        "under_by": round(expected_top - line["top"], 3),
                    }
                )
            if line["bottom"] > expected_bottom + tolerance_points:
                violations.append(
                    {
                        "page_number": line.get("page_number"),
                        "text": line.get("text", "")[:120],
                        "bottom": round(line["bottom"], 3),
                        "expected_bottom": round(expected_bottom, 3),
                        "over_by": round(line["bottom"] - expected_bottom, 3),
                    }
                )

        return {
            "status": "pass" if not violations and checked > 0 else "fail",
            "message": (
                f"Checked {checked} rendered text line(s) against the section margin envelope."
                if not violations and checked > 0
                else f"{len(violations)} rendered text line(s) exceeded the section margin envelope."
            ),
            "details": {
                "section_index": margins["section_index"],
                "left_margin_pt": round(margins["left"], 3),
                "right_margin_pt": round(margins["right"], 3),
                "top_margin_pt": round(margins["top"], 3),
                "bottom_margin_pt": round(margins["bottom"], 3),
                "lines_checked": checked,
                "violations": violations[:20],
            },
        }

    @staticmethod
    def _pdf_lines_from_blocks(block_payload: Dict[str, Any]) -> List[Dict[str, Any]]:
        lines = []
        for page in block_payload.get("pages", []):
            for block in page.get("blocks", []):
                if block.get("type") != "text":
                    continue
                for line in block.get("lines", []):
                    text = re.sub(r"\s+", " ", str(line.get("text", "") or "")).strip()
                    if not text:
                        continue
                    bbox = line.get("bbox") or {}
                    lines.append(
                        {
                            "page_index": page.get("page_index"),
                            "page_number": page.get("page_number"),
                            "page_width": float(page.get("width") or 0),
                            "page_height": float(page.get("height") or 0),
                            "line_id": line.get("line_id"),
                            "text": text,
                            "bbox": bbox,
                            "left": float(bbox.get("left") or 0),
                            "right": float(bbox.get("right") or 0),
                            "top": float(bbox.get("top") or 0),
                            "bottom": float(bbox.get("bottom") or 0),
                        }
                    )
        return lines

    def _line_matches_reference_start(self, line_text: str, reference_text: str) -> bool:
        line_norm = self._normalize_text(line_text).lower()
        ref_norm = self._normalize_text(reference_text).lower()
        if not line_norm or not ref_norm:
            return False
        probe_len = min(len(line_norm), len(ref_norm), 48)
        if probe_len >= 24 and ref_norm.startswith(line_norm[:probe_len]):
            return True
        return line_norm[:32] in ref_norm[:120] if len(line_norm) >= 32 else False

    def rendered_layout_audit(
        self,
        file_path: str,
        *,
        pdf_path: Optional[str] = None,
        tolerance_points: float = 6.0,
        profile: str = "auto",
        trusted_pdf: bool = False,
    ) -> Dict[str, Any]:
        """Audit OnlyOffice-rendered PDF geometry against DOCX layout intent."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        temp_pdf = None
        try:
            requested_profile = str(profile or "auto").strip().lower()
            profile_aliases = {
                "apa": "apa-references",
                "reference": "apa-references",
                "references": "apa-references",
            }
            requested_profile = profile_aliases.get(requested_profile, requested_profile)
            valid_profiles = {"auto", "generic", "apa-references"}
            if requested_profile not in valid_profiles:
                return {
                    "success": False,
                    "error": (
                        "Unsupported render audit profile: "
                        f"{profile}. Use auto, generic, or apa-references."
                    ),
                }

            abs_path = str(Path(file_path).resolve())
            if not os.path.exists(abs_path):
                return {"success": False, "error": f"File not found: {file_path}"}

            hash_before = self._sha256_file(abs_path)
            doc = Document(abs_path)
            sections = self.host._document_sections_summary(doc)
            expectations = self._reference_layout_expectations(doc)
            prefix_report = self._docx_ooxml_prefix_report(abs_path)
            detected_references = (
                expectations.get("references_heading_index") is not None
            )
            effective_profile = (
                "apa-references"
                if requested_profile == "auto" and detected_references
                else "generic"
                if requested_profile == "auto"
                else requested_profile
            )

            conversion = None
            if pdf_path is None:
                fd, temp_pdf = tempfile.mkstemp(
                    prefix="onlyoffice-render-audit-", suffix=".pdf"
                )
                os.close(fd)
                conversion = self.host._office_to_pdf(abs_path, output_path=temp_pdf)
                if not conversion.get("success"):
                    return conversion
                audit_pdf_path = temp_pdf
                pdf_trusted = True
            else:
                audit_pdf_path = str(Path(pdf_path).resolve())
                pdf_trusted = bool(trusted_pdf)

            block_payload = self.host.pdf_read_blocks(
                audit_pdf_path,
                include_spans=False,
                include_images=False,
                include_empty=False,
            )
            if not block_payload.get("success"):
                return block_payload

            lines = self._pdf_lines_from_blocks(block_payload)
            checks: List[Dict[str, Any]] = []

            def add_check(
                name: str,
                status: str,
                message: str,
                details: Optional[Dict[str, Any]] = None,
            ) -> None:
                checks.append(
                    {
                        "name": name,
                        "status": status,
                        "message": message,
                        "details": details or {},
                    }
                )

            ooxml_ready = not prefix_report.get("submission_blocking")
            if ooxml_ready:
                ooxml_message = "All WordprocessingML parts use canonical w: prefixes."
            elif prefix_report.get("parse_error_part_count"):
                ooxml_message = (
                    f"{prefix_report['parse_error_part_count']} WordprocessingML "
                    "part(s) could not be parsed for rendered-layout readiness."
                )
            else:
                ooxml_message = (
                    f"{prefix_report['noncanonical_part_count']} WordprocessingML "
                    "part(s) use non-canonical prefixes that OnlyOffice x2t may not render reliably."
                )
            add_check(
                "ooxml_word_prefix",
                "pass" if ooxml_ready else "fail",
                ooxml_message,
                prefix_report,
            )

            add_check(
                "render_profile",
                "pass",
                (
                    "Using APA References rendered-layout checks."
                    if effective_profile == "apa-references"
                    else "Using generic rendered-layout checks; APA References-specific checks are not required for this document."
                ),
                {
                    "requested_profile": requested_profile,
                    "effective_profile": effective_profile,
                    "references_heading_detected": detected_references,
                },
            )
            add_check(
                "pdf_provenance",
                "pass" if pdf_trusted else "warn",
                (
                    "PDF was generated by this audit/conversion pipeline."
                    if pdf_trusted
                    else "PDF was supplied externally; geometry can be inspected, but submission-ready status is blocked because provenance is unverified."
                ),
                {
                    "trusted_pdf": pdf_trusted,
                    "temporary_pdf_used": pdf_path is None,
                    "supplied_pdf": str(Path(pdf_path).resolve()) if pdf_path else None,
                },
            )
            add_check(
                "rendered_pages",
                "pass" if int(block_payload.get("pages_scanned") or 0) > 0 else "warn",
                (
                    "Rendered PDF page blocks were read."
                    if int(block_payload.get("pages_scanned") or 0) > 0
                    else "No rendered PDF pages were scanned."
                ),
                {
                    "total_pages": block_payload.get("total_pages"),
                    "pages_scanned": block_payload.get("pages_scanned"),
                },
            )
            add_check(
                "pdf_text_rendered",
                "pass" if lines else "warn",
                (
                    "Rendered PDF contains extractable text lines."
                    if lines
                    else "Rendered PDF did not expose extractable text lines for geometry checks."
                ),
                {"line_count": len(lines)},
            )

            def finalize(
                rendered_heading: Optional[Dict[str, Any]],
                reference_paragraphs: Optional[List[Dict[str, Any]]],
            ) -> Dict[str, Any]:
                hash_after = self._sha256_file(abs_path)
                add_check(
                    "source_unchanged",
                    "pass" if hash_before == hash_after else "fail",
                    (
                        "Rendered audit did not modify the source DOCX."
                        if hash_before == hash_after
                        else "Rendered audit changed the source DOCX; submission-ready status is blocked."
                    ),
                    {
                        "source_hash_before": hash_before,
                        "source_hash_after": hash_after,
                    },
                )
                overall = "pass"
                if any(check["status"] == "fail" for check in checks):
                    overall = "fail"
                elif any(check["status"] == "warn" for check in checks):
                    overall = "warn"
                readiness_blockers = [
                    {
                        "name": check["name"],
                        "status": check["status"],
                        "message": check["message"],
                        "details": check.get("details", {}),
                    }
                    for check in checks
                    if check["status"] in {"fail", "warn"}
                ]
                refs = reference_paragraphs or []
                return {
                    "success": True,
                    "file": abs_path,
                    "pdf_file": str(Path(pdf_path).resolve()) if pdf_path else None,
                    "temporary_pdf_used": pdf_path is None,
                    "conversion": conversion,
                    "profile": effective_profile,
                    "requested_profile": requested_profile,
                    "source_hash_before": hash_before,
                    "source_hash_after": hash_after,
                    "source_unchanged": hash_before == hash_after,
                    "overall_status": overall,
                    "submission_ready": overall == "pass" and hash_before == hash_after,
                    "readiness_blockers": readiness_blockers,
                    "tolerance_points": tolerance_points,
                    "checks": checks,
                    "sections": sections,
                    "expected_layout": {
                        key: value
                        for key, value in expectations.items()
                        if key != "reference_paragraphs"
                    },
                    "reference_paragraph_count": len(refs),
                    "rendered": {
                        "total_pages": block_payload.get("total_pages"),
                        "pages_scanned": block_payload.get("pages_scanned"),
                        "line_count": len(lines),
                        "references_heading": rendered_heading,
                    },
                    "ooxml": prefix_report,
                }

            if effective_profile == "generic":
                generic_margin = self._audit_generic_rendered_margins(
                    lines, sections, tolerance_points
                )
                add_check(
                    "generic_margin_envelope",
                    generic_margin["status"],
                    generic_margin["message"],
                    generic_margin["details"],
                )
                return finalize(None, [])

            ref_line_index = next(
                (
                    idx
                    for idx, line in enumerate(lines)
                    if self._normalize_text(line["text"]).lower() == "references"
                ),
                None,
            )
            if ref_line_index is None:
                add_check(
                    "references_heading_rendered",
                    "warn",
                    "Could not find a rendered References heading in the PDF.",
                )
                rendered_heading = None
                after_reference_lines: List[Dict[str, Any]] = []
            else:
                rendered_heading = lines[ref_line_index]
                after_reference_lines = lines[ref_line_index + 1 :]
                add_check(
                    "references_heading_rendered",
                    "pass",
                    "Found rendered References heading.",
                    {
                        "page_number": rendered_heading["page_number"],
                        "top": round(rendered_heading["top"], 3),
                        "bbox": rendered_heading["bbox"],
                    },
                )

            heading_margins = self._section_margin_points(
                sections, expectations.get("references_heading_section_index")
            )
            top_margin_pt = heading_margins["top"]

            if rendered_heading is not None and expectations.get(
                "references_heading_expected_page_break"
            ):
                same_page_before = [
                    {
                        "text": line["text"][:100],
                        "top": round(line["top"], 3),
                        "bbox": line["bbox"],
                    }
                    for line in lines[:ref_line_index]
                    if line.get("page_index") == rendered_heading.get("page_index")
                ]
                top_ok = rendered_heading["top"] <= top_margin_pt + 48.0
                add_check(
                    "references_page_break",
                    "pass" if top_ok and not same_page_before else "fail",
                    (
                        "References heading starts on a fresh rendered page."
                        if top_ok and not same_page_before
                        else "DOCX requests a page break before References, but rendered PDF has prior same-page text or a low heading position."
                    ),
                    {
                        "expected_page_break": True,
                        "rendered_heading": {
                            "page_number": rendered_heading["page_number"],
                            "top": round(rendered_heading["top"], 3),
                        },
                        "same_page_text_before_count": len(same_page_before),
                        "same_page_text_before_examples": same_page_before[:5],
                    },
                )
            else:
                add_check(
                    "references_page_break",
                    "pass",
                    "No explicit DOCX page break before References was expected, or heading was not rendered.",
                    {
                        "expected_page_break": expectations.get(
                            "references_heading_expected_page_break"
                        )
                    },
                )

            reference_paragraphs = expectations.get("reference_paragraphs", [])
            default_reference_margins = self._section_margin_points(
                sections,
                reference_paragraphs[0].get("section_index")
                if reference_paragraphs
                else expectations.get("references_heading_section_index"),
            )
            if not reference_paragraphs:
                add_check(
                    "expected_reference_paragraphs",
                    "warn",
                    "No DOCX reference paragraphs with text were found after the References heading.",
                )
            else:
                add_check(
                    "expected_reference_paragraphs",
                    "pass",
                    f"Found {len(reference_paragraphs)} DOCX reference paragraph(s) after the References heading.",
                )

            matches = []
            search_pos = (ref_line_index + 1) if ref_line_index is not None else 0
            for ref in reference_paragraphs:
                found = None
                for idx in range(search_pos, len(lines)):
                    if self._line_matches_reference_start(lines[idx]["text"], ref["text"]):
                        found = idx
                        break
                if found is None:
                    matches.append({"reference": ref, "line_index": None})
                    continue
                matches.append({"reference": ref, "line_index": found})
                search_pos = found + 1

            hanging_failures = []
            horizontal_shift_failures = []
            matched_checks = []
            line_expectations: Dict[int, Dict[str, Any]] = {}
            found_matches = [
                match for match in matches if match.get("line_index") is not None
            ]
            for match_index, match in enumerate(found_matches):
                idx = match["line_index"]
                ref = match["reference"]
                first_line = lines[idx]
                next_start = (
                    found_matches[match_index + 1]["line_index"]
                    if match_index + 1 < len(found_matches)
                    else len(lines)
                )
                continuation_lines = lines[idx + 1 : next_start]
                ref_margins = self._section_margin_points(
                    sections, ref.get("section_index")
                )
                expected_right_bound = (
                    first_line["page_width"]
                    - ref_margins["right"]
                    - ref.get("expected_right_indent_pt", 0.0)
                )
                expected_first_left = (
                    ref_margins["left"] + ref["expected_first_line_offset_pt"]
                )
                expected_cont_left = (
                    ref_margins["left"] + ref["expected_continuation_offset_pt"]
                )
                check_entry = {
                    "paragraph_index": ref["paragraph_index"],
                    "section_index": ref.get("section_index"),
                    "section_left_margin_pt": round(ref_margins["left"], 3),
                    "section_right_margin_pt": round(ref_margins["right"], 3),
                    "first_line": {
                        "text": first_line["text"][:120],
                        "page_number": first_line["page_number"],
                        "left": round(first_line["left"], 3),
                        "right": round(first_line["right"], 3),
                    },
                    "expected_first_left": round(expected_first_left, 3),
                    "expected_continuation_left": round(expected_cont_left, 3),
                    "expected_hanging_indent_pt": ref["expected_hanging_indent_pt"],
                    "expected_right_indent_pt": ref.get("expected_right_indent_pt", 0.0),
                    "expected_right_bound": round(expected_right_bound, 3),
                }
                for line_idx in range(idx, next_start):
                    line_expectations[line_idx] = {
                        "reference": ref,
                        "margins": ref_margins,
                    }
                if abs(first_line["left"] - expected_first_left) > tolerance_points:
                    horizontal_shift_failures.append(check_entry)
                if continuation_lines and ref["expected_hanging_indent_pt"] > 0:
                    continuation_checks = []
                    bad_hanging = []
                    bad_continuation_left = []
                    for continuation_line in continuation_lines:
                        delta = continuation_line["left"] - first_line["left"]
                        continuation_entry = {
                            "text": continuation_line["text"][:120],
                            "page_number": continuation_line["page_number"],
                            "left": round(continuation_line["left"], 3),
                            "right": round(continuation_line["right"], 3),
                            "rendered_hanging_delta_pt": round(delta, 3),
                        }
                        continuation_checks.append(continuation_entry)
                        if (
                            abs(delta - ref["expected_hanging_indent_pt"])
                            > tolerance_points
                        ):
                            bad_hanging.append(continuation_entry)
                        if (
                            abs(continuation_line["left"] - expected_cont_left)
                            > tolerance_points
                        ):
                            bad_continuation_left.append(continuation_entry)
                    check_entry["continuation_lines"] = continuation_checks
                    if bad_hanging:
                        failure = dict(check_entry)
                        failure["bad_continuation_lines"] = bad_hanging
                        hanging_failures.append(failure)
                    if bad_continuation_left:
                        failure = dict(check_entry)
                        failure["bad_continuation_lines"] = bad_continuation_left
                        horizontal_shift_failures.append(failure)
                matched_checks.append(check_entry)

            unmatched = [
                {
                    "paragraph_index": match["reference"]["paragraph_index"],
                    "text_preview": match["reference"]["text_preview"],
                }
                for match in matches
                if match.get("line_index") is None
            ]
            hanging_status = "pass"
            hanging_message = "Rendered continuation lines preserve expected hanging indents."
            if hanging_failures:
                hanging_status = "fail"
                hanging_message = (
                    f"{len(hanging_failures)} reference paragraph(s) do not preserve the expected hanging indent in the rendered PDF."
                )
            elif unmatched:
                hanging_status = "warn"
                hanging_message = (
                    f"Could not match {len(unmatched)} expected reference paragraph(s) to rendered PDF lines."
                )
            add_check(
                "hanging_indents_rendered",
                hanging_status,
                hanging_message,
                {
                    "matched_reference_count": len(found_matches),
                    "unmatched_references": unmatched[:10],
                    "failures": hanging_failures[:10],
                    "examples": matched_checks[:10],
                },
            )

            add_check(
                "horizontal_alignment",
                "pass" if not horizontal_shift_failures else "fail",
                (
                    "Reference line starts align with expected DOCX margins and indents."
                    if not horizontal_shift_failures
                    else f"{len(horizontal_shift_failures)} reference line start(s) are shifted from expected DOCX margins/indents."
                ),
                {"failures": horizontal_shift_failures[:10]},
            )

            margin_violations = []
            margin_lines_checked = 0
            start_idx = (ref_line_index + 1) if ref_line_index is not None else 0
            for line_idx in range(start_idx, len(lines)):
                line = lines[line_idx]
                line_mapping = line_expectations.get(line_idx, {})
                margins = line_mapping.get("margins", default_reference_margins)
                ref = line_mapping.get("reference", {})
                left_margin_pt = margins["left"]
                right_margin_pt = margins["right"]
                top_margin_pt = margins["top"]
                bottom_margin_pt = margins["bottom"]
                expected_right = (
                    line["page_width"]
                    - right_margin_pt
                    - ref.get("expected_right_indent_pt", 0.0)
                )
                expected_bottom = line["page_height"] - bottom_margin_pt
                margin_lines_checked += 1
                if line["right"] > expected_right + tolerance_points:
                    margin_violations.append(
                        {
                            "section_index": margins["section_index"],
                            "paragraph_index": ref.get("paragraph_index"),
                            "page_number": line["page_number"],
                            "text": line["text"][:120],
                            "right": round(line["right"], 3),
                            "expected_right": round(expected_right, 3),
                            "over_by": round(line["right"] - expected_right, 3),
                        }
                    )
                if line["left"] < left_margin_pt - tolerance_points:
                    margin_violations.append(
                        {
                            "section_index": margins["section_index"],
                            "paragraph_index": ref.get("paragraph_index"),
                            "page_number": line["page_number"],
                            "text": line["text"][:120],
                            "left": round(line["left"], 3),
                            "expected_left": round(left_margin_pt, 3),
                            "under_by": round(left_margin_pt - line["left"], 3),
                        }
                    )
                if line["top"] < top_margin_pt - tolerance_points:
                    margin_violations.append(
                        {
                            "section_index": margins["section_index"],
                            "paragraph_index": ref.get("paragraph_index"),
                            "page_number": line["page_number"],
                            "text": line["text"][:120],
                            "top": round(line["top"], 3),
                            "expected_top": round(top_margin_pt, 3),
                            "under_by": round(top_margin_pt - line["top"], 3),
                        }
                    )
                if line["bottom"] > expected_bottom + tolerance_points:
                    margin_violations.append(
                        {
                            "section_index": margins["section_index"],
                            "paragraph_index": ref.get("paragraph_index"),
                            "page_number": line["page_number"],
                            "text": line["text"][:120],
                            "bottom": round(line["bottom"], 3),
                            "expected_bottom": round(expected_bottom, 3),
                            "over_by": round(line["bottom"] - expected_bottom, 3),
                        }
                    )

            add_check(
                "rendered_margins",
                "pass" if not margin_violations else "fail",
                (
                    "Rendered reference text stays within expected page margins."
                    if not margin_violations
                    else f"{len(margin_violations)} rendered reference line(s) exceed expected margins."
                ),
                {
                    "left_margin_pt": round(default_reference_margins["left"], 3),
                    "right_margin_pt": round(default_reference_margins["right"], 3),
                    "top_margin_pt": round(default_reference_margins["top"], 3),
                    "bottom_margin_pt": round(default_reference_margins["bottom"], 3),
                    "section_index": default_reference_margins["section_index"],
                    "lines_checked": margin_lines_checked,
                    "section_specific_mapped_line_count": len(line_expectations),
                    "violations": margin_violations[:20],
                },
            )

            return finalize(rendered_heading, reference_paragraphs)
        except Exception as e:
            return {"success": False, "error": str(e)}
        finally:
            if temp_pdf:
                try:
                    os.unlink(temp_pdf)
                except OSError:
                    pass

    def sanitize_document(
        self,
        file_path: str,
        *,
        remove_comments: bool = False,
        accept_revisions: bool = False,
        clear_metadata: bool = False,
        remove_custom_xml: bool = False,
        set_remove_personal_information: bool = False,
        canonicalize_ooxml: bool = False,
        author: Optional[str] = None,
        title: Optional[str] = None,
        subject: Optional[str] = None,
        keywords: Optional[str] = None,
        output_path: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Sanitize a DOCX for submission by removing hidden data and normalising metadata."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        if not any(
            [
                remove_comments,
                accept_revisions,
                clear_metadata,
                remove_custom_xml,
                set_remove_personal_information,
                canonicalize_ooxml,
                author is not None,
                title is not None,
                subject is not None,
                keywords is not None,
            ]
        ):
            return {"success": False, "error": "No sanitization options provided"}
        try:
            abs_path = str(Path(file_path).resolve())
            out_path = str(Path(output_path).resolve()) if output_path else abs_path
            overwrite = out_path == abs_path
            with self.host._file_locks(abs_path, out_path):
                backup_target = abs_path if overwrite else out_path
                backup = (
                    self.host._snapshot_backup(backup_target)
                    if Path(backup_target).exists()
                    else None
                )
                before = self.inspect_hidden_data(abs_path)
                if not before.get("success"):
                    return before

                files, package_preflight, preflight_error = self._read_docx_zip_parts(
                    abs_path
                )
                if preflight_error:
                    return preflight_error

                stats: Dict[str, int] = {}
                canonicalization = None

                if remove_comments or accept_revisions:
                    for part_name in self.host._docx_story_xml_parts(list(files.keys())):
                        root = ET.fromstring(files[part_name])
                        self.host._rewrite_story_tree(
                            root,
                            remove_comment_nodes=remove_comments,
                            accept_revisions=accept_revisions,
                            stats=stats,
                        )
                        if accept_revisions:
                            self._drop_revision_property_changes(root, stats)
                        files[part_name] = serialize_ooxml_element(root)

                if remove_comments:
                    removed_parts = []
                    comment_relationship_parts = sorted(
                        name
                        for name in files
                        if (
                            name.lower().startswith("word/_rels/comments")
                            and name.lower().endswith(".xml.rels")
                        )
                        or name.lower() == "word/_rels/people.xml.rels"
                    )
                    comment_related_targets = []
                    for rel_part in comment_relationship_parts:
                        comment_related_targets.extend(
                            self._docx_relationship_internal_targets(
                                rel_part, files[rel_part]
                            )
                        )
                    comment_related_targets = sorted(set(comment_related_targets))
                    for part_name in sorted(
                        name
                        for name in files
                        if (
                            name.lower().startswith("word/comments")
                            and name.lower().endswith(".xml")
                        )
                        or name.lower() == "word/people.xml"
                    ):
                        if part_name in files:
                            removed_parts.append(part_name)
                            del files[part_name]
                    removed_relationship_parts = []
                    for part_name in comment_relationship_parts:
                        if part_name in files:
                            removed_relationship_parts.append(part_name)
                            del files[part_name]

                    self.host._strip_docx_relationship_targets(
                        files,
                        lambda target, rel_type: "comments" in rel_type
                        or "comments" in target.lower()
                        or target.lower().endswith("people.xml"),
                    )
                    self.host._strip_docx_content_types(
                        files,
                        lambda child: "comments" in child.get("PartName", "").lower()
                        or child.get("PartName", "").lower().endswith("/people.xml"),
                    )
                    remaining_targets = set(
                        self._docx_all_relationship_internal_targets(files)
                    )
                    removed_related_targets = []
                    for target in comment_related_targets:
                        if (
                            target in files
                            and target not in remaining_targets
                            and self._is_comment_resource_part(target)
                        ):
                            removed_related_targets.append(target)
                            del files[target]
                    if removed_parts:
                        stats["comment_parts_removed"] = len(removed_parts)
                    if removed_relationship_parts:
                        stats["comment_relationship_parts_removed"] = len(
                            removed_relationship_parts
                        )
                    if removed_related_targets:
                        stats["comment_related_targets_removed"] = len(
                            removed_related_targets
                        )

                if remove_custom_xml:
                    custom_xml = self.host._custom_xml_summary(list(files.keys()))
                    removed_custom = list(custom_xml["all_parts"])
                    for name in removed_custom:
                        del files[name]
                    self.host._strip_docx_relationship_targets(
                        files,
                        lambda target, rel_type: target.startswith("customXml/")
                        or "/customXml" in rel_type,
                    )
                    self.host._strip_docx_content_types(
                        files,
                        lambda child: child.get("PartName", "").startswith("/customXml/"),
                    )
                    stats["custom_xml_items_removed"] = len(custom_xml["payload_parts"])
                    stats["custom_xml_support_parts_removed"] = len(
                        custom_xml["support_parts"]
                    )
                    stats["custom_xml_parts_removed"] = len(removed_custom)

                if clear_metadata or any(
                    opt is not None for opt in [author, title, subject, keywords]
                ):
                    core_name = "docProps/core.xml"
                    if core_name in files:
                        root = ET.fromstring(files[core_name])
                        field_map = {
                            "title": ("dc", "title", title),
                            "subject": ("dc", "subject", subject),
                            "creator": ("dc", "creator", author),
                            "keywords": ("cp", "keywords", keywords),
                            "description": ("dc", "description", ""),
                            "category": ("cp", "category", ""),
                            "contentStatus": ("cp", "contentStatus", ""),
                            "created": ("dcterms", "created", ""),
                            "modified": ("dcterms", "modified", ""),
                            "identifier": ("dc", "identifier", ""),
                            "language": ("dc", "language", ""),
                            "lastModifiedBy": ("cp", "lastModifiedBy", ""),
                            "version": ("cp", "version", ""),
                        }
                        for _, (prefix, local, explicit_value) in field_map.items():
                            node = root.find(f"{prefix}:{local}", OOXML_NS)
                            if node is None:
                                continue
                            if explicit_value is not None:
                                node.text = explicit_value
                            elif clear_metadata:
                                node.text = ""
                        revision = root.find("cp:revision", OOXML_NS)
                        if revision is not None and clear_metadata:
                            revision.text = "1"
                        files[core_name] = serialize_ooxml_element(root)

                    app_name = "docProps/app.xml"
                    if app_name in files and clear_metadata:
                        root = ET.fromstring(files[app_name])
                        for tag in [
                            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Manager",
                            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}Company",
                            "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}HyperlinkBase",
                        ]:
                            node = root.find(tag)
                            if node is not None:
                                node.text = ""
                        files[app_name] = serialize_ooxml_element(root)

                    custom_props_name = "docProps/custom.xml"
                    if clear_metadata and custom_props_name in files:
                        del files[custom_props_name]
                        stats["custom_document_properties_removed"] = 1
                        self.host._strip_docx_relationship_targets(
                            files,
                            lambda target, rel_type: target.lower().endswith(
                                "docprops/custom.xml"
                            )
                            or target.lower().endswith("custom.xml")
                            or "custom-properties" in rel_type.lower(),
                        )
                        self.host._strip_docx_content_types(
                            files,
                            lambda child: child.get("PartName", "").lower()
                            == "/docprops/custom.xml",
                        )

                if set_remove_personal_information or clear_metadata:
                    settings_name = "word/settings.xml"
                    if settings_name in files:
                        root = ET.fromstring(files[settings_name])
                        flag = root.find("w:removePersonalInformation", OOXML_NS)
                        if flag is None:
                            flag = ET.Element(
                                f"{{{OOXML_NS['w']}}}removePersonalInformation"
                            )
                            flag.set(f"{{{OOXML_NS['w']}}}val", "true")
                            insert_at = 1 if len(root) >= 1 else 0
                            root.insert(insert_at, flag)
                        else:
                            flag.set(f"{{{OOXML_NS['w']}}}val", "true")
                        files[settings_name] = serialize_ooxml_element(root)

                if canonicalize_ooxml:
                    canonicalization = self._canonicalize_wordprocessingml_parts(files)
                    stats["ooxml_parts_seen"] = canonicalization["parts_seen"]
                    stats["ooxml_parts_rewritten"] = canonicalization["parts_rewritten"]
                    stats["ooxml_parts_skipped"] = canonicalization["parts_skipped"]
                    if canonicalization.get("critical_skipped"):
                        return {
                            "success": False,
                            "error": (
                                "Could not safely canonicalize word/document.xml; "
                                "repair/regeneration is required before submission-ready claims."
                            ),
                            "canonicalization": canonicalization,
                            "stats": stats,
                            "backup": backup or None,
                        }

                self.host._atomic_zip_write(files, out_path)

                after = self.inspect_hidden_data(out_path)
                if not after.get("success"):
                    return after

            return {
                "success": True,
                "file": out_path,
                "input_file": abs_path,
                "output_file": out_path,
                "remove_comments": remove_comments,
                "accept_revisions": accept_revisions,
                "clear_metadata": clear_metadata,
                "remove_custom_xml": remove_custom_xml,
                "set_remove_personal_information": set_remove_personal_information
                or clear_metadata,
                "canonicalize_ooxml": canonicalize_ooxml,
                "canonicalization": canonicalization,
                "stats": stats,
                "before": before,
                "after": after,
                "ooxml_preflight": package_preflight,
                "backup": backup or None,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def extract_images_from_docx(
        self,
        file_path: str,
        output_dir: str,
        fmt: str = "png",
        prefix: str = "image",
    ) -> Dict[str, Any]:
        """Extract all embedded images from a .docx file."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            fmt_info, fmt_error = self._normalize_image_format(fmt)
            if fmt_error:
                return fmt_error
            fmt_ext, pillow_format = fmt_info
            safe_prefix, prefix_preflight, prefix_error = self._validate_image_output_prefix(prefix)
            if prefix_error:
                return prefix_error
            from PIL import Image as PILImage

            output_root = self._prepare_output_dir(output_dir)
            doc = Document(file_path)
            extracted = []
            idx = 0
            truncated = False
            warnings_list = []
            resource_limits = self._image_extract_resource_limits()
            for image_entry in self._docx_related_image_parts(doc):
                if idx >= self.MAX_EXTRACT_IMAGES:
                    truncated = True
                    warning = f"Stopped after {self.MAX_EXTRACT_IMAGES} images"
                    if warning not in warnings_list:
                        warnings_list.append(warning)
                    break
                try:
                    image_part = image_entry["part"]
                    image_bytes = image_part.blob
                    content_type = image_part.content_type
                    ext_map = {
                        "image/png": "png",
                        "image/jpeg": "jpg",
                        "image/gif": "gif",
                        "image/bmp": "bmp",
                        "image/tiff": "tiff",
                        "image/x-emf": "emf",
                        "image/x-wmf": "wmf",
                    }
                    src_ext = ext_map.get(content_type, "png")
                    if content_type in self.VECTOR_IMAGE_CONTENT_TYPES:
                        if len(image_bytes) > self.MAX_EXTRACT_IMAGE_COMPRESSED_BYTES:
                            extracted.append(
                                {
                                    "index": idx,
                                    "skipped": True,
                                    "source_part": image_entry["source_part"],
                                    "partname": image_entry["partname"],
                                    "error": (
                                        f"Embedded vector image is {len(image_bytes)} bytes, exceeding "
                                        f"the safe limit {self.MAX_EXTRACT_IMAGE_COMPRESSED_BYTES}"
                                    ),
                                    "error_code": "image_compressed_bytes_limit_exceeded",
                                    "original_format": src_ext,
                                    "compressed_size_bytes": len(image_bytes),
                                    "max_compressed_size_bytes": self.MAX_EXTRACT_IMAGE_COMPRESSED_BYTES,
                                }
                            )
                        else:
                            raw_name = f"{safe_prefix}_{idx:03d}.{src_ext}"
                            raw_path = self._safe_child_path(output_root, raw_name)
                            self._atomic_write_bytes(raw_path, image_bytes)
                            extracted.append(
                                {
                                    "index": idx,
                                    "file": str(raw_path),
                                    "source_part": image_entry["source_part"],
                                    "partname": image_entry["partname"],
                                    "format": src_ext,
                                    "original_format": src_ext,
                                    "compressed_size_bytes": len(image_bytes),
                                    "size_bytes": os.path.getsize(str(raw_path)),
                                    "note": f"Vector format saved as .{src_ext} (cannot convert to {fmt_ext})",
                                }
                            )
                    else:
                        out_name = f"{safe_prefix}_{idx:03d}.{fmt_ext}"
                        out_path = self._safe_child_path(output_root, out_name)
                        img, image_meta, skip_entry = self._load_bounded_image(
                            PILImage,
                            image_bytes,
                        )
                        if skip_entry:
                            skip_entry.update(
                                {
                                    "index": idx,
                                    "original_format": src_ext,
                                    "source_part": image_entry["source_part"],
                                    "partname": image_entry["partname"],
                                }
                            )
                            extracted.append(skip_entry)
                        else:
                            try:
                                self._save_bounded_image(
                                    img,
                                    out_path,
                                    fmt_ext,
                                    pillow_format,
                                )
                            finally:
                                img.close()
                            extracted.append(
                                {
                                    "index": idx,
                                    "file": str(out_path),
                                    "source_part": image_entry["source_part"],
                                    "partname": image_entry["partname"],
                                    "format": fmt_ext,
                                    "original_format": src_ext,
                                    "width": image_meta["width"],
                                    "height": image_meta["height"],
                                    "pixels": image_meta["pixels"],
                                    "compressed_size_bytes": image_meta["compressed_size_bytes"],
                                    "size_bytes": os.path.getsize(str(out_path)),
                                }
                            )
                    idx += 1
                except Exception as img_err:
                    extracted.append(
                        {
                            "index": idx,
                            "skipped": True,
                            "source_part": image_entry.get("source_part"),
                            "partname": image_entry.get("partname"),
                            "error": str(img_err),
                            "error_code": "image_extract_failed",
                        }
                    )
                    idx += 1
            return {
                "success": True,
                "file": file_path,
                "output_dir": str(output_root),
                "prefix": safe_prefix,
                "images_extracted": len([entry for entry in extracted if "file" in entry]),
                "images_skipped": len([entry for entry in extracted if entry.get("skipped")]),
                "images": extracted,
                "truncated": truncated,
                "warnings": warnings_list,
                "resource_limits": resource_limits,
                "preflight": prefix_preflight,
            }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def doc_render_map(self, file_path: str) -> Dict[str, Any]:
        """Build deterministic paragraph/table-cell to rendered PDF page/bbox mappings."""
        if not DOCX_AVAILABLE:
            return {"success": False, "error": "python-docx not installed"}
        try:
            abs_path = str(Path(file_path).resolve())
            if not os.path.exists(abs_path):
                return {"success": False, "error": f"File not found: {file_path}"}

            with tempfile.TemporaryDirectory(prefix="onlyoffice-doc-render-map-") as tmpdir:
                temp_docx = os.path.join(tmpdir, Path(abs_path).name)
                temp_pdf = os.path.join(tmpdir, f"{Path(abs_path).stem}.pdf")
                shutil.copy2(abs_path, temp_docx)

                doc = Document(temp_docx)
                paragraph_anchors = []
                table_cell_anchors = []

                paragraph_index = 0
                for paragraph in doc.paragraphs:
                    original_text = re.sub(r"\s+", " ", paragraph.text or "").strip()
                    if not original_text:
                        continue
                    paragraph_index += 1
                    anchor_id = f"SLOANE_P_{paragraph_index:04d}"
                    paragraph.text = f"{anchor_id} {original_text}"
                    paragraph_anchors.append(
                        {
                            "anchor_id": anchor_id,
                            "paragraph_index": paragraph_index,
                            "text_preview": original_text[:220],
                        }
                    )

                for table_number, table in enumerate(doc.tables, start=1):
                    for row_number, row in enumerate(table.rows, start=1):
                        for column_number, cell in enumerate(row.cells, start=1):
                            cell_text = re.sub(
                                r"\s+",
                                " ",
                                " ".join(
                                    paragraph.text or "" for paragraph in cell.paragraphs
                                ),
                            ).strip()
                            if not cell_text:
                                continue
                            anchor_id = (
                                f"SLOANE_T{table_number}R{row_number}C{column_number}"
                            )
                            cell.text = f"{anchor_id} {cell_text}"
                            table_cell_anchors.append(
                                {
                                    "anchor_id": anchor_id,
                                    "table_number": table_number,
                                    "row_number": row_number,
                                    "column_number": column_number,
                                    "cell_ref": f"T{table_number}R{row_number}C{column_number}",
                                    "text_preview": cell_text[:220],
                                }
                            )

                doc.save(temp_docx)

                pdf_result = self.host.doc_to_pdf(temp_docx, output_path=temp_pdf)
                if not pdf_result.get("success"):
                    return pdf_result

                block_payload = self.host.pdf_read_blocks(
                    temp_pdf,
                    include_spans=True,
                    include_images=False,
                    include_empty=False,
                )
                if not block_payload.get("success"):
                    return block_payload

                marker_hits = {}
                marker_pattern = re.compile(r"SLOANE_(?:P_\d{4}|T\d+R\d+C\d+)")

                for page in block_payload.get("pages", []):
                    for block in page.get("blocks", []):
                        if block.get("type") != "text":
                            continue
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                span_text = str(span.get("text", "") or "")
                                for match in marker_pattern.finditer(span_text):
                                    anchor_id = match.group(0)
                                    if anchor_id in marker_hits:
                                        continue
                                    marker_hits[anchor_id] = {
                                        "page_index": page.get("page_index"),
                                        "page_number": page.get("page_number"),
                                        "block_id": block.get("block_id"),
                                        "line_id": line.get("line_id"),
                                        "span_id": span.get("span_id"),
                                        "bbox": span.get("bbox") or block.get("bbox"),
                                        "block_bbox": block.get("bbox"),
                                        "text": block.get("text"),
                                    }

                paragraph_mappings = []
                table_cell_mappings = []
                unresolved_anchor_ids = []

                for anchor in paragraph_anchors:
                    hit = marker_hits.get(anchor["anchor_id"])
                    if not hit:
                        unresolved_anchor_ids.append(anchor["anchor_id"])
                        continue
                    paragraph_mappings.append({**anchor, **hit})

                for anchor in table_cell_anchors:
                    hit = marker_hits.get(anchor["anchor_id"])
                    if not hit:
                        unresolved_anchor_ids.append(anchor["anchor_id"])
                        continue
                    table_cell_mappings.append({**anchor, **hit})

                return {
                    "success": True,
                    "file": abs_path,
                    "paragraph_count": len(paragraph_anchors),
                    "table_cell_count": len(table_cell_anchors),
                    "mapped_paragraph_count": len(paragraph_mappings),
                    "mapped_table_cell_count": len(table_cell_mappings),
                    "pages": block_payload.get("pages_scanned", 0),
                    "paragraphs": paragraph_mappings,
                    "table_cells": table_cell_mappings,
                    "unresolved_anchor_ids": unresolved_anchor_ids,
                }
        except Exception as e:
            return {"success": False, "error": str(e)}

    def doc_to_pdf(
        self,
        file_path: str,
        output_path: str = None,
        *,
        layout_warnings: bool = False,
        render_profile: str = "auto",
    ) -> Dict[str, Any]:
        """Convert a .docx file to PDF via OnlyOffice x2t."""
        result = self.host._office_to_pdf(file_path, output_path=output_path)
        if not result.get("success") or not layout_warnings:
            return result

        audit = self.rendered_layout_audit(
            file_path,
            pdf_path=result.get("output_file"),
            profile=render_profile,
            trusted_pdf=True,
        )
        if not audit.get("success"):
            result["layout_audit_status"] = "error"
            result["layout_warnings"] = [
                {
                    "name": "rendered_layout_audit",
                    "status": "warn",
                    "message": audit.get("error", "Rendered layout audit failed."),
                    "details": audit,
                }
            ]
            result["layout_audit"] = audit
            result["submission_ready"] = False
            result["readiness_blockers"] = result["layout_warnings"]
            return result
        result["layout_audit_status"] = audit.get("overall_status")
        result["layout_warnings"] = [
            {
                "name": check.get("name"),
                "status": check.get("status"),
                "message": check.get("message"),
                "details": check.get("details", {}),
            }
            for check in audit.get("checks", [])
            if check.get("status") in {"warn", "fail"}
        ]
        result["layout_audit"] = audit
        result["submission_ready"] = audit.get("submission_ready") is True
        result["readiness_blockers"] = audit.get("readiness_blockers", [])
        return result

    def preview_document(
        self,
        file_path: str,
        output_dir: str,
        pages: str = None,
        dpi: int = 150,
        fmt: str = "png",
    ) -> Dict[str, Any]:
        """Render DOCX pages as images via OnlyOffice conversion + PyMuPDF."""
        return self.host._preview_via_pdf(
            file_path, output_dir, self.host.doc_to_pdf, pages=pages, dpi=dpi, fmt=fmt
        )
